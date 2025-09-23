# --- PARTE 1 ---
import pandas as pd
from pandas.tseries.offsets import DateOffset
import os
from datetime import datetime, timezone, timedelta
import tkinter as tk
from tkinter import messagebox, ttk, font as tkFont
from collections import Counter, defaultdict
import logging
import pyperclip # type: ignore
import subprocess
import sys
from functools import partial
import re
import shutil # Aggiunto per copiare i file master

from docx import Document # Per report Word
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

# Import per manipolazione file Excel
from openpyxl import load_workbook # Usato per .xlsx
# xlrd è usato da pandas come engine per .xls, quindi non serve import diretto se si usa pandas


# --- INIZIO CONFIGURAZIONE GLOBALE ---

# Variabili per i percorsi che saranno letti dal file parametri.xlsm
FILE_REGISTRO_STRUMENTI = None
FOLDER_PATH_DEFAULT = None
FILE_DATI_COMPILAZIONE_SCHEDE = None
FILE_MASTER_DIGITALE_XLSX = None # Nuovo: Percorso master per schede digitali
FILE_MASTER_ANALOGICO_XLSX = None # Nuovo: Percorso master per schede analogiche

try:
    # Determina la directory in cui si trova lo script Python
    try:
        SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    except NameError:
        SCRIPT_DIR = os.getcwd() 

    # Configurazione per leggere il file parametri.xlsm
    NOME_FILE_PARAMETRI = "parametri.xlsm"
    PATH_FILE_PARAMETRI = os.path.join(SCRIPT_DIR, NOME_FILE_PARAMETRI)
    NOME_FOGLIO_PARAMETRI = "parametri"

    print(f"INFO: Tentativo di lettura percorsi da: {PATH_FILE_PARAMETRI}")

    if not os.path.exists(PATH_FILE_PARAMETRI):
        critical_error_msg = f"ERRORE CRITICO: File parametri '{PATH_FILE_PARAMETRI}' non trovato. Impossibile configurare i percorsi."
        print(critical_error_msg, file=sys.stderr)
        raise FileNotFoundError(critical_error_msg)

    df_params = pd.read_excel(PATH_FILE_PARAMETRI, sheet_name=NOME_FOGLIO_PARAMETRI, header=None, engine='openpyxl', dtype=str)

    # Cella B2: FILE_REGISTRO_STRUMENTI
    try:
        path_registro_letto = df_params.iloc[1, 1] # Riga 2, Colonna B (0-indexed)
        if pd.isna(path_registro_letto) or str(path_registro_letto).strip() == "":
            error_msg_reg = f"ERRORE CRITICO: Cella B2 (FILE_REGISTRO_STRUMENTI) nel file parametri è vuota o non valida." # Reso più critico
            print(error_msg_reg, file=sys.stderr)
            raise ValueError(error_msg_reg)
        FILE_REGISTRO_STRUMENTI = str(path_registro_letto).strip()
        if not os.path.exists(FILE_REGISTRO_STRUMENTI):
            raise FileNotFoundError(f"File registro strumenti specificato in B2 non trovato: {FILE_REGISTRO_STRUMENTI}")
        print(f"INFO: FILE_REGISTRO_STRUMENTI impostato da parametri.xlsm: {FILE_REGISTRO_STRUMENTI}")
    except IndexError:
        error_msg_reg_idx = f"ERRORE CRITICO: Cella B2 non trovata nel foglio '{NOME_FOGLIO_PARAMETRI}'. Verificare '{PATH_FILE_PARAMETRI}'."
        print(error_msg_reg_idx, file=sys.stderr)
        raise ValueError(error_msg_reg_idx)

    # Cella B3: FOLDER_PATH_DEFAULT
    try:
        path_schede_letto = df_params.iloc[2, 1] # Riga 3, Colonna B
        if pd.isna(path_schede_letto) or str(path_schede_letto).strip() == "":
            error_msg_fld = f"ERRORE CRITICO: Cella B3 (FOLDER_PATH_DEFAULT) nel file parametri è vuota o non valida." # Reso più critico
            print(error_msg_fld, file=sys.stderr)
            raise ValueError(error_msg_fld)
        FOLDER_PATH_DEFAULT = str(path_schede_letto).strip()
        if not os.path.isdir(FOLDER_PATH_DEFAULT):
            raise NotADirectoryError(f"Cartella schede specificata in B3 non trovata o non è una cartella: {FOLDER_PATH_DEFAULT}")
        print(f"INFO: FOLDER_PATH_DEFAULT impostato da parametri.xlsm: {FOLDER_PATH_DEFAULT}")
    except IndexError:
        error_msg_fld_idx = f"ERRORE CRITICO: Cella B3 non trovata nel foglio '{NOME_FOGLIO_PARAMETRI}'. Verificare '{PATH_FILE_PARAMETRI}'."
        print(error_msg_fld_idx, file=sys.stderr)
        raise ValueError(error_msg_fld_idx)

    # Cella B4: FILE_DATI_COMPILAZIONE_SCHEDE (Opzionale)
    try:
        path_compilazione_letto = df_params.iloc[3, 1] # Riga 4, Colonna B
        if pd.isna(path_compilazione_letto) or str(path_compilazione_letto).strip() == "":
            warn_msg_comp = f"AVVISO: Cella B4 (FILE_DATI_COMPILAZIONE_SCHEDE) nel file parametri è vuota. La compilazione automatica dei campi anagrafici non sarà disponibile."
            print(warn_msg_comp, file=sys.stderr)
            FILE_DATI_COMPILAZIONE_SCHEDE = None
        else:
            FILE_DATI_COMPILAZIONE_SCHEDE = str(path_compilazione_letto).strip()
            if not os.path.exists(FILE_DATI_COMPILAZIONE_SCHEDE):
                print(f"AVVISO: File dati compilazione specificato in B4 non trovato: {FILE_DATI_COMPILAZIONE_SCHEDE}. La compilazione automatica dei campi anagrafici potrebbe non funzionare.", file=sys.stderr)
                FILE_DATI_COMPILAZIONE_SCHEDE = None # Tratta come non specificato se il file non esiste
            else:
                print(f"INFO: FILE_DATI_COMPILAZIONE_SCHEDE impostato da parametri.xlsm: {FILE_DATI_COMPILAZIONE_SCHEDE}")
    except IndexError:
        warn_msg_comp_idx = f"AVVISO: Cella B4 non trovata nel foglio '{NOME_FOGLIO_PARAMETRI}'. Il percorso per FILE_DATI_COMPILAZIONE_SCHEDE non è stato impostato."
        print(warn_msg_comp_idx, file=sys.stderr)
        FILE_DATI_COMPILAZIONE_SCHEDE = None

    # Cella B5: FILE_MASTER_DIGITALE_XLSX (Nuovo, Opzionale per conversione .xls)
    try:
        path_master_dig_letto = df_params.iloc[4, 1] # Riga 5, Colonna B
        if pd.isna(path_master_dig_letto) or str(path_master_dig_letto).strip() == "":
            print(f"AVVISO: Cella B5 (FILE_MASTER_DIGITALE_XLSX) nel file parametri è vuota. La conversione di schede digitali .xls tramite master non sarà possibile.", file=sys.stderr)
            FILE_MASTER_DIGITALE_XLSX = None
        else:
            FILE_MASTER_DIGITALE_XLSX = str(path_master_dig_letto).strip()
            if not os.path.exists(FILE_MASTER_DIGITALE_XLSX):
                print(f"AVVISO: File master digitale specificato in B5 non trovato: {FILE_MASTER_DIGITALE_XLSX}. La conversione di schede digitali .xls tramite master non sarà possibile.", file=sys.stderr)
                FILE_MASTER_DIGITALE_XLSX = None
            elif not FILE_MASTER_DIGITALE_XLSX.lower().endswith(".xlsx"):
                print(f"AVVISO: File master digitale specificato in B5 ({FILE_MASTER_DIGITALE_XLSX}) non è un file .xlsx. La conversione tramite master non sarà possibile.", file=sys.stderr)
                FILE_MASTER_DIGITALE_XLSX = None
            else:
                print(f"INFO: FILE_MASTER_DIGITALE_XLSX impostato da parametri.xlsm: {FILE_MASTER_DIGITALE_XLSX}")
    except IndexError:
        print(f"AVVISO: Cella B5 non trovata nel foglio '{NOME_FOGLIO_PARAMETRI}'. Il percorso per FILE_MASTER_DIGITALE_XLSX non è stato impostato.", file=sys.stderr)
        FILE_MASTER_DIGITALE_XLSX = None

    # Cella B6: FILE_MASTER_ANALOGICO_XLSX (Nuovo, Opzionale per conversione .xls)
    try:
        path_master_ana_letto = df_params.iloc[5, 1] # Riga 6, Colonna B
        if pd.isna(path_master_ana_letto) or str(path_master_ana_letto).strip() == "":
            print(f"AVVISO: Cella B6 (FILE_MASTER_ANALOGICO_XLSX) nel file parametri è vuota. La conversione di schede analogiche .xls tramite master non sarà possibile.", file=sys.stderr)
            FILE_MASTER_ANALOGICO_XLSX = None
        else:
            FILE_MASTER_ANALOGICO_XLSX = str(path_master_ana_letto).strip()
            if not os.path.exists(FILE_MASTER_ANALOGICO_XLSX):
                print(f"AVVISO: File master analogico specificato in B6 non trovato: {FILE_MASTER_ANALOGICO_XLSX}. La conversione di schede analogiche .xls tramite master non sarà possibile.", file=sys.stderr)
                FILE_MASTER_ANALOGICO_XLSX = None
            elif not FILE_MASTER_ANALOGICO_XLSX.lower().endswith(".xlsx"):
                print(f"AVVISO: File master analogico specificato in B6 ({FILE_MASTER_ANALOGICO_XLSX}) non è un file .xlsx. La conversione tramite master non sarà possibile.", file=sys.stderr)
                FILE_MASTER_ANALOGICO_XLSX = None
            else:
                print(f"INFO: FILE_MASTER_ANALOGICO_XLSX impostato da parametri.xlsm: {FILE_MASTER_ANALOGICO_XLSX}")
    except IndexError:
        print(f"AVVISO: Cella B6 non trovata nel foglio '{NOME_FOGLIO_PARAMETRI}'. Il percorso per FILE_MASTER_ANALOGICO_XLSX non è stato impostato.", file=sys.stderr)
        FILE_MASTER_ANALOGICO_XLSX = None

    # Validazione finale dei percorsi critici (Registro e Cartella Schede)
    if not FILE_REGISTRO_STRUMENTI or not FOLDER_PATH_DEFAULT:
        # Questo blocco è ridondante se gli errori sopra sono ValueError, ma lo lascio per sicurezza
        final_config_error = "ERRORE CRITICO FINALE: Uno o entrambi i percorsi FILE_REGISTRO_STRUMENTI o FOLDER_PATH_DEFAULT non sono stati configurati correttamente da parametri.xlsm."
        print(final_config_error, file=sys.stderr)
        raise ValueError(final_config_error)

except (FileNotFoundError, NotADirectoryError, ValueError) as e_config: # ValueError per errori di configurazione manuali
    print(f"ERRORE FATALE DURANTE CONFIGURAZIONE PERCORSI: {e_config}", file=sys.stderr)
    print("L'applicazione non può continuare. Correggere 'parametri.xlsm' e riavviare.", file=sys.stderr)
    # Qui si potrebbe voler uscire esplicitamente se questo codice è il punto di ingresso principale
    # In un contesto GUI, si solleva l'eccezione per essere gestita dal blocco try-except principale dell'app.
    raise # Rilancia l'eccezione per bloccare l'avvio se critico

LOG_FILENAME_TEMPLATE = "log_analisi_schede_{timestamp}.txt"
ANALYSIS_DATETIME = datetime.now(timezone.utc)

REGISTRO_COL_IDX_MODELLO_STRUM_CAMPIONE = 6
REGISTRO_COL_IDX_ID_CERT_CAMPIONE = 16
REGISTRO_COL_IDX_RANGE_CAMPIONE = 12
REGISTRO_COL_IDX_SCADENZA_CAMPIONE = 18
REGISTRO_RIGA_INIZIO_DATI = 7
REGISTRO_FOGLIO_NOME = "strumenti campione ISAB SUD"
SOGLIA_PER_SUGGERIMENTO_ALTERNATIVO = 5

def _determine_log_filepath():
    now_local = ANALYSIS_DATETIME.astimezone()
    timestamp_str = now_local.strftime("%Y%m%d_%H%M%S")
    log_filename = LOG_FILENAME_TEMPLATE.format(timestamp=timestamp_str)
    try:
        documents_folder = os.path.join(os.path.expanduser("~"), "Documents", "AnalisiSchedeLogs")
        if not os.path.exists(documents_folder):
            os.makedirs(documents_folder, exist_ok=True)
        filepath = os.path.join(documents_folder, log_filename)
        with open(filepath, "a", encoding="utf-8") as f_test: 
            f_test.write("") 
        return filepath
    except Exception:
        script_dir_for_log = SCRIPT_DIR 
        filepath = os.path.join(script_dir_for_log, log_filename)
        try:
            with open(filepath, "a", encoding="utf-8") as f_test: f_test.write("")
        except Exception:
            return None 
        return filepath

LOG_FILEPATH = _determine_log_filepath()

if LOG_FILEPATH:
    print(f"INFO: Il file di log sarà: {LOG_FILEPATH}")
else:
    print("ATTENZIONE: Percorso del file di log non determinato.", file=sys.stderr)

logger = logging.getLogger('SchedeAnalyzer')
if not logger.handlers: 
    logger.propagate = False 
    logger.setLevel(logging.DEBUG) 

    formatter_file = logging.Formatter('%(asctime)s - %(levelname)s - [%(funcName)s:%(lineno)d] - %(message)s')
    formatter_console = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')

    if LOG_FILEPATH:
        try:
            log_dir = os.path.dirname(LOG_FILEPATH)
            if not os.path.exists(log_dir):
                os.makedirs(log_dir, exist_ok=True)
            file_handler = logging.FileHandler(LOG_FILEPATH, encoding='utf-8', mode='w') 
            file_handler.setFormatter(formatter_file)
            logger.addHandler(file_handler)
        except Exception as e_fh_setup:
            print(f"ERRORE CRITICO: Impossibile creare FileHandler per il log in '{LOG_FILEPATH}'. Errore: {e_fh_setup}. I log andranno solo su console.", file=sys.stderr)
    else:
        print("ERRORE CRITICO: Percorso del file di log non determinato. I log andranno solo su console.", file=sys.stderr)
    
    console_handler_main = logging.StreamHandler(sys.stdout) 
    console_handler_main.setFormatter(formatter_console)
    console_handler_main.setLevel(logging.INFO) 
    logger.addHandler(console_handler_main)
    
    if LOG_FILEPATH and any(isinstance(h, logging.FileHandler) for h in logger.handlers):
        logger.info(f"Logger 'SchedeAnalyzer' configurato. Log su file: {LOG_FILEPATH}")
    else:
        logger.warning("Logger 'SchedeAnalyzer' configurato solo per output su console.")

if FOLDER_PATH_DEFAULT:
    logger.info(f"Cartella schede da analizzare (da parametri.xlsm): {FOLDER_PATH_DEFAULT}")
if FILE_REGISTRO_STRUMENTI:
    logger.info(f"File registro strumenti (da parametri.xlsm): {FILE_REGISTRO_STRUMENTI}")
if FILE_DATI_COMPILAZIONE_SCHEDE:
    logger.info(f"File dati compilazione schede (da parametri.xlsm): {FILE_DATI_COMPILAZIONE_SCHEDE}")
else:
    logger.warning("File dati compilazione schede (B4) non specificato in parametri.xlsm. La compilazione automatica dei campi anagrafici non sarà disponibile.")

if FILE_MASTER_DIGITALE_XLSX:
    logger.info(f"File master digitale .xlsx (B5): {FILE_MASTER_DIGITALE_XLSX}")
else:
    logger.warning("File master digitale .xlsx (B5) non specificato o non valido. Conversione .xls digitali tramite master non possibile.")
if FILE_MASTER_ANALOGICO_XLSX:
    logger.info(f"File master analogico .xlsx (B6): {FILE_MASTER_ANALOGICO_XLSX}")
else:
    logger.warning("File master analogico .xlsx (B6) non specificato o non valido. Conversione .xls analogici tramite master non possibile.")


SCHEDA_DIG_CELL_TIPOLOGIA_STRUM = "N10"
SCHEDA_DIG_CELL_RANGE_UM_PROCESSO = "D22"
SCHEDA_ANA_CELL_TIPOLOGIA_STRUM = "N9"
SCHEDA_ANA_CELL_MODELLO_STRUM = "L9"
SCHEDA_ANA_CELL_RANGE_INGRESSO = "A9"; SCHEDA_ANA_CELL_UM_INGRESSO = "C9"
SCHEDA_ANA_CELL_RANGE_USCITA = "D12"; SCHEDA_ANA_CELL_UM_USCITA = "F12"
SCHEDA_ANA_CELL_RANGE_DCS = "D9"; SCHEDA_ANA_CELL_UM_DCS = "F9"
SCHEDA_ANA_CELL_ODC = "L50"
SCHEDA_ANA_CELL_DATA_COMPILAZIONE = "B50"
SCHEDA_ANA_CELL_PDL = "F50"
SCHEDA_ANA_CELL_ESECUTORE = "F52"
SCHEDA_ANA_CELL_SUPERVISORE_ISAB = "L52"
SCHEDA_ANA_CELL_CONTRATTO_COEMI = "B52"
SCHEDA_DIG_CELL_ODC = "L45"
SCHEDA_DIG_CELL_DATA_COMPILAZIONE = "B45"
SCHEDA_DIG_CELL_PDL = "F45"
SCHEDA_DIG_CELL_ESECUTORE = "F47"
SCHEDA_DIG_CELL_SUPERVISORE_ISAB = "L47"
SCHEDA_DIG_CELL_CONTRATTO_COEMI = "B47"
VALORE_ATTESO_CONTRATTO_COEMI = "COEMI 4600002254"
VALORE_ATTESO_CONTRATTO_COEMI_VARIANTE_NUMERICA = "4600002254"
NOME_FOGLIO_DATI_COMPILAZIONE = "RIASSUNTO"
COL_IDX_COMP_DATA = 0      
COL_IDX_COMP_ESECUTORE = 1  
COL_IDX_COMP_SUPERVISORE = 3
COL_IDX_COMP_ODC = 4        
COL_IDX_COMP_PDL = 5        
KEY_SP_VUOTO = "SP_VUOTO_O_ILLEGGIBILE"
KEY_L9_VUOTO = "L9_VUOTO_O_ILLEGGIBILE"
KEY_L9_SKINPOINT_INCOMPLETO = "L9_SKINPOINT_INCOMPLETO"
KEY_CELL_RANGE_UM_NON_LEGGIBILE = "RANGE_UM_CELL_ILLEGGIBILE"
KEY_ERR_ANA_TEMP_CONV_C9F9_UM_DIVERSE = "ERR_ANA_TEMP_CONV_C9F9_UM_DIVERSE"
KEY_ERR_ANA_TEMP_CONV_F12_UM_NON_MA = "ERR_ANA_TEMP_CONV_F12_UM_NON_MA"
KEY_ERR_ANA_TEMP_CONV_A9D9_RANGE_DIVERSI = "ERR_ANA_TEMP_CONV_A9D9_RANGE_DIVERSI"
KEY_ERR_ANA_TEMP_CONV_D12_RANGE_NON_4_20 = "ERR_ANA_TEMP_CONV_D12_RANGE_NON_4_20"
KEY_ERR_ANA_TEMP_NOCONV_UM_NON_COINCIDENTI = "ERR_ANA_TEMP_NOCONV_UM_NON_COINCIDENTI"
KEY_ERR_ANA_TEMP_NOCONV_RANGE_NON_COINCIDENTI = "ERR_ANA_TEMP_NOCONV_RANGE_NON_COINCIDENTI"
KEY_ERR_ANA_LIVELLO_DP_C9_UM_NON_PRESSIONE = "ERR_ANA_LIVELLO_DP_C9_UM_NON_PRESSIONE"
KEY_ERR_ANA_LIVELLO_DP_D9_RANGE_NON_0_100 = "ERR_ANA_LIVELLO_DP_D9_RANGE_NON_0_100"
KEY_ERR_ANA_LIVELLO_DP_F9_UM_NON_PERCENTO = "ERR_ANA_LIVELLO_DP_F9_UM_NON_PERCENTO"
KEY_ERR_ANA_LIVELLO_DP_D12_RANGE_NON_4_20 = "ERR_ANA_LIVELLO_DP_D12_RANGE_NON_4_20"
KEY_ERR_ANA_LIVELLO_DP_F12_UM_NON_MA = "ERR_ANA_LIVELLO_DP_F12_UM_NON_MA"
KEY_ERR_ANA_LIVELLO_TORS_C9_UM_INVALIDA = "ERR_ANA_LIVELLO_TORS_C9_UM_INVALIDA" 
KEY_ERR_ANA_LIVELLO_TORS_D9_RANGE_NON_0_100 = "ERR_ANA_LIVELLO_TORS_D9_RANGE_NON_0_100"
KEY_ERR_ANA_LIVELLO_TORS_F9_UM_NON_PERCENTO = "ERR_ANA_LIVELLO_TORS_F9_UM_NON_PERCENTO"
KEY_ERR_ANA_LIVELLO_TORS_ELETTR_D12_RANGE_NON_4_20 = "ERR_ANA_LIVELLO_TORS_ELETTR_D12_RANGE_NON_4_20"
KEY_ERR_ANA_LIVELLO_TORS_ELETTR_F12_UM_NON_MA = "ERR_ANA_LIVELLO_TORS_ELETTR_F12_UM_NON_MA"
KEY_ERR_ANA_LIVELLO_TORS_LOCALE_D12_RANGE_NON_VUOTO = "ERR_ANA_LIVELLO_TORS_LOCALE_D12_RANGE_NON_VUOTO" 
KEY_ERR_ANA_LIVELLO_TORS_LOCALE_F12_UM_NON_VUOTA = "ERR_ANA_LIVELLO_TORS_LOCALE_F12_UM_NON_VUOTA" 
KEY_ERR_ANA_LIVELLO_RADARULTR_D9_RANGE_NON_0_100 = "ERR_ANA_LIVELLO_RADARULTR_D9_RANGE_NON_0_100"
KEY_ERR_ANA_LIVELLO_RADARULTR_F9_UM_NON_PERCENTO = "ERR_ANA_LIVELLO_RADARULTR_F9_UM_NON_PERCENTO"
KEY_ERR_ANA_LIVELLO_RADARULTR_D12_RANGE_NON_4_20 = "ERR_ANA_LIVELLO_RADARULTR_D12_RANGE_NON_4_20"
KEY_ERR_ANA_LIVELLO_RADARULTR_F12_UM_NON_MA = "ERR_ANA_LIVELLO_RADARULTR_F12_UM_NON_MA"
KEY_ERR_ANA_PRESS_DP_TX_C9F9_UM_DIVERSE = "ERR_ANA_PRESS_DP_TX_C9F9_UM_DIVERSE"
KEY_ERR_ANA_PRESS_DP_TX_F12_UM_NON_MA = "ERR_ANA_PRESS_DP_TX_F12_UM_NON_MA"
KEY_ERR_ANA_PRESS_DP_TX_A9D9_RANGE_DIVERSI = "ERR_ANA_PRESS_DP_TX_A9D9_RANGE_DIVERSI"
KEY_ERR_ANA_PRESS_DP_TX_D12_RANGE_NON_4_20 = "ERR_ANA_PRESS_DP_TX_D12_RANGE_NON_4_20"
KEY_ERR_ANA_PORTATA_DP_D12_RANGE_NON_4_20 = "ERR_ANA_PORTATA_DP_D12_RANGE_NON_4_20"
KEY_ERR_ANA_PORTATA_DP_F12_UM_NON_MA = "ERR_ANA_PORTATA_DP_F12_UM_NON_MA"
KEY_ERR_DIG_PRESS_D22_UM_NON_PRESSIONE = "ERR_DIG_PRESS_D22_UM_NON_PRESSIONE"
KEY_ERR_DIG_LIVELLO_D22_UM_NON_PERCENTO = "ERR_DIG_LIVELLO_D22_UM_NON_PERCENTO"
KEY_COMP_ANA_ODC_MANCANTE = "COMP_ANA_ODC_MANCANTE"
KEY_COMP_ANA_DATA_COMP_MANCANTE = "COMP_ANA_DATA_COMP_MANCANTE"
KEY_COMP_ANA_PDL_MANCANTE = "COMP_ANA_PDL_MANCANTE"
KEY_COMP_ANA_ESECUTORE_MANCANTE = "COMP_ANA_ESECUTORE_MANCANTE"
KEY_COMP_ANA_SUPERVISORE_MANCANTE = "COMP_ANA_SUPERVISORE_MANCANTE"
KEY_COMP_ANA_CONTRATTO_MANCANTE = "COMP_ANA_CONTRATTO_MANCANTE"
KEY_COMP_ANA_CONTRATTO_DIVERSO = "COMP_ANA_CONTRATTO_DIVERSO" 
KEY_COMP_DIG_ODC_MANCANTE = "COMP_DIG_ODC_MANCANTE"
KEY_COMP_DIG_DATA_COMP_MANCANTE = "COMP_DIG_DATA_COMP_MANCANTE"
KEY_COMP_DIG_PDL_MANCANTE = "COMP_DIG_PDL_MANCANTE"
KEY_COMP_DIG_ESECUTORE_MANCANTE = "COMP_DIG_ESECUTORE_MANCANTE"
KEY_COMP_DIG_SUPERVISORE_MANCANTE = "COMP_DIG_SUPERVISORE_MANCANTE"
KEY_COMP_DIG_CONTRATTO_MANCANTE = "COMP_DIG_CONTRATTO_MANCANTE"
KEY_COMP_DIG_CONTRATTO_DIVERSO = "COMP_DIG_CONTRATTO_DIVERSO" 
KEY_COMP_CAMPI_MANCANTI_NON_LEGGIBILI = "COMP_CAMPI_MANCANTI_NON_LEGGIBILI"
MAPPA_SP_TIPOLOGIA = {"SP 11/04":"LIVELLO","SP 11-04":"LIVELLO","SP 11/03":"TEMPERATURA","SP 11-03":"TEMPERATURA","SP 11/02":"PRESSIONE","SP 11-02":"PRESSIONE","SP 11/01":"PORTATA","SP 11-01":"PORTATA"}
MAPPA_L9_SOTTOTIPO_NORMALIZZATA = {"DP":["PRESSIONE","PORTATA","LIVELLO"],"CAPILLARE":["PRESSIONE","PORTATA","LIVELLO"],"TX":["PRESSIONE"],"TX PRESSIONE":["PRESSIONE"],"TX DI PRESSIONE":["PRESSIONE"],"TORSIONALE":["LIVELLO"],"TORSIONALE PNEUMATICO":["LIVELLO"],"TORSIONALE LOCALE":["LIVELLO"],"BARRA DI TORSIONE":["LIVELLO"],"ONDA GUIDATA":["LIVELLO"],"RADAR":["LIVELLO"],"ULTRASUONI":["LIVELLO","PORTATA"],"K":["TEMPERATURA_TERMOCOPPIA"],"J":["TEMPERATURA_TERMOCOPPIA"],"SKIN POINT K":["TEMPERATURA_TERMOCOPPIA"],"SKIN POINT J":["TEMPERATURA_TERMOCOPPIA"],"TERMOCOPPIA K":["TEMPERATURA_TERMOCOPPIA"],"TERMOCOPPIA J":["TEMPERATURA_TERMOCOPPIA"],"TERMOCOPPIA":["TEMPERATURA_TERMOCOPPIA"],"TEMOCOPPIA J":["TEMPERATURA_TERMOCOPPIA"],"TC K":["TEMPERATURA_TERMOCOPPIA"],"TC J":["TEMPERATURA_TERMOCOPPIA"],"K TC":["TEMPERATURA_TERMOCOPPIA"],"J TC":["TEMPERATURA_TERMOCOPPIA"],"TERMOCOPPIA TIPO K":["TEMPERATURA_TERMOCOPPIA"],"TERMOCOPPIA TIPO J":["TEMPERATURA_TERMOCOPPIA"],"RTD 3W":["TEMPERATURA_RTD"],"RTD":["TEMPERATURA_RTD"],"RTD 2W":["TEMPERATURA_RTD"],"TERMORESISTENZA":["TEMPERATURA_RTD"],"PT100":["TEMPERATURA_RTD"],"CONVERTITORE":["TEMPERATURA_CONVERTITORE"],"INDICATORE LOCALE":["LIVELLO","PORTATA","PRESSIONE","TEMPERATURA"]}
REGOLE_CONGRUITA_CERTIFICATI_NORMALIZZATE = {"TEMPERATURA":{"modelli_campione_congrui":["CALIBR. TEMPERATURA"],"sottotipi_l9":{"TEMPERATURA_TERMOCOPPIA":["TERMOCOPPIA CAMPIONE","MULTIMETRO DIGITALE"],"TEMPERATURA_RTD":["TERMORESISTENZA CAMPIONE","MULTIMETRO DIGITALE"],"TEMPERATURA_CONVERTITORE":["MULTIMETRO DIGITALE","CALIBRATORE DI LOOP"]},"modelli_campione_incongrui":["MANOMETRO DIGITALE","CALIBRATORE DI LOOP"],"eccezioni_l9_incongrui":{"TEMPERATURA_CONVERTITORE":["MANOMETRO DIGITALE","TERMOCOPPIA CAMPIONE","TERMORESISTENZA CAMPIONE"]}},"PRESSIONE":{"modelli_campione_congrui":["MANOMETRO DIGITALE","MULTIMETRO DIGITALE","CALIBRATORE DI LOOP"],"modelli_campione_incongrui":["CALIBR. TEMPERATURA","TERMOCOPPIA CAMPIONE","TERMORESISTENZA CAMPIONE"]},"PORTATA":{"modelli_campione_congrui":["MANOMETRO DIGITALE","MULTIMETRO DIGITALE","CALIBRATORE DI LOOP"],"modelli_campione_incongrui":["CALIBR. TEMPERATURA","TERMOCOPPIA CAMPIONE","TERMORESISTENZA CAMPIONE"]},"LIVELLO":{"modelli_campione_congrui":["MULTIMETRO DIGITALE","CALIBRATORE DI LOOP","COMPARATORE", "MANOMETRO DIGITALE"],"modelli_campione_incongrui":["CALIBR. TEMPERATURA","TERMOCOPPIA CAMPIONE","TERMORESISTENZA CAMPIONE"]}}
for tipologia_regola, item_regola in REGOLE_CONGRUITA_CERTIFICATI_NORMALIZZATE.items():
    if "modelli_campione_congrui" in item_regola: item_regola["modelli_campione_congrui"] = [m.strip().upper() for m in item_regola.get("modelli_campione_congrui", [])]
    if "modelli_campione_incongrui" in item_regola: item_regola["modelli_campione_incongrui"] = [m.strip().upper() for m in item_regola.get("modelli_campione_incongrui", [])]
    if "sottotipi_l9" in item_regola and isinstance(item_regola["sottotipi_l9"], dict):
        for sottotipo_key, modelli_sottotipo_list in item_regola["sottotipi_l9"].items():
            if isinstance(modelli_sottotipo_list, list): item_regola["sottotipi_l9"][sottotipo_key] = [m.strip().upper() for m in modelli_sottotipo_list]
    if "eccezioni_l9_incongrui" in item_regola and isinstance(item_regola["eccezioni_l9_incongrui"], dict):
        for eccezione_l9_key, modelli_eccezione_list in item_regola["eccezioni_l9_incongrui"].items():
            if isinstance(modelli_eccezione_list, list): item_regola["eccezioni_l9_incongrui"][eccezione_l9_key] = [m.strip().upper() for m in modelli_eccezione_list]
LISTA_UM_PRESSIONE_RICONOSCIUTE = sorted(["bar","barg","bara","mbar","mbarg","mbara","pa","kpa","mpa","psi","psig","psia","mmh2o","cmh2o","mh2o","mmhg","cmhg","mhg","kg/cm2"])
MAPPA_NORMALIZZAZIONE_UM = {"mm h2o":"mmh2o","mmh₂o":"mmh2o","mm H₂O":"mmh2o","mm H2O":"mmh2o","kg/cm²":"kg/cm2","kg/cm^2":"kg/cm2","milliampere":"ma","milli ampere":"ma","milliamperes":"ma","mamp":"ma","percent":"%","percentage":"%"}
RANGE_0_100_NORMALIZZATO="0-100";RANGE_4_20_NORMALIZZATO="4-20";UM_MA_NORMALIZZATA="ma";UM_PERCENTO_NORMALIZZATA="%";UM_MMH2O_NORMALIZZATA="mmh2o";UM_MM_NORMALIZZATA="mm";UM_PSI_NORMALIZZATA="psi"
human_error_messages_map_descriptive = {
    KEY_SP_VUOTO: f"Codice SP (Tipologia Strumento da N9/N10) mancante o illeggibile.", KEY_L9_VUOTO: f"Modello Strumento (L9, solo per analogici) mancante o illeggibile.", KEY_L9_SKINPOINT_INCOMPLETO: f"Modello L9 (analogico) 'SKIN POINT' incompleto (manca tipo K, J, ecc.).", KEY_CELL_RANGE_UM_NON_LEGGIBILE: f"Impossibile leggere una o più celle necessarie per Range/UM (A9,C9,D9,F9,D12,F12 per analogici; D22 per digitali).",
    KEY_ERR_ANA_TEMP_CONV_C9F9_UM_DIVERSE: f"Temp./Convertitore (ANA): Unità Ingresso ({SCHEDA_ANA_CELL_UM_INGRESSO}) diversa da Unità DCS ({SCHEDA_ANA_CELL_UM_DCS}). Devono coincidere.", KEY_ERR_ANA_TEMP_CONV_F12_UM_NON_MA: f"Temp./Convertitore (ANA): Unità Uscita ({SCHEDA_ANA_CELL_UM_USCITA}) deve essere '{UM_MA_NORMALIZZATA}'.", KEY_ERR_ANA_TEMP_CONV_A9D9_RANGE_DIVERSI: f"Temp./Convertitore (ANA): Range Ingresso ({SCHEDA_ANA_CELL_RANGE_INGRESSO}) diverso da Range DCS ({SCHEDA_ANA_CELL_RANGE_DCS}). Devono coincidere.", KEY_ERR_ANA_TEMP_CONV_D12_RANGE_NON_4_20: f"Temp./Convertitore (ANA): Range Uscita ({SCHEDA_ANA_CELL_RANGE_USCITA}) deve essere '{RANGE_4_20_NORMALIZZATO}'.", KEY_ERR_ANA_TEMP_NOCONV_UM_NON_COINCIDENTI: f"Temp./NO-Convertitore (RTD, TC) (ANA): Unità Ingresso ({SCHEDA_ANA_CELL_UM_INGRESSO}), DCS ({SCHEDA_ANA_CELL_UM_DCS}) e Uscita ({SCHEDA_ANA_CELL_UM_USCITA}) devono coincidere.", KEY_ERR_ANA_TEMP_NOCONV_RANGE_NON_COINCIDENTI: f"Temp./NO-Convertitore (RTD, TC) (ANA): Range Ingresso ({SCHEDA_ANA_CELL_RANGE_INGRESSO}), DCS ({SCHEDA_ANA_CELL_RANGE_DCS}) e Uscita ({SCHEDA_ANA_CELL_RANGE_USCITA}) devono coincidere.",
    KEY_ERR_ANA_LIVELLO_DP_C9_UM_NON_PRESSIONE: f"Livello/DP (ANA): Unità Ingresso ({SCHEDA_ANA_CELL_UM_INGRESSO}) deve essere un'unità di misura di pressione valida (es. mmH2O, bar).", KEY_ERR_ANA_LIVELLO_DP_D9_RANGE_NON_0_100: f"Livello/DP (ANA): Range DCS ({SCHEDA_ANA_CELL_RANGE_DCS}) deve essere '{RANGE_0_100_NORMALIZZATO}'.", KEY_ERR_ANA_LIVELLO_DP_F9_UM_NON_PERCENTO: f"Livello/DP (ANA): Unità DCS ({SCHEDA_ANA_CELL_UM_DCS}) deve essere '{UM_PERCENTO_NORMALIZZATA}'.", KEY_ERR_ANA_LIVELLO_DP_D12_RANGE_NON_4_20: f"Livello/DP (ANA): Range Uscita ({SCHEDA_ANA_CELL_RANGE_USCITA}) deve essere '{RANGE_4_20_NORMALIZZATO}'.", KEY_ERR_ANA_LIVELLO_DP_F12_UM_NON_MA: f"Livello/DP (ANA): Unità Uscita ({SCHEDA_ANA_CELL_UM_USCITA}) deve essere '{UM_MA_NORMALIZZATA}'.", KEY_ERR_ANA_LIVELLO_TORS_C9_UM_INVALIDA: f"Livello/BarraTors. (ANA): Unità Ingresso ({SCHEDA_ANA_CELL_UM_INGRESSO}) deve essere '{UM_MMH2O_NORMALIZZATA}' o '{UM_MM_NORMALIZZATA}'.", KEY_ERR_ANA_LIVELLO_TORS_D9_RANGE_NON_0_100: f"Livello/BarraTors. (ANA): Range DCS ({SCHEDA_ANA_CELL_RANGE_DCS}) deve essere '{RANGE_0_100_NORMALIZZATO}'.", KEY_ERR_ANA_LIVELLO_TORS_F9_UM_NON_PERCENTO: f"Livello/BarraTors. (ANA): Unità DCS ({SCHEDA_ANA_CELL_UM_DCS}) deve essere '{UM_PERCENTO_NORMALIZZATA}'.", KEY_ERR_ANA_LIVELLO_TORS_ELETTR_D12_RANGE_NON_4_20: f"Livello/BarraTors.Elettronico (ANA): Range Uscita ({SCHEDA_ANA_CELL_RANGE_USCITA}) deve essere '{RANGE_4_20_NORMALIZZATO}'.", KEY_ERR_ANA_LIVELLO_TORS_ELETTR_F12_UM_NON_MA: f"Livello/BarraTors.Elettronico (ANA): Unità Uscita ({SCHEDA_ANA_CELL_UM_USCITA}) deve essere '{UM_MA_NORMALIZZATA}'.", KEY_ERR_ANA_LIVELLO_TORS_LOCALE_D12_RANGE_NON_VUOTO: f"Livello/BarraTors.Locale (ANA): Range Uscita ({SCHEDA_ANA_CELL_RANGE_USCITA}) deve essere vuoto (o '{UM_PSI_NORMALIZZATA}' se pneumatico).", KEY_ERR_ANA_LIVELLO_TORS_LOCALE_F12_UM_NON_VUOTA: f"Livello/BarraTors.Locale (ANA): Unità Uscita ({SCHEDA_ANA_CELL_UM_USCITA}) deve essere vuota (o '{UM_PSI_NORMALIZZATA}' se pneumatico).", KEY_ERR_ANA_LIVELLO_RADARULTR_D9_RANGE_NON_0_100: f"Livello/Radar-Ultrasuoni (ANA): Range DCS ({SCHEDA_ANA_CELL_RANGE_DCS}) deve essere '{RANGE_0_100_NORMALIZZATO}'.", KEY_ERR_ANA_LIVELLO_RADARULTR_F9_UM_NON_PERCENTO: f"Livello/Radar-Ultrasuoni (ANA): Unità DCS ({SCHEDA_ANA_CELL_UM_DCS}) deve essere '{UM_PERCENTO_NORMALIZZATA}'.", KEY_ERR_ANA_LIVELLO_RADARULTR_D12_RANGE_NON_4_20: f"Livello/Radar-Ultrasuoni (ANA): Range Uscita ({SCHEDA_ANA_CELL_RANGE_USCITA}) deve essere '{RANGE_4_20_NORMALIZZATO}'.", KEY_ERR_ANA_LIVELLO_RADARULTR_F12_UM_NON_MA: f"Livello/Radar-Ultrasuoni (ANA): Unità Uscita ({SCHEDA_ANA_CELL_UM_USCITA}) deve essere '{UM_MA_NORMALIZZATA}'.",
    KEY_ERR_ANA_PRESS_DP_TX_C9F9_UM_DIVERSE: f"Pressione/DP-TX-Capillare (ANA): Unità Ingresso ({SCHEDA_ANA_CELL_UM_INGRESSO}) e DCS ({SCHEDA_ANA_CELL_UM_DCS}) devono coincidere.", KEY_ERR_ANA_PRESS_DP_TX_F12_UM_NON_MA: f"Pressione/DP-TX-Capillare (ANA): Unità Uscita ({SCHEDA_ANA_CELL_UM_USCITA}) deve essere '{UM_MA_NORMALIZZATA}'.", KEY_ERR_ANA_PRESS_DP_TX_A9D9_RANGE_DIVERSI: f"Pressione/DP-TX-Capillare (ANA): Range Ingresso ({SCHEDA_ANA_CELL_RANGE_INGRESSO}) e DCS ({SCHEDA_ANA_CELL_RANGE_DCS}) devono coincidere.", KEY_ERR_ANA_PRESS_DP_TX_D12_RANGE_NON_4_20: f"Pressione/DP-TX-Capillare (ANA): Range Uscita ({SCHEDA_ANA_CELL_RANGE_USCITA}) deve essere '{RANGE_4_20_NORMALIZZATO}'.",
    KEY_ERR_ANA_PORTATA_DP_D12_RANGE_NON_4_20: f"Portata/DP (ANA): Range Uscita ({SCHEDA_ANA_CELL_RANGE_USCITA}) deve essere '{RANGE_4_20_NORMALIZZATO}'.", KEY_ERR_ANA_PORTATA_DP_F12_UM_NON_MA: f"Portata/DP (ANA): Unità Uscita ({SCHEDA_ANA_CELL_UM_USCITA}) deve essere '{UM_MA_NORMALIZZATA}'.",
    KEY_ERR_DIG_PRESS_D22_UM_NON_PRESSIONE: f"Digitale/Pressione: Unità Processo ({SCHEDA_DIG_CELL_RANGE_UM_PROCESSO}) non è UM di pressione valida.", KEY_ERR_DIG_LIVELLO_D22_UM_NON_PERCENTO: f"Digitale/Livello: Unità Processo ({SCHEDA_DIG_CELL_RANGE_UM_PROCESSO}) deve essere '{UM_PERCENTO_NORMALIZZATA}'.",
    KEY_COMP_ANA_ODC_MANCANTE: f"Scheda Analogica: ODC ({SCHEDA_ANA_CELL_ODC}) mancante.", KEY_COMP_ANA_DATA_COMP_MANCANTE: f"Scheda Analogica: Data Compilazione ({SCHEDA_ANA_CELL_DATA_COMPILAZIONE}) mancante.", KEY_COMP_ANA_PDL_MANCANTE: f"Scheda Analogica: Numero PDL ({SCHEDA_ANA_CELL_PDL}) mancante.", KEY_COMP_ANA_ESECUTORE_MANCANTE: f"Scheda Analogica: Esecutore ({SCHEDA_ANA_CELL_ESECUTORE}) mancante.", KEY_COMP_ANA_SUPERVISORE_MANCANTE: f"Scheda Analogica: Supervisore ISAB ({SCHEDA_ANA_CELL_SUPERVISORE_ISAB}) mancante.", KEY_COMP_ANA_CONTRATTO_MANCANTE: f"Scheda Analogica: Contratto Coemi ({SCHEDA_ANA_CELL_CONTRATTO_COEMI}) mancante.", KEY_COMP_ANA_CONTRATTO_DIVERSO: f"Scheda Analogica: Contratto Coemi ({SCHEDA_ANA_CELL_CONTRATTO_COEMI}) diverso da '{VALORE_ATTESO_CONTRATTO_COEMI}' o '{VALORE_ATTESO_CONTRATTO_COEMI_VARIANTE_NUMERICA}'.",
    KEY_COMP_DIG_ODC_MANCANTE: f"Scheda Digitale: ODC ({SCHEDA_DIG_CELL_ODC}) mancante.", KEY_COMP_DIG_DATA_COMP_MANCANTE: f"Scheda Digitale: Data Compilazione ({SCHEDA_DIG_CELL_DATA_COMPILAZIONE}) mancante.", KEY_COMP_DIG_PDL_MANCANTE: f"Scheda Digitale: Numero PDL ({SCHEDA_DIG_CELL_PDL}) mancante.", KEY_COMP_DIG_ESECUTORE_MANCANTE: f"Scheda Digitale: Esecutore ({SCHEDA_DIG_CELL_ESECUTORE}) mancante.", KEY_COMP_DIG_SUPERVISORE_MANCANTE: f"Scheda Digitale: Supervisore ISAB ({SCHEDA_DIG_CELL_SUPERVISORE_ISAB}) mancante.", KEY_COMP_DIG_CONTRATTO_MANCANTE: f"Scheda Digitale: Contratto Coemi ({SCHEDA_DIG_CELL_CONTRATTO_COEMI}) mancante.", KEY_COMP_DIG_CONTRATTO_DIVERSO: f"Scheda Digitale: Contratto Coemi ({SCHEDA_DIG_CELL_CONTRATTO_COEMI}) diverso da '{VALORE_ATTESO_CONTRATTO_COEMI}' o '{VALORE_ATTESO_CONTRATTO_COEMI_VARIANTE_NUMERICA}'.", KEY_COMP_CAMPI_MANCANTI_NON_LEGGIBILI: "Impossibile leggere uno o più campi per la verifica di completezza (ODC, Data, PDL, etc.).",
}
def excel_coord_to_indices(coord_str):
    match=re.match(r"([A-Z]+)([0-9]+)",coord_str.upper());
    if not match:raise ValueError(f"Coordinata Excel non valida: {coord_str}")
    col_s,row_s=match.groups();col_idx=0
    for char_i,char_v in enumerate(reversed(col_s)):col_idx+=(ord(char_v)-ord('A')+1)*(26**char_i)
    return int(row_s)-1,col_idx-1
IDX_DIG_TIPOLOGIA_STRUM=excel_coord_to_indices(SCHEDA_DIG_CELL_TIPOLOGIA_STRUM);IDX_ANA_TIPOLOGIA_STRUM=excel_coord_to_indices(SCHEDA_ANA_CELL_TIPOLOGIA_STRUM)
IDX_ANA_MODELLO_STRUM=excel_coord_to_indices(SCHEDA_ANA_CELL_MODELLO_STRUM);IDX_ANA_RANGE_INGRESSO=excel_coord_to_indices(SCHEDA_ANA_CELL_RANGE_INGRESSO)
IDX_ANA_UM_INGRESSO=excel_coord_to_indices(SCHEDA_ANA_CELL_UM_INGRESSO);IDX_ANA_RANGE_USCITA=excel_coord_to_indices(SCHEDA_ANA_CELL_RANGE_USCITA)
IDX_ANA_UM_USCITA=excel_coord_to_indices(SCHEDA_ANA_CELL_UM_USCITA);IDX_ANA_RANGE_DCS=excel_coord_to_indices(SCHEDA_ANA_CELL_RANGE_DCS)
IDX_ANA_UM_DCS=excel_coord_to_indices(SCHEDA_ANA_CELL_UM_DCS);IDX_DIG_RANGE_UM_PROCESSO=excel_coord_to_indices(SCHEDA_DIG_CELL_RANGE_UM_PROCESSO)
IDX_ANA_UM_USCITA_ROW, IDX_ANA_UM_USCITA_COL = IDX_ANA_UM_USCITA
IDX_ANA_ODC = excel_coord_to_indices(SCHEDA_ANA_CELL_ODC)
IDX_ANA_DATA_COMPILAZIONE = excel_coord_to_indices(SCHEDA_ANA_CELL_DATA_COMPILAZIONE)
IDX_ANA_PDL = excel_coord_to_indices(SCHEDA_ANA_CELL_PDL)
IDX_ANA_ESECUTORE = excel_coord_to_indices(SCHEDA_ANA_CELL_ESECUTORE)
IDX_ANA_SUPERVISORE_ISAB = excel_coord_to_indices(SCHEDA_ANA_CELL_SUPERVISORE_ISAB)
IDX_ANA_CONTRATTO_COEMI = excel_coord_to_indices(SCHEDA_ANA_CELL_CONTRATTO_COEMI)
IDX_DIG_ODC = excel_coord_to_indices(SCHEDA_DIG_CELL_ODC)
IDX_DIG_DATA_COMPILAZIONE = excel_coord_to_indices(SCHEDA_DIG_CELL_DATA_COMPILAZIONE)
IDX_DIG_PDL = excel_coord_to_indices(SCHEDA_DIG_CELL_PDL)
IDX_DIG_ESECUTORE = excel_coord_to_indices(SCHEDA_DIG_CELL_ESECUTORE)
IDX_DIG_SUPERVISORE_ISAB = excel_coord_to_indices(SCHEDA_DIG_CELL_SUPERVISORE_ISAB)
IDX_DIG_CONTRATTO_COEMI = excel_coord_to_indices(SCHEDA_DIG_CELL_CONTRATTO_COEMI)
def log_message_gui(log_text_widget, message, level="INFO"):
    if log_text_widget and log_text_widget.winfo_exists():
        try:
            log_text_widget.config(state=tk.NORMAL); tag=None; font_options={}; color="black"
            if level=="ERROR": tag,color,font_options="error_tag","red",{'weight':'bold'}
            elif level=="WARNING": tag,color="warning_tag","darkorange"
            elif level=="SUCCESS": tag,color="success_tag","green"
            elif level=="DEBUG": tag,color="debug_tag","gray50"
            if tag and tag not in log_text_widget.tag_names():
                current_font=tkFont.Font(font=log_text_widget.cget("font"))
                custom_font=tkFont.Font(family=current_font.actual("family"),size=current_font.actual("size"),
                                          weight=font_options.get('weight',current_font.actual("weight")),
                                          slant=font_options.get('slant',current_font.actual("slant")))
                log_text_widget.tag_configure(tag,foreground=color,font=custom_font)
            log_text_widget.insert(tk.END,f"{message}\n",tag if tag else None)
            log_text_widget.see(tk.END); log_text_widget.config(state=tk.DISABLED); log_text_widget.update_idletasks()
        except tk.TclError as e_tcl: logger.debug(f"TclError in log_message_gui (widget likely destroyed): {e_tcl}")
        except Exception as e_gui_log:
            if logger.handlers: logger.error(f"Errore in log_message_gui: {e_gui_log}",exc_info=True)
            else: print(f"Errore in log_message_gui (logger non pronto): {e_gui_log}", file=sys.stderr)
def log_to_all(log_widget, message, level="INFO", exc_info=False):
    log_level_map={"ERROR":logging.ERROR,"WARNING":logging.WARNING,"INFO":logging.INFO,"SUCCESS":logging.INFO,"DEBUG":logging.DEBUG}
    if logger.hasHandlers():
        logger.log(log_level_map.get(level.upper(),logging.INFO),message,exc_info=exc_info)
    else: 
        print(f"{level}: {message}", file=sys.stderr if level in ["ERROR", "WARNING"] else sys.stdout)
        if exc_info:
            import traceback
            traceback.print_exc(file=sys.stderr)
    log_message_gui(log_widget,message,level.upper())
def parse_date_robust(date_str_val, context_filename="N/A"):
    if pd.isna(date_str_val): return None
    if isinstance(date_str_val,datetime): return date_str_val
    if isinstance(date_str_val,pd.Timestamp): return date_str_val.to_pydatetime()
    s_date_str = str(date_str_val).strip()
    if not s_date_str: return None
    expected_formats = ['%d/%m/%Y','%d-%m-%Y','%d.%m.%Y','%Y-%m-%d %H:%M:%S','%Y-%m-%d']
    common_formats = ['%m/%d/%Y','%Y/%m/%d'] 
    all_formats_to_try = expected_formats + [f for f in common_formats if f not in expected_formats]
    for fmt in all_formats_to_try:
        try:
            return datetime.strptime(s_date_str.split(' ')[0],fmt)
        except ValueError: continue
    try:
        if isinstance(date_str_val,(int,float)) or \
           (s_date_str.replace('.','',1).isdigit() and ('.' in s_date_str or float(s_date_str) > 1000)):
            numeric_val = float(s_date_str)
            if 1 < numeric_val < 200000: 
                return pd.to_datetime(numeric_val,unit='D',origin='1899-12-30').to_pydatetime()
    except (ValueError,TypeError,OverflowError) as e_num: 
        logger.debug(f"File: {context_filename} - Parse numerico Excel fallito per '{s_date_str}': {e_num}")
    logger.warning(f"File: {context_filename} - Data '{s_date_str}' (raw: '{date_str_val}') non riconosciuta con formati standard o Excel serial.")
    return None
def normalize_sp_code(sp_code_raw):
    if pd.isna(sp_code_raw): return ""
    s_norm = str(sp_code_raw).strip().upper()
    s_norm = s_norm.replace("S.P.","SP").replace(".","")
    s_norm = s_norm.replace("-","/")
    s_norm = " ".join(s_norm.split())
    s_norm = re.sub(r'\s*SP\s*(\d+)\s*/\s*(\d+)',r'SP \1/\2',s_norm)
    return s_norm
def normalize_um(um_str_raw): 
    if pd.isna(um_str_raw): return ""
    s_norm = str(um_str_raw).strip().lower(); s_norm = " ".join(s_norm.split()) 
    for k_map,v_map in MAPPA_NORMALIZZAZIONE_UM.items(): s_norm = s_norm.replace(k_map,v_map)
    s_norm = s_norm.replace(" ","") 
    return s_norm
def normalize_range_string(range_str_raw):
    if pd.isna(range_str_raw): return ""
    if isinstance(range_str_raw,(int,float)):
        range_str_raw = str(int(range_str_raw)) if range_str_raw == int(range_str_raw) else str(range_str_raw)
    norm_str = str(range_str_raw).lower(); norm_str = " ".join(norm_str.split()) 
    norm_str = re.sub(r'\s*[\/÷]\s*','-',norm_str); norm_str = re.sub(r'\s*to\s*','-',norm_str,flags=re.IGNORECASE) 
    norm_str = re.sub(r'\s*-\s*','-',norm_str); norm_str = re.sub(r'\s+','',norm_str) 
    return norm_str
def is_valid_expected_range(actual_range_norm,expected_range_norm_target):
    return actual_range_norm == expected_range_norm_target
def is_um_pressione_valida(um_str_normalized_input):
    return um_str_normalized_input in LISTA_UM_PRESSIONE_RICONOSCIUTE
def is_cell_value_empty(cell_value):
    if pd.isna(cell_value): return True
    if isinstance(cell_value,str) and not cell_value.strip(): return True
    if isinstance(cell_value,str) and cell_value.strip().lower()=="nan": return True 
    return False
# --- FINE PARTE 1 ---
# --- PARTE 2 ---
def leggi_registro_strumenti(log_widget_gui):
    if not FILE_REGISTRO_STRUMENTI:
        log_to_all(log_widget_gui, "Percorso FILE_REGISTRO_STRUMENTI non configurato. Impossibile leggere il registro.", "ERROR")
        return None

    logger.info(f"Tentativo lettura registro strumenti: {FILE_REGISTRO_STRUMENTI}")
    if not os.path.exists(FILE_REGISTRO_STRUMENTI):
        log_to_all(log_widget_gui, f"File registro strumenti NON TROVATO: {FILE_REGISTRO_STRUMENTI}", "ERROR")
        return None
    try:
        # Mappa nomi desiderati a indici colonna (0-based)
        cols_to_read_indices_map = {
            'modello_strumento_campione': REGISTRO_COL_IDX_MODELLO_STRUM_CAMPIONE, # G -> 6
            'id_cert_campione': REGISTRO_COL_IDX_ID_CERT_CAMPIONE,             # Q -> 16
            'range_campione': REGISTRO_COL_IDX_RANGE_CAMPIONE,               # M -> 12
            'scadenza_cert_campione': REGISTRO_COL_IDX_SCADENZA_CAMPIONE      # S -> 18
        }
        # Ordina le colonne da leggere per indice per pd.read_excel
        sorted_map_items = sorted(cols_to_read_indices_map.items(), key=lambda item: item[1])
        sorted_col_indices = [item[1] for item in sorted_map_items]
        sorted_col_names = [item[0] for item in sorted_map_items]

        df_registro = pd.read_excel(FILE_REGISTRO_STRUMENTI, 
                                    sheet_name=REGISTRO_FOGLIO_NOME,
                                    header=None, # Nessun header nel file Excel alle righe che leggiamo
                                    skiprows=REGISTRO_RIGA_INIZIO_DATI - 1, # Salta le righe sopra i dati effettivi
                                    usecols=sorted_col_indices, # Leggi solo le colonne specificate
                                    engine='openpyxl',
                                    dtype=str # Leggi tutto come stringa inizialmente
                                    )
        df_registro.columns = sorted_col_names # Assegna i nomi corretti alle colonne lette

        # Rimuovi righe dove 'id_cert_campione' è mancante o vuoto dopo lo strip
        df_registro.dropna(subset=['id_cert_campione'], inplace=True)
        df_registro = df_registro[df_registro['id_cert_campione'].astype(str).str.strip() != ""]
        
        strumenti_campione = []
        for _, row in df_registro.iterrows():
            modello_strum = str(row['modello_strumento_campione']).strip().upper()
            id_cert_strum = str(row['id_cert_campione']).strip()
            range_strum = str(row['range_campione']).strip() # Normalizzazione range non necessaria qui
            scadenza_val = row['scadenza_cert_campione'] # Può essere data, stringa, numero
            
            scadenza_dt = parse_date_robust(scadenza_val, FILE_REGISTRO_STRUMENTI)

            data_emissione_dt = None
            if scadenza_dt:
                try:
                    # Tentativo con DateOffset per gestire correttamente anni bisestili, ecc.
                    data_emissione_dt = scadenza_dt - DateOffset(years=1)
                except Exception as e_offset: # pandas <-> datetime a volte problematico
                    logger.warning(f"Errore calcolo data emissione per {id_cert_strum} (scad: {scadenza_dt}) con DateOffset: {e_offset}. Tentativo con timedelta.")
                    try:
                        data_emissione_dt = scadenza_dt - timedelta(days=365) # Approssimazione
                    except Exception as e_delta:
                        logger.error(f"Fallito anche calcolo data emissione con timedelta per {id_cert_strum}: {e_delta}")
                        data_emissione_dt = None # Non si può calcolare


            if id_cert_strum: # Assicura che ci sia un ID certificato
                strumenti_campione.append({
                    'modello_strumento': modello_strum if modello_strum and modello_strum.lower() != 'nan' else "N/D",
                    'id_certificato': id_cert_strum,
                    'range': range_strum if range_strum and range_strum.lower() != 'nan' else "N/D",
                    'scadenza': scadenza_dt, # Oggetto datetime o None
                    'scadenza_raw': scadenza_val if scadenza_val and str(scadenza_val).lower() != 'nan' else "N/D",
                    'data_emissione': data_emissione_dt # Oggetto datetime o None
                })
        log_to_all(log_widget_gui, f"Letti {len(strumenti_campione)} strumenti validi dal registro.", "SUCCESS" if strumenti_campione else "WARNING")
        return strumenti_campione
    except ValueError as ve: # Es. nome foglio non trovato
        log_to_all(log_widget_gui, f"Errore lettura registro (ValueError): {ve}. Verifica nome foglio, indici e nomi colonne.", "ERROR")
        return None
    except Exception as e:
        log_to_all(log_widget_gui, f"Errore imprevisto durante lettura registro strumenti: {e}", "ERROR", exc_info=True)
        return None

def analyze_excel_file(file_path, log_widget_gui, strumenti_campione_global_list):
    base_filename = os.path.basename(file_path)
    file_ext = os.path.splitext(file_path)[1].lower()
    engine = None
    human_error_keys_this_file = [] # Lista per le chiavi di errore "umano" specifiche per questo file

    # Dizionario per raccogliere dati utili alla compilazione automatica
    dati_scheda_per_compilazione = {
        "file_path": file_path, 
        "base_filename": base_filename,
        "file_type": None, # 'analogico' o 'digitale'
        "campi_mancanti": set(), # Set di chiavi errore tipo KEY_COMP_...
        "pdl_val": None, # Valore PDL letto dalla scheda
        "odc_val_scheda": None # Valore ODC letto dalla scheda
    }

    if file_ext == '.xlsx': engine = 'openpyxl'
    elif file_ext == '.xls': engine = 'xlrd' # xlrd per .xls
    else:
        log_to_all(log_widget_gui, f"{base_filename}: Tipo file non supportato '{file_ext}'. Sarà saltato.", "WARNING")
        return None, "Tipo file non supportato", [], human_error_keys_this_file, dati_scheda_per_compilazione


    try:
        # Leggi tutti i dati come stringhe per evitare conversioni automatiche di pandas che potrebbero alterare formati
        df = pd.read_excel(file_path, header=None, sheet_name=0, engine=engine, dtype=str)
    except Exception as e:
        if "Workbook is encrypted" in str(e) and engine == 'xlrd': # xlrd non gestisce file .xls protetti
             log_to_all(log_widget_gui, f"{base_filename}: IMPOSSIBILE APRIRE - File Excel (.xls) PROTETTO DA PASSWORD. Errore: {e}", "ERROR")
             return None, f"File protetto da password ({engine})", [], human_error_keys_this_file, dati_scheda_per_compilazione
        log_to_all(log_widget_gui, f"{base_filename}: IMPOSSIBILE APRIRE o leggere il foglio (engine: {engine}). Errore: {e}", "ERROR", exc_info=True)
        return None, f"Errore apertura ({engine})", [], human_error_keys_this_file, dati_scheda_per_compilazione
    
    file_type = None; card_date = None; extracted_certs_data = []
    tipologia_strumento_scheda = "N/D_INIZIALE"; modello_l9_scheda_normalizzato = "N/A" # Solo per analogici
    sp_code_normalizzato_letto = "N/A" # SP Code da N9/N10 dopo normalizzazione
    # Variabili per Range/UM (analogici)
    range_ing_raw, um_ing_raw, range_usc_raw, um_usc_raw, range_dcs_raw, um_dcs_raw = "","","","","",""
    range_ing_norm, um_ing_norm, range_usc_norm, um_usc_norm, range_dcs_norm, um_dcs_norm = "","","","","",""
    

    # Determina tipo scheda (Digitale/Analogico) dalla cella E2 (indice 1,4)
    try:
        model_val_raw_e2 = df.iloc[1, 4] # E2
        model_indicator_e2 = str(model_val_raw_e2).strip().upper()
        if not model_indicator_e2 : # Se E2 è vuota
            log_to_all(log_widget_gui, f"{base_filename}: IGNORATO - Cella E2 (tipo scheda) vuota.", "WARNING")
            return None, "E2 vuota", [], human_error_keys_this_file, dati_scheda_per_compilazione
    except IndexError:
        log_to_all(log_widget_gui, f"{base_filename}: IGNORATO - Cella E2 non trovata (IndexError). Il file potrebbe essere corrotto o avere una struttura inattesa.", "WARNING")
        return None, "E2 non trovata", [], human_error_keys_this_file, dati_scheda_per_compilazione
    except Exception as e_e2: # Altri errori leggendo E2
        val_e2_debug = model_val_raw_e2 if 'model_val_raw_e2' in locals() else "NON LETTA"
        log_to_all(log_widget_gui, f"{base_filename}: IGNORATO - Errore leggendo E2: '{str(e_e2)}'. Raw: '{val_e2_debug}'", "WARNING", exc_info=True)
        return None, f"Errore E2: {e_e2}", [], human_error_keys_this_file, dati_scheda_per_compilazione

    # Variabili per coordinate campi compilazione
    idx_sp_code_scheda, cella_sp_code_nome = None, None
    idx_odc, cell_odc_name = None, "N/A"
    idx_data_comp, cell_data_comp_name = None, "N/A"
    idx_pdl, cell_pdl_name = None, "N/A"
    idx_esec, cell_esec_name = None, "N/A"
    idx_sup_isab, cell_sup_isab_name = None, "N/A"
    idx_contr_coemi, cell_contr_coemi_name = None, "N/A"
    # Chiavi di errore per compilazione (specifiche per tipo scheda)
    key_odc_manc_loc, key_data_manc_loc, key_pdl_manc_loc, key_esec_manc_loc, key_sup_manc_loc, key_contr_manc_loc, key_contr_div_loc = ("N/A_LOCAL",)*7


    if "STRUMENTI DIGITALI" in model_indicator_e2:
        file_type = "digitale"; dr, dc = 44, 1 # dr,dc per data scheda (B45)
        dati_scheda_per_compilazione["file_type"] = "digitale"
        # Celle certificati per digitali (riga, colonna) - 0-based
        ccells = {
            'ids': [(17, 2), (17, 4), (17, 6)],      # C18, E18, G18
            'expiries': [(18, 2), (18, 4), (18, 6)], # C19, E19, G19
            'models': [(12, 2), (12, 4), (12, 6)],   # C13, E13, G13
            'ranges': [(15, 2), (15, 4), (15, 6)]    # C16, E16, G16
        }
        idx_sp_code_scheda = IDX_DIG_TIPOLOGIA_STRUM # N10
        cella_sp_code_nome = SCHEDA_DIG_CELL_TIPOLOGIA_STRUM
        
        idx_odc, cell_odc_name = IDX_DIG_ODC, SCHEDA_DIG_CELL_ODC
        idx_data_comp, cell_data_comp_name = IDX_DIG_DATA_COMPILAZIONE, SCHEDA_DIG_CELL_DATA_COMPILAZIONE
        idx_pdl, cell_pdl_name = IDX_DIG_PDL, SCHEDA_DIG_CELL_PDL
        idx_esec, cell_esec_name = IDX_DIG_ESECUTORE, SCHEDA_DIG_CELL_ESECUTORE
        idx_sup_isab, cell_sup_isab_name = IDX_DIG_SUPERVISORE_ISAB, SCHEDA_DIG_CELL_SUPERVISORE_ISAB
        idx_contr_coemi, cell_contr_coemi_name = IDX_DIG_CONTRATTO_COEMI, SCHEDA_DIG_CELL_CONTRATTO_COEMI
        key_odc_manc_loc, key_data_manc_loc, key_pdl_manc_loc = KEY_COMP_DIG_ODC_MANCANTE, KEY_COMP_DIG_DATA_COMP_MANCANTE, KEY_COMP_DIG_PDL_MANCANTE
        key_esec_manc_loc, key_sup_manc_loc = KEY_COMP_DIG_ESECUTORE_MANCANTE, KEY_COMP_DIG_SUPERVISORE_MANCANTE
        key_contr_manc_loc, key_contr_div_loc = KEY_COMP_DIG_CONTRATTO_MANCANTE, KEY_COMP_DIG_CONTRATTO_DIVERSO
        logger.debug(f"  {base_filename}: Rilevata scheda DIGITALE. Controllo campi compilatore specifici.")

    elif "STRUMENTI ANALOGICI" in model_indicator_e2:
        file_type = "analogico"; dr, dc = 49, 1 # dr,dc per data scheda (B50)
        dati_scheda_per_compilazione["file_type"] = "analogico"
        # Celle certificati per analogici (riga, colonna) - 0-based
        ccells = {
            'ids': [(42, 10), (43, 10), (44, 10)],   # K43, K44, K45
            'expiries': [(42, 12), (43, 12), (44, 12)],# M43, M44, M45
            'models': [(42, 0), (43, 0), (44, 0)],   # A43, A44, A45
            'ranges': [(42, 6), (43, 6), (44, 6)]    # G43, G44, G45
        }
        idx_sp_code_scheda = IDX_ANA_TIPOLOGIA_STRUM # N9
        cella_sp_code_nome = SCHEDA_ANA_CELL_TIPOLOGIA_STRUM
        
        idx_odc, cell_odc_name = IDX_ANA_ODC, SCHEDA_ANA_CELL_ODC
        idx_data_comp, cell_data_comp_name = IDX_ANA_DATA_COMPILAZIONE, SCHEDA_ANA_CELL_DATA_COMPILAZIONE
        idx_pdl, cell_pdl_name = IDX_ANA_PDL, SCHEDA_ANA_CELL_PDL
        idx_esec, cell_esec_name = IDX_ANA_ESECUTORE, SCHEDA_ANA_CELL_ESECUTORE
        idx_sup_isab, cell_sup_isab_name = IDX_ANA_SUPERVISORE_ISAB, SCHEDA_ANA_CELL_SUPERVISORE_ISAB
        idx_contr_coemi, cell_contr_coemi_name = IDX_ANA_CONTRATTO_COEMI, SCHEDA_ANA_CELL_CONTRATTO_COEMI
        key_odc_manc_loc, key_data_manc_loc, key_pdl_manc_loc = KEY_COMP_ANA_ODC_MANCANTE, KEY_COMP_ANA_DATA_COMP_MANCANTE, KEY_COMP_ANA_PDL_MANCANTE
        key_esec_manc_loc, key_sup_manc_loc = KEY_COMP_ANA_ESECUTORE_MANCANTE, KEY_COMP_ANA_SUPERVISORE_MANCANTE
        key_contr_manc_loc, key_contr_div_loc = KEY_COMP_ANA_CONTRATTO_MANCANTE, KEY_COMP_ANA_CONTRATTO_DIVERSO
        logger.debug(f"  {base_filename}: Rilevata scheda ANALOGICA. Controllo campi compilatore specifici.")
    else:
        log_to_all(log_widget_gui, f"{base_filename}: IGNORATO - Valore E2='{model_val_raw_e2}', tipo scheda non riconosciuto.", "WARNING")
        return None, f"Tipo non riconosciuto E2='{model_val_raw_e2}'", [], human_error_keys_this_file, dati_scheda_per_compilazione

    # Leggi SP Code (Tipologia Strumento dalla scheda)
    try:
        sp_code_raw_val = df.iloc[idx_sp_code_scheda[0], idx_sp_code_scheda[1]]
        sp_code_normalizzato_letto = normalize_sp_code(sp_code_raw_val)
        if not sp_code_normalizzato_letto or sp_code_normalizzato_letto.upper() == 'NAN': # "NAN" come stringa
            tipologia_strumento_scheda = f"SP MANCANTE ({cella_sp_code_nome} vuoto/NAN)"
            log_to_all(log_widget_gui, f"  {base_filename}: Cella tipo strumento ({cella_sp_code_nome}) VUOTA o 'NAN'. (Errore Umano)", "WARNING")
            human_error_keys_this_file.append(KEY_SP_VUOTO)
        else:
            tipologia_strumento_scheda = MAPPA_SP_TIPOLOGIA.get(sp_code_normalizzato_letto, f"SP NON MAPPATO: {sp_code_raw_val}")
            if tipologia_strumento_scheda.startswith("SP NON MAPPATO"):
                log_to_all(log_widget_gui, f"  {base_filename}: {tipologia_strumento_scheda} (da {cella_sp_code_nome}='{sp_code_raw_val}', Norm='{sp_code_normalizzato_letto}') - ANOMALIA: Aggiornare MAPPA_SP_TIPOLOGIA.", "WARNING")
    except IndexError:
        tipologia_strumento_scheda = f"Errore Indice {cella_sp_code_nome}"
        log_to_all(log_widget_gui, f"  {base_filename}: {tipologia_strumento_scheda} - ANOMALIA: Cella SP ({cella_sp_code_nome}) non trovata.", "WARNING")
        human_error_keys_this_file.append(KEY_SP_VUOTO) # Se non trovo la cella, considero l'SP vuoto
    except Exception as e_sp:
        tipologia_strumento_scheda = f"Errore SP {cella_sp_code_nome}"
        log_to_all(log_widget_gui, f"  {base_filename}: {tipologia_strumento_scheda}: {e_sp}", "WARNING", exc_info=True)
        human_error_keys_this_file.append(KEY_SP_VUOTO)


    # Leggi Modello L9 (solo per analogici)
    if file_type == "analogico":
        try:
            modello_l9_raw_value = df.iloc[IDX_ANA_MODELLO_STRUM[0], IDX_ANA_MODELLO_STRUM[1]] # L9
            modello_l9_temp = str(modello_l9_raw_value).strip().upper()
            # Pre-normalizzazioni specifiche per L9 prima di usare MAPPA_L9_SOTTOTIPO_NORMALIZZATA
            modello_l9_temp = modello_l9_temp.replace('ΔP', 'DP').replace('DELTA P', 'DP').replace("SKINPOINT", "SKIN POINT")
            modello_l9_scheda_normalizzato = " ".join(modello_l9_temp.split()) # Normalizza spazi
            
            if not modello_l9_scheda_normalizzato or modello_l9_scheda_normalizzato == 'NAN':
                log_to_all(log_widget_gui, f"  {base_filename}: Modello L9 ({SCHEDA_ANA_CELL_MODELLO_STRUM}) VUOTO o 'NAN'. (Errore Umano)", "WARNING")
                human_error_keys_this_file.append(KEY_L9_VUOTO); modello_l9_scheda_normalizzato = "L9 VUOTO" # Placeholder per report
            elif modello_l9_scheda_normalizzato == "SKIN POINT": # Caso specifico di incompletezza
                log_to_all(log_widget_gui, f"  {base_filename}: Modello L9 '{modello_l9_raw_value}' incompleto (es. manca K/J). (Errore Umano)", "WARNING")
                human_error_keys_this_file.append(KEY_L9_SKINPOINT_INCOMPLETO)
        except IndexError:
            modello_l9_scheda_normalizzato = f"Errore Indice {SCHEDA_ANA_CELL_MODELLO_STRUM}"
            log_to_all(log_widget_gui, f"  {base_filename}: {modello_l9_scheda_normalizzato} - ANOMALIA.", "WARNING"); human_error_keys_this_file.append(KEY_L9_VUOTO)
        except Exception as e_l9:
            modello_l9_scheda_normalizzato = f"Errore L9 {SCHEDA_ANA_CELL_MODELLO_STRUM}"
            log_to_all(log_widget_gui, f"  {base_filename}: {modello_l9_scheda_normalizzato}: {e_l9}", "WARNING", exc_info=True); human_error_keys_this_file.append(KEY_L9_VUOTO)

    log_to_all(log_widget_gui, f"  {base_filename}: Tipo Scheda='{file_type}', Tip.Strum.(SP)='{tipologia_strumento_scheda}', SP Code Norm='{sp_code_normalizzato_letto}', Mod.L9(Ana)='{modello_l9_scheda_normalizzato if file_type == 'analogico' else 'N/A'}'", "DEBUG")
    # Data scheda (da B45 o B50)
    card_date_val_raw = "N/D"
    try:
        card_date_val_raw = df.iloc[dr, dc] 
        card_date = parse_date_robust(card_date_val_raw, base_filename)
        assert card_date is not None, f"Data scheda '{card_date_val_raw}' non interpretabile."
    except (IndexError, AssertionError) as e_date:
        log_to_all(log_widget_gui, f"{base_filename} ({file_type}): IGNORATO - Errore data scheda (r{dr+1},c{dc+1}): Val='{card_date_val_raw}', Err='{e_date}'", "WARNING")
        return None, f"Errore data: {e_date}", [], human_error_keys_this_file, dati_scheda_per_compilazione
    except Exception as e_date_unexp: # Catch-all per altri errori sulla data
        log_to_all(log_widget_gui, f"{base_filename} ({file_type}): IGNORATO - Errore imprevisto data scheda: {e_date_unexp}", "WARNING", exc_info=True)
        return None, f"Errore data: {e_date_unexp}", [], human_error_keys_this_file, dati_scheda_per_compilazione

    # Lettura e verifica campi per compilazione automatica
    logger.debug(f"  {base_filename}: Inizio analisi campi compilatore (Tipo: {file_type}).")
    try:
        val_odc = df.iloc[idx_odc[0], idx_odc[1]]
        dati_scheda_per_compilazione["odc_val_scheda"] = str(val_odc).strip() if not pd.isna(val_odc) else None
        is_odc_empty = is_cell_value_empty(val_odc)
        logger.debug(f"    Campo ODC ({cell_odc_name} -> {idx_odc}): Valore='{val_odc}', Vuoto={is_odc_empty}")
        if is_odc_empty: 
            human_error_keys_this_file.append(key_odc_manc_loc)
            dati_scheda_per_compilazione["campi_mancanti"].add(key_odc_manc_loc)


        val_data_comp = df.iloc[idx_data_comp[0], idx_data_comp[1]]
        is_data_comp_empty = is_cell_value_empty(val_data_comp)
        logger.debug(f"    Campo Data Comp. ({cell_data_comp_name} -> {idx_data_comp}): Valore='{val_data_comp}', Vuoto={is_data_comp_empty}")
        if is_data_comp_empty:
            human_error_keys_this_file.append(key_data_manc_loc)
            dati_scheda_per_compilazione["campi_mancanti"].add(key_data_manc_loc)
        
        val_pdl = df.iloc[idx_pdl[0], idx_pdl[1]]
        dati_scheda_per_compilazione["pdl_val"] = str(val_pdl).strip() if not pd.isna(val_pdl) else None
        is_pdl_empty = is_cell_value_empty(val_pdl)
        logger.debug(f"    Campo PDL ({cell_pdl_name} -> {idx_pdl}): Valore='{val_pdl}', Vuoto={is_pdl_empty}")
        if is_pdl_empty:
            human_error_keys_this_file.append(key_pdl_manc_loc)
            dati_scheda_per_compilazione["campi_mancanti"].add(key_pdl_manc_loc)


        val_esec = df.iloc[idx_esec[0], idx_esec[1]]
        is_esec_empty = is_cell_value_empty(val_esec)
        logger.debug(f"    Campo Esecutore ({cell_esec_name} -> {idx_esec}): Valore='{val_esec}', Vuoto={is_esec_empty}")
        if is_esec_empty:
            human_error_keys_this_file.append(key_esec_manc_loc)
            dati_scheda_per_compilazione["campi_mancanti"].add(key_esec_manc_loc)
        
        val_sup_isab = df.iloc[idx_sup_isab[0], idx_sup_isab[1]]
        is_sup_isab_empty = is_cell_value_empty(val_sup_isab)
        logger.debug(f"    Campo Supervisore ISAB ({cell_sup_isab_name} -> {idx_sup_isab}): Valore='{val_sup_isab}', Vuoto={is_sup_isab_empty}")
        if is_sup_isab_empty:
            human_error_keys_this_file.append(key_sup_manc_loc)
            dati_scheda_per_compilazione["campi_mancanti"].add(key_sup_manc_loc)

        val_contr_coemi = df.iloc[idx_contr_coemi[0], idx_contr_coemi[1]]
        is_contr_coemi_empty = is_cell_value_empty(val_contr_coemi)
        contr_coemi_val_stripped = str(val_contr_coemi).strip() if not pd.isna(val_contr_coemi) else ""
        logger.debug(f"    Campo Contratto Coemi ({cell_contr_coemi_name} -> {idx_contr_coemi}): Valore='{val_contr_coemi}', Stripped='{contr_coemi_val_stripped}', Vuoto={is_contr_coemi_empty}")
        if is_contr_coemi_empty:
            human_error_keys_this_file.append(key_contr_manc_loc)
            dati_scheda_per_compilazione["campi_mancanti"].add(key_contr_manc_loc)
            logger.debug(f"      -> Contratto Coemi MANCANTE (key: {key_contr_manc_loc})")
        elif not (contr_coemi_val_stripped == VALORE_ATTESO_CONTRATTO_COEMI or \
                  contr_coemi_val_stripped == VALORE_ATTESO_CONTRATTO_COEMI_VARIANTE_NUMERICA):
            human_error_keys_this_file.append(key_contr_div_loc)
            dati_scheda_per_compilazione["campi_mancanti"].add(key_contr_div_loc) # Aggiungi anche se diverso per possibile correzione
            logger.debug(f"      -> Contratto Coemi DIVERSO (key: {key_contr_div_loc}). Trovato: '{contr_coemi_val_stripped}', Attesi: '{VALORE_ATTESO_CONTRATTO_COEMI}' o '{VALORE_ATTESO_CONTRATTO_COEMI_VARIANTE_NUMERICA}'")

    except IndexError:
        log_to_all(log_widget_gui,f"  {base_filename}: Errore Indice lettura campi per compilatore (ODC, Data, PDL, etc.). Alcuni campi potrebbero mancare nel file.", "WARNING")
        human_error_keys_this_file.append(KEY_COMP_CAMPI_MANCANTI_NON_LEGGIBILI)
        dati_scheda_per_compilazione["campi_mancanti"].add(KEY_COMP_CAMPI_MANCANTI_NON_LEGGIBILI) # Segna errore generico
    except Exception as e_comp_fields:
        log_to_all(log_widget_gui,f"  {base_filename}: Errore lettura campi per compilatore: {e_comp_fields}", "ERROR", exc_info=True)
        human_error_keys_this_file.append(KEY_COMP_CAMPI_MANCANTI_NON_LEGGIBILI)
        dati_scheda_per_compilazione["campi_mancanti"].add(KEY_COMP_CAMPI_MANCANTI_NON_LEGGIBILI)


    # Lettura Range/UM e validazione specifica (se tipo scheda e SP sono validi)
    if file_type == "analogico":
        try:
            range_ing_raw = df.iloc[IDX_ANA_RANGE_INGRESSO[0], IDX_ANA_RANGE_INGRESSO[1]]
            um_ing_raw = df.iloc[IDX_ANA_UM_INGRESSO[0], IDX_ANA_UM_INGRESSO[1]]
            range_usc_raw = df.iloc[IDX_ANA_RANGE_USCITA[0], IDX_ANA_RANGE_USCITA[1]]
            um_usc_raw = df.iloc[IDX_ANA_UM_USCITA[0], IDX_ANA_UM_USCITA[1]]
            range_dcs_raw = df.iloc[IDX_ANA_RANGE_DCS[0], IDX_ANA_RANGE_DCS[1]]
            um_dcs_raw = df.iloc[IDX_ANA_UM_DCS[0], IDX_ANA_UM_DCS[1]]

            range_ing_norm = normalize_range_string(range_ing_raw)
            um_ing_norm = normalize_um(um_ing_raw)
            range_usc_norm = normalize_range_string(range_usc_raw)
            um_usc_norm = normalize_um(um_usc_raw)
            range_dcs_norm = normalize_range_string(range_dcs_raw)
            um_dcs_norm = normalize_um(um_dcs_raw)
            logger.debug(f"  {base_filename} (ANA) Raw: ING='{range_ing_raw}'({um_ing_raw}), USC='{range_usc_raw}'({um_usc_raw}), DCS='{range_dcs_raw}'({um_dcs_raw})")
            logger.debug(f"  {base_filename} (ANA) Norm: ING='{range_ing_norm}'({um_ing_norm}), USC='{range_usc_norm}'({um_usc_norm}), DCS='{range_dcs_norm}'({um_dcs_norm})")
        except IndexError:
            log_to_all(log_widget_gui,f"  {base_filename}: Errore Indice Range/UM analogiche preliminare. Alcune celle potrebbero mancare.","WARNING"); human_error_keys_this_file.append(KEY_CELL_RANGE_UM_NON_LEGGIBILE)
        except Exception as e_ana_read: # Altri errori inattesi
            log_to_all(log_widget_gui,f"  {base_filename}: Errore lettura Range/UM analogiche preliminare: {e_ana_read}","ERROR",exc_info=True); human_error_keys_this_file.append(KEY_CELL_RANGE_UM_NON_LEGGIBILE)
    
    # Validazione specifica Range/UM basata su Tipologia Strumento e Modello L9
    # Solo se non ci sono stati errori critici prima (SP valido, celle Range/UM leggibili)
    if not tipologia_strumento_scheda.startswith(("SP MANCANTE", "SP NON MAPPATO", "Errore Indice", "Errore SP", "N/D_INIZIALE")):
        if file_type == "analogico":
            if KEY_CELL_RANGE_UM_NON_LEGGIBILE not in human_error_keys_this_file: # Solo se le celle sono state lette
                try:
                    if tipologia_strumento_scheda == "TEMPERATURA":
                        if modello_l9_scheda_normalizzato == "CONVERTITORE":
                            if um_ing_norm != um_dcs_norm: human_error_keys_this_file.append(KEY_ERR_ANA_TEMP_CONV_C9F9_UM_DIVERSE)
                            if um_usc_norm != UM_MA_NORMALIZZATA: human_error_keys_this_file.append(KEY_ERR_ANA_TEMP_CONV_F12_UM_NON_MA)
                            if range_ing_norm != range_dcs_norm: human_error_keys_this_file.append(KEY_ERR_ANA_TEMP_CONV_A9D9_RANGE_DIVERSI)
                            if not is_valid_expected_range(range_usc_norm, RANGE_4_20_NORMALIZZATO): human_error_keys_this_file.append(KEY_ERR_ANA_TEMP_CONV_D12_RANGE_NON_4_20)
                        elif not modello_l9_scheda_normalizzato.startswith("L9 VUOTO"): # Per RTD, TC (non convertitori)
                            if not (um_ing_norm == um_dcs_norm and um_dcs_norm == um_usc_norm): human_error_keys_this_file.append(KEY_ERR_ANA_TEMP_NOCONV_UM_NON_COINCIDENTI)
                            if not (range_ing_norm == range_dcs_norm and range_dcs_norm == range_usc_norm): human_error_keys_this_file.append(KEY_ERR_ANA_TEMP_NOCONV_RANGE_NON_COINCIDENTI)
                    elif tipologia_strumento_scheda == "LIVELLO":
                        if modello_l9_scheda_normalizzato == "DP":
                            if not is_um_pressione_valida(um_ing_norm): human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_DP_C9_UM_NON_PRESSIONE)
                            if not is_valid_expected_range(range_dcs_norm, RANGE_0_100_NORMALIZZATO): human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_DP_D9_RANGE_NON_0_100)
                            if um_dcs_norm != UM_PERCENTO_NORMALIZZATA: human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_DP_F9_UM_NON_PERCENTO)
                            if not is_valid_expected_range(range_usc_norm, RANGE_4_20_NORMALIZZATO): human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_DP_D12_RANGE_NON_4_20)
                            if um_usc_norm != UM_MA_NORMALIZZATA: human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_DP_F12_UM_NON_MA)
                        elif "BARRA DI TORSIONE" in modello_l9_scheda_normalizzato or \
                             ("TORSIONALE" in modello_l9_scheda_normalizzato and "PNEUMATICO" not in modello_l9_scheda_normalizzato and "LOCALE" not in modello_l9_scheda_normalizzato and "CAPILLARE" not in modello_l9_scheda_normalizzato) : # Torsionale Elettronico generico
                            if not (um_ing_norm == UM_MMH2O_NORMALIZZATA or um_ing_norm == UM_MM_NORMALIZZATA): human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_TORS_C9_UM_INVALIDA)
                            if not is_valid_expected_range(range_dcs_norm, RANGE_0_100_NORMALIZZATO): human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_TORS_D9_RANGE_NON_0_100)
                            if um_dcs_norm != UM_PERCENTO_NORMALIZZATA: human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_TORS_F9_UM_NON_PERCENTO)
                            # Per Torsionale Elettronico, ci si aspetta uscita 4-20mA
                            if "ELETTRONICO" in modello_l9_scheda_normalizzato or "ELETTRONICA" in modello_l9_scheda_normalizzato or ("TORSIONALE" == modello_l9_scheda_normalizzato and um_usc_norm == UM_MA_NORMALIZZATA): # Assumiamo 4-20mA se elettronico o solo "TORSIONALE" con mA
                                if not is_valid_expected_range(range_usc_norm, RANGE_4_20_NORMALIZZATO): human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_TORS_ELETTR_D12_RANGE_NON_4_20)
                                if um_usc_norm != UM_MA_NORMALIZZATA: human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_TORS_ELETTR_F12_UM_NON_MA)
                        elif "TORSIONALE LOCALE" in modello_l9_scheda_normalizzato : # o TORSIONALE PNEUMATICO (che spesso è locale)
                            # Per locale, D12 e F12 dovrebbero essere vuoti o indicare unità locali come PSI se pneumatici.
                            # Se F12 è PSI, D12 potrebbe avere un range. Altrimenti vuoti.
                            if not (um_usc_norm == UM_PSI_NORMALIZZATA or is_cell_value_empty(um_usc_raw)): # F12 non è PSI e non è vuota
                                human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_TORS_LOCALE_F12_UM_NON_VUOTA)
                            if um_usc_norm != UM_PSI_NORMALIZZATA and not is_cell_value_empty(range_usc_raw): # Se F12 non è PSI, D12 deve essere vuota
                                human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_TORS_LOCALE_D12_RANGE_NON_VUOTO)
                        elif modello_l9_scheda_normalizzato in ["RADAR", "ULTRASUONI", "ONDA GUIDATA"]:
                            if not is_valid_expected_range(range_dcs_norm, RANGE_0_100_NORMALIZZATO): human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_RADARULTR_D9_RANGE_NON_0_100)
                            if um_dcs_norm != UM_PERCENTO_NORMALIZZATA: human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_RADARULTR_F9_UM_NON_PERCENTO)
                            if not is_valid_expected_range(range_usc_norm, RANGE_4_20_NORMALIZZATO): human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_RADARULTR_D12_RANGE_NON_4_20)
                            if um_usc_norm != UM_MA_NORMALIZZATA: human_error_keys_this_file.append(KEY_ERR_ANA_LIVELLO_RADARULTR_F12_UM_NON_MA)
                    elif tipologia_strumento_scheda == "PRESSIONE":
                        if modello_l9_scheda_normalizzato in ["DP", "TX", "TX PRESSIONE", "TX DI PRESSIONE", "CAPILLARE"]:
                            if um_ing_norm != um_dcs_norm: human_error_keys_this_file.append(KEY_ERR_ANA_PRESS_DP_TX_C9F9_UM_DIVERSE)
                            if um_usc_norm != UM_MA_NORMALIZZATA: human_error_keys_this_file.append(KEY_ERR_ANA_PRESS_DP_TX_F12_UM_NON_MA)
                            if range_ing_norm != range_dcs_norm: human_error_keys_this_file.append(KEY_ERR_ANA_PRESS_DP_TX_A9D9_RANGE_DIVERSI)
                            if not is_valid_expected_range(range_usc_norm, RANGE_4_20_NORMALIZZATO): human_error_keys_this_file.append(KEY_ERR_ANA_PRESS_DP_TX_D12_RANGE_NON_4_20)
                    elif tipologia_strumento_scheda == "PORTATA":
                         if modello_l9_scheda_normalizzato == "DP" or "CAPILLARE" in modello_l9_scheda_normalizzato: # O altri modelli di portata con uscita 4-20 mA
                            if not is_valid_expected_range(range_usc_norm, RANGE_4_20_NORMALIZZATO): human_error_keys_this_file.append(KEY_ERR_ANA_PORTATA_DP_D12_RANGE_NON_4_20)
                            if um_usc_norm != UM_MA_NORMALIZZATA: human_error_keys_this_file.append(KEY_ERR_ANA_PORTATA_DP_F12_UM_NON_MA)
                except IndexError: # Se una cella Range/UM non è trovata durante la validazione specifica
                    log_to_all(log_widget_gui,f"  {base_filename}: Errore Indice durante validazione Range/UM analogiche specifiche.","WARNING"); human_error_keys_this_file.append(KEY_CELL_RANGE_UM_NON_LEGGIBILE)
                except Exception as e_ana: # Altri errori inattesi
                    log_to_all(log_widget_gui,f"  {base_filename}: Errore durante validazione Range/UM analogiche specifiche: {e_ana}","ERROR",exc_info=True); human_error_keys_this_file.append(KEY_CELL_RANGE_UM_NON_LEGGIBILE)
        elif file_type == "digitale":
            try:
                range_um_proc_raw = df.iloc[IDX_DIG_RANGE_UM_PROCESSO[0], IDX_DIG_RANGE_UM_PROCESSO[1]] # D22
                # Per digitali, l'UM è spesso parte della stringa del range. Cerchiamo di estrarla.
                # Questa è una semplificazione: l'UM potrebbe essere alla fine della stringa.
                # Esempio: "0-100 barg". Vogliamo "barg".
                # Una regex più robusta potrebbe essere necessaria se i formati variano molto.
                um_proc_norm_parts = normalize_um(range_um_proc_raw).split() # Splitta e normalizza
                um_proc_norm = um_proc_norm_parts[-1] if um_proc_norm_parts else "" # Prendi l'ultima parte come potenziale UM

                logger.debug(f"  {base_filename} (DIG) UM Processo ({SCHEDA_DIG_CELL_RANGE_UM_PROCESSO}) Raw:'{range_um_proc_raw}', Norm(tentativo):'{um_proc_norm}'")
                if tipologia_strumento_scheda == "PRESSIONE":
                    if not is_um_pressione_valida(um_proc_norm): # Tenta di validare l'UM estratta
                         # Se l'UM estratta non è valida, prova a normalizzare l'intera stringa e vedere se corrisponde a un'UM di pressione
                         # Questo gestisce casi come D22 = "barg" o "MMH2O" senza un range numerico.
                        full_string_as_um_norm = normalize_um(range_um_proc_raw)
                        if not is_um_pressione_valida(full_string_as_um_norm):
                            human_error_keys_this_file.append(KEY_ERR_DIG_PRESS_D22_UM_NON_PRESSIONE)
                elif tipologia_strumento_scheda == "LIVELLO":
                    # Per livello digitale, ci si aspetta tipicamente '%'
                    # Se D22 contiene "0-100 %" o solo "%"
                    if UM_PERCENTO_NORMALIZZATA not in normalize_um(range_um_proc_raw):
                        human_error_keys_this_file.append(KEY_ERR_DIG_LIVELLO_D22_UM_NON_PERCENTO)
            except IndexError:
                log_to_all(log_widget_gui,f"  {base_filename}: Errore Indice {SCHEDA_DIG_CELL_RANGE_UM_PROCESSO} (digitale).","WARNING"); human_error_keys_this_file.append(KEY_CELL_RANGE_UM_NON_LEGGIBILE)
            except Exception as e_dig:
                 log_to_all(log_widget_gui,f"  {base_filename}: Errore validazione Range/UM digitali: {e_dig}","ERROR",exc_info=True); human_error_keys_this_file.append(KEY_CELL_RANGE_UM_NON_LEGGIBILE)

        # Log riepilogativo errori Range/UM per il file corrente
        current_file_range_um_errors = [k for k in human_error_keys_this_file if k.startswith(("KEY_ERR_ANA_","KEY_ERR_DIG_")) or k==KEY_CELL_RANGE_UM_NON_LEGGIBILE]
        if current_file_range_um_errors: # Se ci sono errori di questo tipo
            unique_errors_for_this_log = len(list(set(current_file_range_um_errors))) # Conta i tipi unici di errore
            log_to_all(log_widget_gui, f"  {base_filename}: Rilevati {unique_errors_for_this_log} tipi di errori di compilazione Range/UM.", "WARNING")


    # Estrazione dati certificati usati
    for i in range(3): # Itera sui 3 slot certificati
        try:
            cert_id_raw = df.iloc[ccells['ids'][i][0], ccells['ids'][i][1]]
            cert_id = str(cert_id_raw).strip()
            # Salta se ID certificato è vuoto, NaN, o solo spazi
            if not cert_id or cert_id.lower() == 'nan' or cert_id == "" or cert_id.isspace():
                continue

            exp_raw = df.iloc[ccells['expiries'][i][0], ccells['expiries'][i][1]]
            mod_raw_card = df.iloc[ccells['models'][i][0], ccells['models'][i][1]]
            ran_raw_card = df.iloc[ccells['ranges'][i][0], ccells['ranges'][i][1]]

            cert_exp_dt = parse_date_robust(exp_raw, base_filename)
            is_exp = bool(cert_exp_dt and card_date and cert_exp_dt < card_date) # True se scaduto all'uso

            # Verifica congruità
            is_congr = None; congr_notes = "Verifica non iniziata."; mod_camp_reg = "N/D_NonTrovatoRegistro"
            used_before_em = False; sott_l9_eff = "N/A" # Sottotipo L9 effettivo dopo mappatura

            if not strumenti_campione_global_list:
                is_congr, congr_notes = None, "Registro campioni non disponibile."
            else:
                found_camp = next((sc for sc in strumenti_campione_global_list if sc['id_certificato'] == cert_id), None)
                if not found_camp:
                    is_congr, congr_notes = None, f"Cert.ID '{cert_id}' NON TROVATO nel registro."
                else:
                    mod_camp_reg = found_camp.get('modello_strumento', "N/D_ModMancanteRegistro").strip().upper()
                    dt_em_camp = found_camp.get('data_emissione') # Oggetto datetime o None
                    
                    if dt_em_camp and card_date and card_date < dt_em_camp:
                        used_before_em = True; is_congr = False # Se usato prima dell'emissione, è incongruo
                        congr_notes = f"Dettaglio: errato per EMISSIONE. Cert.'{cert_id}'({mod_camp_reg}) usato il {card_date:%d/%m/%Y} ma il certificato è stato emesso il {dt_em_camp:%d/%m/%Y}."
                    
                    if not used_before_em: # Prosegui con altre verifiche solo se non usato prima dell'emissione
                        if mod_camp_reg.startswith("N/D_"): # Modello campione non definito nel registro
                            is_congr, congr_notes = None, f"Cert.ID '{cert_id}' trovato, ma modello campione N/D nel registro."
                        elif tipologia_strumento_scheda.startswith(("SP MANCANTE", "SP NON MAPPATO", "Errore Indice", "Errore SP", "N/D_INIZIALE")):
                            is_congr, congr_notes = None, f"Tipologia strumento scheda ('{tipologia_strumento_scheda}') non valida o non mappata."
                        elif tipologia_strumento_scheda not in REGOLE_CONGRUITA_CERTIFICATI_NORMALIZZATE:
                            is_congr, congr_notes = None, f"Regole congruità non definite per tipologia '{tipologia_strumento_scheda}'."
                        else:
                            reg_tip = REGOLE_CONGRUITA_CERTIFICATI_NORMALIZZATE[tipologia_strumento_scheda]
                            
                            # Determina sottotipo L9 effettivo per analogici
                            if file_type == 'analogico' and modello_l9_scheda_normalizzato not in ["N/A", "L9 VUOTO", "SKIN POINT"] and not modello_l9_scheda_normalizzato.startswith(("Errore Indice", "Errore L9")):
                                if modello_l9_scheda_normalizzato in MAPPA_L9_SOTTOTIPO_NORMALIZZATA:
                                    poss_l9_val = MAPPA_L9_SOTTOTIPO_NORMALIZZATA[modello_l9_scheda_normalizzato]
                                    # Il valore può essere stringa o lista, assicurati sia una lista
                                    poss_l9_list = [poss_l9_val] if isinstance(poss_l9_val, str) else poss_l9_val
                                    for cand_l9 in poss_l9_list:
                                        # Un sottotipo L9 è valido se la tipologia SP è contenuta nel nome del sottotipo L9 (es. SP=TEMP, L9_sott=TEMP_TERMOCOPPIA)
                                        # o se sono uguali (es. SP=PRESSIONE, L9_sott=PRESSIONE)
                                        if tipologia_strumento_scheda in cand_l9 or cand_l9 == tipologia_strumento_scheda:
                                            sott_l9_eff = cand_l9; break
                                elif modello_l9_scheda_normalizzato and sott_l9_eff == "N/A": # Match parziale se non esatto
                                    matched_key_len = 0
                                    for l9_key_map in sorted(MAPPA_L9_SOTTOTIPO_NORMALIZZATA.keys(), key=len, reverse=True):
                                        if l9_key_map in modello_l9_scheda_normalizzato:
                                            if len(l9_key_map) > matched_key_len:
                                                poss_l9_cand_val = MAPPA_L9_SOTTOTIPO_NORMALIZZATA[l9_key_map]
                                                poss_l9_cand_list = [poss_l9_cand_val] if isinstance(poss_l9_cand_val, str) else poss_l9_cand_val
                                                for cand_st_partial in poss_l9_cand_list:
                                                    if tipologia_strumento_scheda in cand_st_partial or cand_st_partial == tipologia_strumento_scheda:
                                                        sott_l9_eff = cand_st_partial
                                                        matched_key_len = len(l9_key_map)
                                                        break 
                                            if sott_l9_eff != "N/A" and matched_key_len == len(l9_key_map): break # Trovato il match più lungo possibile

                            # Logica di congruità effettiva
                            is_congr_prov, congr_notes_prov = False, f"INCONGRUO (default): '{mod_camp_reg}' per {tipologia_strumento_scheda} (L9:'{modello_l9_scheda_normalizzato}',SottL9Eff:'{sott_l9_eff}')."

                            # Caso speciale: LIVELLO con MANOMETRO DIGITALE
                            if tipologia_strumento_scheda == "LIVELLO" and mod_camp_reg == "MANOMETRO DIGITALE":
                                if file_type == 'digitale': # Per digitali è OK
                                    is_congr_prov, congr_notes_prov = True, "OK (LIV digitale con MAN DIG)."
                                elif file_type == 'analogico':
                                    mod_l9_norm_local_liv = modello_l9_scheda_normalizzato # Usare il normalizzato dalla scheda
                                    # Condizioni specifiche per analogici di livello con manometro digitale
                                    cond_dp_liv = (mod_l9_norm_local_liv == "DP")
                                    cond_tors_pneu_liv = ("TORSIONALE PNEUMATICO" in mod_l9_norm_local_liv and um_usc_norm == UM_PSI_NORMALIZZATA) # Uscita PSI
                                    cond_tors_locale_liv = ("TORSIONALE LOCALE" in mod_l9_norm_local_liv and um_usc_norm == UM_PSI_NORMALIZZATA) # Uscita PSI
                                    cond_capillare_liv = ("CAPILLARE" in mod_l9_norm_local_liv and um_usc_norm == UM_MA_NORMALIZZATA ) # Uscita mA

                                    if cond_dp_liv:
                                        is_congr_prov, congr_notes_prov = True, "OK (LIV DP analogico con MAN DIG)."
                                    elif cond_tors_pneu_liv:
                                        is_congr_prov, congr_notes_prov = True, f"OK (LIV {mod_l9_norm_local_liv} con MAN DIG e UM Uscita ({SCHEDA_ANA_CELL_UM_USCITA})='{um_usc_norm.upper()}')."
                                    elif cond_tors_locale_liv:
                                         is_congr_prov, congr_notes_prov = True, f"OK (LIV {mod_l9_norm_local_liv} con MAN DIG e UM Uscita ({SCHEDA_ANA_CELL_UM_USCITA})='{um_usc_norm.upper()}')."
                                    elif cond_capillare_liv: # Un TX di livello con capillari e uscita mA può essere testato come un TX di pressione
                                        is_congr_prov, congr_notes_prov = True, f"OK (LIV {mod_l9_norm_local_liv} con MAN DIG e UM Uscita ({SCHEDA_ANA_CELL_UM_USCITA})='{um_usc_norm.upper()}')."
                                    else:
                                        error_details_liv_man_list = []
                                        if "TORSIONALE PNEUMATICO" in mod_l9_norm_local_liv and um_usc_norm != UM_PSI_NORMALIZZATA:
                                            error_details_liv_man_list.append(f"per L9 '{mod_l9_norm_local_liv}' UM Uscita ({SCHEDA_ANA_CELL_UM_USCITA}) deve essere '{UM_PSI_NORMALIZZATA.upper()}' (trovato: '{um_usc_norm.upper() if um_usc_norm else 'VUOTO'}')")
                                        if "TORSIONALE LOCALE" in mod_l9_norm_local_liv and um_usc_norm != UM_PSI_NORMALIZZATA:
                                            error_details_liv_man_list.append(f"per L9 '{mod_l9_norm_local_liv}' UM Uscita ({SCHEDA_ANA_CELL_UM_USCITA}) deve essere '{UM_PSI_NORMALIZZATA.upper()}' (trovato: '{um_usc_norm.upper() if um_usc_norm else 'VUOTO'}')")
                                        if "CAPILLARE" in mod_l9_norm_local_liv and um_usc_norm != UM_MA_NORMALIZZATA :
                                             error_details_liv_man_list.append(f"per L9 '{mod_l9_norm_local_liv}' UM Uscita ({SCHEDA_ANA_CELL_UM_USCITA}) deve essere '{UM_MA_NORMALIZZATA.upper()}' (trovato: '{um_usc_norm.upper() if um_usc_norm else 'VUOTO'}')")

                                        allowed_l9_for_man_dig_str = "'DP', 'TORSIONALE PNEUMATICO' (con F12='PSI'), 'TORSIONALE LOCALE' (con F12='PSI'), 'CAPILLARE' (con F12='mA')"
                                        reason_str = "; ".join(error_details_liv_man_list) if error_details_liv_man_list else f"L9='{mod_l9_norm_local_liv}' non supportato con MAN DIG. Ammessi: {allowed_l9_for_man_dig_str}"
                                        is_congr_prov, congr_notes_prov = False, f"INCONGRUO: MAN DIG per LIV analogico. {reason_str}."
                            
                            # Altre regole di congruità (se non è il caso speciale sopra)
                            elif "eccezioni_l9_incongrui" in reg_tip and sott_l9_eff != "N/A" and sott_l9_eff in reg_tip["eccezioni_l9_incongrui"] and mod_camp_reg in reg_tip["eccezioni_l9_incongrui"][sott_l9_eff]:
                                is_congr_prov, congr_notes_prov = False, f"INCONGRUO (eccL9):'{mod_camp_reg}' per {tipologia_strumento_scheda}({sott_l9_eff})."
                            elif mod_camp_reg in reg_tip.get("modelli_campione_incongrui", []):
                                # Se è incongruo in generale, controlla se il sottotipo L9 lo rende congruo
                                if "sottotipi_l9" in reg_tip and sott_l9_eff != "N/A" and sott_l9_eff in reg_tip["sottotipi_l9"] and mod_camp_reg in reg_tip["sottotipi_l9"][sott_l9_eff]:
                                    is_congr_prov, congr_notes_prov = True, f"OK (sottL9 sovrascrive incongruo gen.):'{mod_camp_reg}' per {tipologia_strumento_scheda}({sott_l9_eff})."
                                else:
                                    is_congr_prov, congr_notes_prov = False, f"INCONGRUO (lista gen):'{mod_camp_reg}' per {tipologia_strumento_scheda}."
                            elif "sottotipi_l9" in reg_tip and sott_l9_eff != "N/A" and sott_l9_eff in reg_tip["sottotipi_l9"] and mod_camp_reg in reg_tip["sottotipi_l9"][sott_l9_eff]:
                                is_congr_prov, congr_notes_prov = True, f"OK (sottL9):'{mod_camp_reg}' per {tipologia_strumento_scheda}({sott_l9_eff})."
                            elif mod_camp_reg in reg_tip.get("modelli_campione_congrui", []):
                                is_congr_prov, congr_notes_prov = True, "OK (regole base)."
                            
                            # Caso specifico per TEMPERATURA/TERMOCOPPIA (senza K/J) e MULTIMETRO
                            # Se L9 è solo "TERMOCOPPIA" e si usa un MULTIMETRO, è OK se le regole per TEMPERATURA_TERMOCOPPIA lo ammettono
                            if not is_congr_prov and \
                               tipologia_strumento_scheda == "TEMPERATURA" and \
                               mod_camp_reg == "MULTIMETRO DIGITALE" and \
                               modello_l9_scheda_normalizzato == "TERMOCOPPIA" and \
                               sott_l9_eff == "N/A": # sott_l9_eff sarebbe TEMPERATURA_TERMOCOPPIA se L9 fosse completo
                                if "sottotipi_l9" in reg_tip and \
                                   "TEMPERATURA_TERMOCOPPIA" in reg_tip["sottotipi_l9"] and \
                                   mod_camp_reg in reg_tip["sottotipi_l9"]["TEMPERATURA_TERMOCOPPIA"]:
                                    # Consideralo congruo ma con una nota che L9 è incompleto
                                    # is_congr_prov = True # Non marcarlo come True qui, lasciamolo come anomalia L9
                                    congr_notes_prov = "Dettaglio: L9='TERMOCOPPIA' incompleto (manca tipo K/J), MULTIMETRO sarebbe OK per Termocoppia completa."


                            is_congr, congr_notes = is_congr_prov, congr_notes_prov
            
            # Log risultato congruità
            log_lvl = "DEBUG" if is_congr is None else "ERROR" if used_before_em else "WARNING" if is_congr is False else "INFO"
            log_to_all(log_widget_gui,f"    {base_filename}(Slot {i+1}): Cert.ID '{cert_id}', Mod.Camp(Reg):'{mod_camp_reg}'. Scad:'{exp_raw}'. Congr:{is_congr}. Note: {congr_notes}",log_lvl)
            
            extracted_certs_data.append({
                'file_name': base_filename, 'file_path': file_path,
                'card_type': file_type, 'card_date': card_date,
                'certificate_id': cert_id,
                'certificate_expiry_raw': exp_raw, 'certificate_expiry': cert_exp_dt,
                'instrument_model_on_card': str(mod_raw_card).strip() if mod_raw_card and str(mod_raw_card).lower() != 'nan' else "N/D",
                'instrument_range_on_card': str(ran_raw_card).strip() if ran_raw_card and str(ran_raw_card).lower() != 'nan' else "N/D",
                'is_expired_at_use': is_exp,
                'tipologia_strumento_scheda': tipologia_strumento_scheda, # Da SP
                'modello_L9_scheda': modello_l9_scheda_normalizzato if file_type == 'analogico' else "N/A", # Da L9 (solo ANA)
                'modello_strumento_campione_usato': mod_camp_reg, # Dal registro
                'is_congruent': is_congr, 'congruency_notes': congr_notes,
                'used_before_emission': used_before_em
            })
        except IndexError: # Errore se le celle dei certificati non esistono (es. file con meno righe/colonne)
            log_to_all(log_widget_gui,f"    {base_filename}: Cert. slot {i+1} - Errore indice cella. Slot potrebbe essere vuoto o file malformato.","WARNING")
        except Exception as e_slot: # Altri errori inattesi per lo slot
            log_to_all(log_widget_gui,f"    {base_filename}: Cert. slot {i+1} - Errore: {e_slot}","ERROR",exc_info=True)
    
    # Ritorna i dati estratti e un set unico delle chiavi di errore "umano"
    return card_date, f"{file_type} - {len(extracted_certs_data)} cert.", extracted_certs_data, list(set(human_error_keys_this_file)), dati_scheda_per_compilazione
# --- FINE PARTE 2 ---
# --- PARTE 3 ---
def trova_strumenti_alternativi(range_richiesto_raw, data_riferimento_scheda, strumenti_campione_list):
    logger.debug(f"Inizio trova_strumenti_alternativi. Range Richiesto (raw): '{range_richiesto_raw}', Data Riferimento: {data_riferimento_scheda}")
    if not strumenti_campione_list:
        logger.warning("Lista strumenti campione vuota o non disponibile in trova_strumenti_alternativi.")
        return []

    alternative_valide = []
    range_richiesto_norm = normalize_range_string(range_richiesto_raw)
    logger.debug(f"Range richiesto normalizzato per ricerca: '{range_richiesto_norm}'")

    # Assicura che la data di riferimento sia 'naive' per il confronto con le date del registro
    # che potrebbero essere state lette come naive o convertite a naive.
    data_riferimento_scheda_naive = None
    if data_riferimento_scheda is None: # Se non fornita, usa la data corrente dell'analisi
        data_riferimento_scheda_naive = ANALYSIS_DATETIME.replace(tzinfo=None) # Usa UTC naive se non specificato
        logger.debug(f"Data riferimento impostata a default (ANALYSIS_DATETIME naive UTC): {data_riferimento_scheda_naive}")
    elif data_riferimento_scheda.tzinfo is not None:
        data_riferimento_scheda_naive = data_riferimento_scheda.astimezone(timezone.utc).replace(tzinfo=None)
    else: # Già naive
        data_riferimento_scheda_naive = data_riferimento_scheda

    for strumento in strumenti_campione_list:
        scadenza_campione_naive = None
        if strumento.get('scadenza'): # 'scadenza' è un oggetto datetime o None
            scadenza_dt_originale = strumento['scadenza']
            # Rendi naive anche la scadenza del campione per confronto omogeneo
            scadenza_campione_naive = scadenza_dt_originale.astimezone(timezone.utc).replace(tzinfo=None) if scadenza_dt_originale.tzinfo is not None else scadenza_dt_originale
        
        data_emissione_campione_naive = None
        if strumento.get('data_emissione'): # 'data_emissione' è un oggetto datetime o None
            emissione_dt_originale = strumento['data_emissione']
            data_emissione_campione_naive = emissione_dt_originale.astimezone(timezone.utc).replace(tzinfo=None) if emissione_dt_originale.tzinfo is not None else emissione_dt_originale
        
        # Verifica validità temporale dello strumento campione
        if data_emissione_campione_naive and scadenza_campione_naive and \
           data_emissione_campione_naive <= data_riferimento_scheda_naive < scadenza_campione_naive:
            # Strumento campione valido temporalmente, controlla il range
            range_campione_norm = normalize_range_string(strumento.get('range', ""))
            if range_richiesto_norm == range_campione_norm:
                alternative_valide.append(strumento)

    # Ordina le alternative per data di scadenza più lontana
    alternative_valide.sort(key=lambda x: (x.get('scadenza') or datetime.min.replace(tzinfo=timezone.utc)), reverse=True)
    logger.debug(f"Trovate {len(alternative_valide)} alternative valide con range '{range_richiesto_norm}'.")
    return alternative_valide

# --- FUNZIONE PER REPORT WORD ---
def crea_e_apri_report_anomalie_word(
    log_widget, parent_window_for_msg,
    errors_list, # Lista di {'file': str, 'key': str, 'path': str} per errori di compilazione Range/UM etc.
    temporal_list, # Lista di dict da all_extracted_data per usi prematuri o scaduti
    incongruent_list, # Lista di dict da all_extracted_data per incongruenze di regole
    candidate_files_count, validated_file_count):

    logger.info("Inizio creazione report anomalie Word.")
    log_to_all(log_widget, "Creazione report anomalie Word in corso...", "INFO")

    # Controlla se ci sono effettivamente anomalie da riportare
    anomalie_da_riportare = bool(errors_list or \
                                 any(item['alert_type'] == 'premature_emission' for item in temporal_list) or \
                                 any(item['alert_type'] == 'expired_at_use' for item in temporal_list) or \
                                 incongruent_list)

    if not anomalie_da_riportare:
        logger.info("Nessuna anomalia significativa da riportare nel file Word.")
        log_to_all(log_widget, "Nessuna anomalia significativa trovata per il report Word.", "INFO")
        if parent_window_for_msg and parent_window_for_msg.winfo_exists():
            messagebox.showinfo("Nessuna Anomalia", "Nessuna anomalia significativa da includere nel report Word.", parent=parent_window_for_msg)
        return

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    title = doc.add_heading("Report Anomalie Analisi Schede Taratura", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data analisi e cartella
    doc.add_paragraph().add_run(f"Data Analisi: {ANALYSIS_DATETIME.astimezone().strftime('%d/%m/%Y %H:%M:%S %Z')}").italic = True
    p_folder = doc.add_paragraph()
    p_folder.add_run(f"Cartella Analizzata: ").italic = True
    p_folder.add_run(f"{FOLDER_PATH_DEFAULT if FOLDER_PATH_DEFAULT else 'N/D - Errore Config.'}").bold = True
    p_folder.add_run(f" (File Candidati: {candidate_files_count}, Schede Validate: {validated_file_count})").italic = True
    doc.add_paragraph() # Spazio

    # Sezione 1: Errori di Compilazione Scheda (Range/UM, SP, L9, etc.)
    # Filtra per non includere gli errori COMP_ qui, verranno dopo se necessario (o in un'altra sezione)
    errori_comp_scheda_strutturali = [err for err in errors_list if not err['key'].startswith("COMP_")]
    if errori_comp_scheda_strutturali:
        doc.add_heading("1. Errori di Compilazione Scheda (Strutturali: Range/UM, SP, L9)", level=1)
        errori_comp_scheda_strutturali.sort(key=lambda x: (x['file'], x['key'])) # Ordina per file, poi per tipo errore
        current_file_comp_err = None
        for error in errori_comp_scheda_strutturali:
            if error['file'] != current_file_comp_err:
                if current_file_comp_err is not None:
                    doc.add_paragraph() # Spazio tra file diversi
                p_file = doc.add_paragraph()
                p_file.add_run("File: ").bold = True
                p_file.add_run(f"{error['file']} (Percorso: {error['path']})")
                current_file_comp_err = error['file']
            
            error_desc = human_error_messages_map_descriptive.get(error['key'], f"Codice Errore Sconosciuto: {error['key']}")
            doc.add_paragraph(f"  • {error_desc}", style='ListBullet')
        doc.add_paragraph() # Spazio dopo la sezione

    # Sezione 2: Certificati Utilizzati Prima dell'Emissione
    premature_uses = [item for item in temporal_list if item['alert_type'] == 'premature_emission']
    if premature_uses:
        doc.add_heading("2. Certificati Campione Utilizzati Prima della Loro Data di Emissione", level=1)
        premature_uses.sort(key=lambda x: (x.get('file_name', ''), x.get('card_date', datetime.min.replace(tzinfo=timezone.utc))))
        for item in premature_uses:
            p = doc.add_paragraph()
            p.add_run("File Scheda: ").bold = True
            p.add_run(f"{item['file_name']} ({item['file_path']})")
            doc.add_paragraph(f"  • Data Scheda: {item['card_date_str']}", style='ListBullet')
            doc.add_paragraph(f"  • Certificato ID: {item['certificate_id']}", style='ListBullet')
            p_cert_detail = doc.add_paragraph(f"  • CERTIFICATO CAMPIONE (da Registro): {item['modello_strumento_campione_usato']}", style='ListBullet')
            p_cert_detail_em = doc.add_paragraph(f"  • Data Emissione Certificato Campione: {item.get('data_emissione_presunta','N/A')}", style='ListBullet')
            run_em_note = p_cert_detail_em.add_run(" - USATO PRIMA DELL'EMISSIONE!")
            run_em_note.bold = True; run_em_note.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Rosso
            # Nota sulla congruità (esclusa la parte di emissione già trattata)
            congr_note_cleaned = item['congruency_notes'].replace(f"Dettaglio: errato per EMISSIONE. Cert.'{item['certificate_id']}'({item['modello_strumento_campione_usato']}) usato il {item['card_date']:%d/%m/%Y} ma il certificato è stato emesso il {item.get('data_emissione_presunta','N/A')}.", "").strip()
            if congr_note_cleaned and congr_note_cleaned != ".": # Evita note vuote o solo un punto
                 doc.add_paragraph(f"  • Altre Note Congruità: {congr_note_cleaned}", style='ListBullet')
            doc.add_paragraph() # Spazio tra item
        doc.add_paragraph() # Spazio dopo la sezione

    # Sezione 3: Certificati Scaduti Utilizzati (non per emissione prematura)
    expired_at_use_pure = [item for item in temporal_list if item['alert_type'] == 'expired_at_use'] # Già filtrati in run_analysis
    if expired_at_use_pure:
        doc.add_heading("3. Certificati Campione Scaduti al Momento dell'Uso", level=1)
        expired_at_use_pure.sort(key=lambda x: (x.get('file_name', ''), x.get('card_date', datetime.min.replace(tzinfo=timezone.utc))))
        for item in expired_at_use_pure:
            p = doc.add_paragraph()
            p.add_run("File Scheda: ").bold = True
            p.add_run(f"{item['file_name']} ({item['file_path']})")
            doc.add_paragraph(f"  • Data Scheda: {item['card_date_str']}", style='ListBullet')
            doc.add_paragraph(f"  • Certificato ID: {item['certificate_id']}", style='ListBullet')
            p_cert_detail_exp = doc.add_paragraph(f"  • CERTIFICATO CAMPIONE (da Registro): {item['modello_strumento_campione_usato']}", style='ListBullet')
            p_cert_detail_scad = doc.add_paragraph(f"  • Data Scadenza Certificato Campione: {item.get('expiry_date_str','N/P')}", style='ListBullet')
            run_scad_note = p_cert_detail_scad.add_run(" - SCADUTO ALL'USO!")
            run_scad_note.bold = True; run_scad_note.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00) # Arancione scuro
            doc.add_paragraph(f"  • Note Congruità (se presenti): {item['congruency_notes']}", style='ListBullet')
            doc.add_paragraph() # Spazio
        doc.add_paragraph()

    # Sezione 4: Certificati Non Congrui per Regole Applicate (non per emissione o scadenza già coperte)
    if incongruent_list: # Questa lista è già filtrata per non includere 'used_before_emission'
        doc.add_heading("4. Certificati Campione Non Congrui (per Regole Tipologia/Modello)", level=1)
        incongruent_list.sort(key=lambda x: (x.get('file_name', ''), x.get('card_date', datetime.min.replace(tzinfo=timezone.utc))))
        for item in incongruent_list:
            p = doc.add_paragraph()
            p.add_run("File Scheda: ").bold = True
            p.add_run(f"{item['file_name']} ({item['file_path']})")
            doc.add_paragraph(f"  • Data Scheda: {item['card_date_str']}", style='ListBullet')
            doc.add_paragraph(f"    Tipologia Strumento Scheda (da SP): {item.get('tipologia_strumento_scheda','N/D')}", style='ListBullet').paragraph_format.left_indent = Inches(0.25)
            if item.get('card_type') == 'analogico':
                doc.add_paragraph(f"    Modello L9 Scheda (Analogico): {item.get('modello_L9_scheda','N/A')}", style='ListBullet').paragraph_format.left_indent = Inches(0.25)

            doc.add_paragraph(f"  • Certificato ID: {item['certificate_id']}", style='ListBullet')
            p_cert_detail_inc = doc.add_paragraph(f"  • CERTIFICATO CAMPIONE (da Registro): {item['modello_strumento_campione_usato']}", style='ListBullet')
            
            # Info scadenza (se scaduto anche, ma non è la causa primaria qui)
            scad_info_text = " (SCADUTO ALL'USO!)" if item.get('is_expired_at_use') else "" # is_expired_at_use è ancora nel dict
            p_cert_detail_scad_inc = doc.add_paragraph(style='ListBullet')
            p_cert_detail_scad_inc.paragraph_format.left_indent = Inches(0.25) # Indenta info scadenza
            p_cert_detail_scad_inc.add_run(f"    Scadenza Certificato Campione: {item.get('certificate_expiry_str','N/P')}")
            if scad_info_text:
                run_extra_scad = p_cert_detail_scad_inc.add_run(scad_info_text)
                run_extra_scad.bold = True
                run_extra_scad.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00) # Arancione scuro
            
            # Motivo non congruità
            congr_note_text = item.get('congruency_notes','')
            if congr_note_text.startswith("Dettaglio:"): # Se la nota è già dettagliata
                p_reason = doc.add_paragraph(f"  • {congr_note_text}", style='ListBullet')
            else:
                p_reason = doc.add_paragraph(f"  • Motivo Non Congruità: ", style='ListBullet')
                run_reason = p_reason.add_run(congr_note_text if congr_note_text else "Non specificato")
                run_reason.bold = True; run_reason.font.color.rgb = RGBColor(0x80, 0x00, 0x80) # Viola
            
            doc.add_paragraph() # Spazio
        doc.add_paragraph()

    # Sezione 5: Errori di Compilazione Campi ODC, Data, PDL (quelli con prefisso COMP_)
    errori_comp_campi_scheda = [err for err in errors_list if err['key'].startswith("COMP_")]
    if errori_comp_campi_scheda:
        doc.add_heading("5. Errori/Mancanze nei Campi Anagrafici Scheda (ODC, Data, PDL, etc.)", level=1)
        errori_comp_campi_scheda.sort(key=lambda x: (x['file'], x['key']))
        current_file_comp_campi_err = None
        for error in errori_comp_campi_scheda:
            if error['file'] != current_file_comp_campi_err:
                if current_file_comp_campi_err is not None:
                    doc.add_paragraph() 
                p_file_campi = doc.add_paragraph()
                p_file_campi.add_run("File: ").bold = True
                p_file_campi.add_run(f"{error['file']} (Percorso: {error['path']})")
                current_file_comp_campi_err = error['file']
            
            error_desc_campi = human_error_messages_map_descriptive.get(error['key'], f"Codice Errore Compilazione Sconosciuto: {error['key']}")
            doc.add_paragraph(f"  • {error_desc_campi}", style='ListBullet')
        doc.add_paragraph()

    # Salvataggio e apertura file Word
    try:
        temp_dir = tempfile.gettempdir() # Cartella temporanea di sistema
        timestamp_rep = ANALYSIS_DATETIME.astimezone().strftime("%Y%m%d_%H%M%S")
        word_filename = f"Report_Anomalie_Schede_{timestamp_rep}.docx"
        word_file_path = os.path.join(temp_dir, word_filename)
        doc.save(word_file_path)
        logger.info(f"Report Word salvato in: {word_file_path}")

        # Tenta di aprire il file
        if sys.platform == "win32":
            os.startfile(word_file_path)
        elif sys.platform == "darwin": # macOS
            subprocess.Popen(["open", word_file_path])
        else: # linux variants
            subprocess.Popen(["xdg-open", word_file_path])
        logger.info(f"Tentativo di apertura report Word: {word_file_path}")
        log_to_all(log_widget, f"Report anomalie Word generato e aperto: {word_file_path}", "SUCCESS")
        if parent_window_for_msg and parent_window_for_msg.winfo_exists(): 
            messagebox.showinfo("Report Generato", f"Report anomalie generato e aperto:\n{word_file_path}", parent=parent_window_for_msg)

    except Exception as e_word:
        logger.error(f"Errore durante la creazione o apertura del report Word: {e_word}", exc_info=True)
        log_to_all(log_widget, f"Errore creazione/apertura report Word: {e_word}", "ERROR")
        if parent_window_for_msg and parent_window_for_msg.winfo_exists():
            messagebox.showerror("Errore Report Word", f"Impossibile generare o aprire il report Word.\nErrore: {e_word}\nControllare il log.", parent=parent_window_for_msg)

# --- FINE PARTE 3 ---
# --- PARTE 4 ---

def run_analysis_and_show_results():
    # Dizionario per tenere traccia dei widget dell'interfaccia del tab suggerimenti,
    # per poterli popolare da altre parti del codice (es. doppio click su un certificato)
    sugg_tab_interface_details = {
        'notebook_widget': None, 'suggerimenti_tab_widget': None,
        'cert_id_entry_widget': None, 'range_entry_widget': None,
        'date_entry_widget': None, 'search_function_widget': None,
        'results_text_widget': None
    }
    # Riferimento alla finestra dei risultati per i messagebox
    result_window_ref_store = {'ref': None}

    # Lista per memorizzare i dati specifici di ogni scheda analizzata, utili per la compilazione
    # Questa variabile è resa globale per essere accessibile da esegui_compilazione_schede definita più avanti
    global schede_analizzate_info_list # Dichiarazione per modifica
    schede_analizzate_info_list = []


    def on_file_click(file_path_to_copy, filename_to_open, open_file_direct=False, event=None):
        nonlocal result_window_ref_store, root # Accedi alle variabili non locali/globali
        parent_win_for_msg = result_window_ref_store.get('ref')
        if not (parent_win_for_msg and parent_win_for_msg.winfo_exists()): # Fallback se la finestra risultati non esiste
            parent_win_for_msg = root

        logger.info(f"on_file_click: Path='{file_path_to_copy}', File='{filename_to_open}', OpenDirect={open_file_direct}")

        if not file_path_to_copy or not isinstance(file_path_to_copy, str) or not file_path_to_copy.strip():
            logger.error(f"Percorso file non valido o vuoto: '{file_path_to_copy}'")
            if parent_win_for_msg and parent_win_for_msg.winfo_exists():
                messagebox.showerror("Errore Percorso", "Percorso file non fornito o non valido.", parent=parent_win_for_msg)
            return

        try:
            normalized_path_to_copy = os.path.normpath(file_path_to_copy)
            pyperclip.copy(normalized_path_to_copy)
            logger.info(f"Percorso '{normalized_path_to_copy}' copiato negli appunti (pyperclip).")
            
            action_text = "Aprire il file" if open_file_direct else f"Aprire la cartella del file '{filename_to_open}'"
            
            # Chiedi conferma all'utente se la finestra dei risultati è disponibile
            action = True # Default se non c'è finestra per chiedere
            if parent_win_for_msg and parent_win_for_msg.winfo_exists():
                action = messagebox.askyesnocancel("Percorso Copiato",
                                                f"Percorso copiato negli appunti:\n{normalized_path_to_copy}\n\n{action_text}?",
                                                parent=parent_win_for_msg, icon='question')
            else:
                logger.warning("Finestra parent per messagebox.askyesnocancel non disponibile. Procedo con l'azione di default (apertura).")

            if action: # True (Sì) o None (se messagebox.askyesnocancel non eseguito per mancanza finestra)
                target_to_open = normalized_path_to_copy if open_file_direct else os.path.dirname(normalized_path_to_copy)
                logger.info(f"Tentativo di aprire: '{target_to_open}' (open_direct: {open_file_direct})")
                try:
                    if not os.path.exists(target_to_open):
                        logger.error(f"Target '{target_to_open}' non esiste.")
                        if parent_win_for_msg and parent_win_for_msg.winfo_exists():
                            messagebox.showerror("Errore Apertura", f"Il percorso specificato non esiste:\n{target_to_open}", parent=parent_win_for_msg)
                        return

                    if sys.platform == "win32":
                        logger.debug(f"Windows: Chiamata a os.startfile('{target_to_open}')")
                        os.startfile(target_to_open)
                    elif sys.platform == "darwin":
                        logger.debug(f"macOS: Chiamata a subprocess.Popen(['open', '{target_to_open}'])")
                        subprocess.Popen(["open", target_to_open])
                    else: # linux variants
                        logger.debug(f"Linux/Unix: Chiamata a subprocess.Popen(['xdg-open', '{target_to_open}'])")
                        subprocess.Popen(["xdg-open", target_to_open])
                    logger.info(f"Comando di apertura per '{target_to_open}' eseguito.")
                except Exception as e_open:
                    logger.error(f"Errore durante l'apertura di '{target_to_open}': {e_open}", exc_info=True)
                    if parent_win_for_msg and parent_win_for_msg.winfo_exists():
                        messagebox.showerror("Errore Apertura", f"Impossibile aprire:\n{target_to_open}\n\nErrore: {e_open}", parent=parent_win_for_msg)
        except pyperclip.PyperclipException as e_clip: # Pyperclip non disponibile o errore
            logger.warning(f"Pyperclip non disponibile o errore durante la copia di '{normalized_path_to_copy}': {e_clip}. Tentativo con Tkinter fallback.")
            try:
                if parent_win_for_msg and parent_win_for_msg.winfo_exists():
                    parent_win_for_msg.clipboard_clear()
                    parent_win_for_msg.clipboard_append(normalized_path_to_copy)
                    parent_win_for_msg.update() # Necessario per Tkinter clipboard
                    logger.info(f"Percorso '{normalized_path_to_copy}' copiato negli appunti (Tkinter fallback).")
                    messagebox.showinfo("Percorso Copiato (Fallback)", f"Pyperclip non disponibile.\nPercorso copiato (fallback Tkinter):\n{normalized_path_to_copy}", parent=parent_win_for_msg)
                else:
                    logger.warning("Finestra parent per clipboard Tkinter fallback non disponibile.")
            except tk.TclError as e_tk_clip:
                logger.error(f"Errore durante la copia con Tkinter fallback per '{normalized_path_to_copy}': {e_tk_clip}")
                # Se anche Tkinter fallisce, mostra solo il percorso
                if parent_win_for_msg and parent_win_for_msg.winfo_exists():
                    messagebox.showinfo("Percorso", f"Impossibile copiare automaticamente.\nPercorso:\n{normalized_path_to_copy}", parent=parent_win_for_msg)
        except Exception as e_gen: # Errore generico non previsto
            logger.error(f"Errore generico in on_file_click per '{file_path_to_copy}': {e_gen}", exc_info=True)
            if parent_win_for_msg and parent_win_for_msg.winfo_exists():
                messagebox.showerror("Errore Imprevisto",f"Si è verificato un errore imprevisto durante la gestione del file:\n{file_path_to_copy}\n\nErrore: {e_gen}", parent=parent_win_for_msg)

    def launch_suggestion_tab_and_search(entry_data_from_alert):
        nonlocal result_window_ref_store # Per messagebox in caso di errore
        # Verifica che tutti i widget necessari del tab suggerimenti siano stati inizializzati
        required_keys = ['notebook_widget', 'suggerimenti_tab_widget', 'cert_id_entry_widget',
                         'range_entry_widget', 'date_entry_widget', 'search_function_widget']
        if not all(sugg_tab_interface_details.get(key) for key in required_keys):
            logger.error("Interfaccia del tab suggerimenti non completamente inizializzata.")
            parent_win = result_window_ref_store.get('ref')
            if parent_win and parent_win.winfo_exists(): #Mostra errore se la finestra dei risultati è disponibile
                messagebox.showerror("Errore Interno", "La funzionalità di suggerimento non è pronta.", parent=parent_win)
            return
        
        # Popola i campi di input del tab suggerimenti
        sugg_tab_interface_details['cert_id_entry_widget'].delete(0, tk.END)
        sugg_tab_interface_details['cert_id_entry_widget'].insert(0, entry_data_from_alert.get('cert_id', ''))
        sugg_tab_interface_details['range_entry_widget'].delete(0, tk.END)
        sugg_tab_interface_details['range_entry_widget'].insert(0, entry_data_from_alert.get('range', '')) # 'range' è già nel dict entry_data
        sugg_tab_interface_details['date_entry_widget'].delete(0, tk.END)
        sugg_tab_interface_details['date_entry_widget'].insert(0, entry_data_from_alert.get('card_date_str', ANALYSIS_DATETIME.astimezone().strftime('%d/%m/%Y')))
        
        # Cambia al tab suggerimenti ed esegui la ricerca
        sugg_tab_interface_details['notebook_widget'].select(sugg_tab_interface_details['suggerimenti_tab_widget'])
        sugg_tab_interface_details['search_function_widget']() # Chiama la funzione di ricerca


    # --- Inizio Effettivo Analisi ---
    root = tk.Tk(); root.withdraw() # Nascondi la finestra root principale

    # Finestra per i log di avanzamento
    progress_window = tk.Toplevel(root); progress_window.title("Progresso Analisi Schede"); progress_window.geometry("950x650")
    log_text_frame = tk.Frame(progress_window); log_text_frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
    log_v_scroll = tk.Scrollbar(log_text_frame); log_v_scroll.pack(side=tk.RIGHT, fill=tk.Y)
    log_h_scroll = tk.Scrollbar(log_text_frame, orient=tk.HORIZONTAL); log_h_scroll.pack(side=tk.BOTTOM, fill=tk.X)
    log_text_widget = tk.Text(log_text_frame, wrap=tk.NONE, relief=tk.SUNKEN, borderwidth=1,
                              yscrollcommand=log_v_scroll.set, xscrollcommand=log_h_scroll.set, font=("Consolas", 10))
    log_text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    log_v_scroll.config(command=log_text_widget.yview); log_h_scroll.config(command=log_text_widget.xview)
    log_text_widget.config(state=tk.DISABLED) # Inizia disabilitato, abilitato solo per inserire testo
    
    log_to_all(log_text_widget, f"--- INIZIO ANALISI ({ANALYSIS_DATETIME.astimezone().strftime('%d/%m/%Y %H:%M:%S %Z')}) ---", "INFO")
    log_to_all(log_text_widget, f"Cartella Script: {SCRIPT_DIR}", "DEBUG")
    if FILE_REGISTRO_STRUMENTI:
        log_to_all(log_text_widget, f"File Registro Strumenti (da parametri.xlsm): {FILE_REGISTRO_STRUMENTI}", "DEBUG")
    else:
        log_to_all(log_text_widget, "ERRORE: FILE_REGISTRO_STRUMENTI non impostato!", "ERROR") 

    if FOLDER_PATH_DEFAULT:
        log_to_all(log_text_widget, f"Cartella Schede (da parametri.xlsm): {FOLDER_PATH_DEFAULT}", "DEBUG")
    else:
        log_to_all(log_text_widget, "ERRORE: FOLDER_PATH_DEFAULT non impostato!", "ERROR") 
    
    if FILE_DATI_COMPILAZIONE_SCHEDE:
        log_to_all(log_text_widget, f"File Dati Compilazione Schede (da parametri.xlsm B4): {FILE_DATI_COMPILAZIONE_SCHEDE}", "DEBUG")
    else:
        log_to_all(log_text_widget, f"File Dati Compilazione Schede (da parametri.xlsm B4) NON IMPOSTATO. La compilazione automatica non sarà disponibile.", "WARNING")


    # Controllo critico percorsi, esci se mancano
    if not FILE_REGISTRO_STRUMENTI or not FOLDER_PATH_DEFAULT:
        error_msg_critical = "ERRORE CRITICO: Percorsi fondamentali non configurati. L'applicazione non può continuare."
        log_to_all(log_text_widget, error_msg_critical, "ERROR")
        messagebox.showerror("Errore Configurazione Percorsi", error_msg_critical + "\nControllare 'parametri.xlsm' e i log.", parent=progress_window if progress_window.winfo_exists() else root)
        if progress_window.winfo_exists(): progress_window.destroy()
        if root.winfo_exists(): root.destroy()
        return # Termina l'esecuzione della funzione

    # Carica registro strumenti
    strumenti_campione_list = leggi_registro_strumenti(log_text_widget)
    if strumenti_campione_list is None: 
        log_to_all(log_text_widget, "ERRORE CRITICO: Registro strumenti campione non caricato. Funzionalità limitate.", "ERROR")
        messagebox.showerror("Errore Registro Strumenti", "Impossibile caricare il registro strumenti campione.\nControllare il log.", parent=progress_window)
        strumenti_campione_list = [] 
    elif not strumenti_campione_list: 
        log_to_all(log_text_widget, "Registro strumenti campione caricato ma VUOTO. Funzionalità limitate.", "WARNING")

    # Verifica cartella schede
    if not os.path.isdir(FOLDER_PATH_DEFAULT): 
        log_to_all(log_text_widget, f"PERCORSO CARTELLA SCHEDE NON VALIDO O INACCESSIBILE (da parametri.xlsm): {FOLDER_PATH_DEFAULT}", "ERROR")
        messagebox.showerror("Errore Percorso Schede", f"Percorso cartella schede non valido o non accessibile:\n{FOLDER_PATH_DEFAULT}\nControllare 'parametri.xlsm'.", parent=progress_window)
        if progress_window.winfo_exists(): progress_window.destroy()
        if root.winfo_exists(): root.destroy(); return
        
    # Inizializzazione liste per risultati
    all_extracted_data = []; validated_file_count = 0; human_errors_details_list = []
    schede_analizzate_info_list.clear() 

    # Trova file candidati
    candidate_files = [] 
    try:
        candidate_files = [f for f in os.listdir(FOLDER_PATH_DEFAULT) if f.lower().endswith(('.xls', '.xlsx')) and not f.startswith('~')] 
        log_to_all(log_text_widget, f"Trovati {len(candidate_files)} file candidati in: {FOLDER_PATH_DEFAULT}")
    except Exception as e:
        log_to_all(log_text_widget, f"ERRORE CRITICO lettura cartella schede {FOLDER_PATH_DEFAULT}: {e}", "ERROR", exc_info=True)
        messagebox.showerror("Errore Cartella Schede", f"Errore lettura cartella schede:\n{FOLDER_PATH_DEFAULT}\n{e}", parent=progress_window)
        if progress_window.winfo_exists(): progress_window.destroy()
        if root.winfo_exists(): root.destroy(); return
    if not candidate_files:
        log_to_all(log_text_widget, "Nessun file candidato trovato nella cartella schede.", "WARNING")
        messagebox.showinfo("Nessun File", f"Nessun file .xls o .xlsx trovato in:\n{FOLDER_PATH_DEFAULT}", parent=progress_window)
        
    # Ciclo principale di analisi dei file
    for filename_loop in candidate_files:
        file_path = os.path.join(FOLDER_PATH_DEFAULT, filename_loop)
        logger.debug(f"Preparazione analisi per file: '{file_path}' (da filename_loop: '{filename_loop}')")
        
        # Analizza il singolo file
        card_date_from_file, status_msg, certs_from_file, human_error_keys_in_file, dati_comp_scheda_corrente = analyze_excel_file(file_path, log_text_widget, strumenti_campione_list)
        
        if card_date_from_file: 
            validated_file_count += 1
            all_extracted_data.extend(certs_from_file or []) 
            if dati_comp_scheda_corrente: 
                schede_analizzate_info_list.append(dati_comp_scheda_corrente)
        
        if human_error_keys_in_file:
            for error_key in list(set(human_error_keys_in_file)): 
                human_errors_details_list.append({'file': filename_loop, 'key': error_key, 'path': file_path})
        log_to_all(log_text_widget, f"--- Fine analisi file: {filename_loop}. Status: {status_msg} ---", "INFO")
        
    # Fine ciclo analisi file
    log_to_all(log_text_widget, f"--- ANALISI FILE COMPLETATA ---", "INFO")
    log_to_all(log_text_widget, f"{validated_file_count} su {len(candidate_files)} schede validate.", "SUCCESS" if validated_file_count > 0 else "WARNING")
    log_to_all(log_text_widget, f"Estratti {len(all_extracted_data)} utilizzi certificati.", "INFO")
    log_to_all(log_text_widget, f"Rilevate {len(human_errors_details_list)} istanze errori compilazione.", "INFO" if not human_errors_details_list else "WARNING")
    log_to_all(log_text_widget, f"Raccolte informazioni per la compilazione da {len(schede_analizzate_info_list)} schede.", "DEBUG")
    log_to_all(log_text_widget, "Aggregazione dati per report e GUI...", "DEBUG")
    
    # Preparazione dati aggregati per GUI e report Word
    report_sections_generale_sintetico = []
    report_sections_generale_sintetico.append(f"Data Analisi: {ANALYSIS_DATETIME.astimezone().strftime('%d/%m/%Y %H:%M:%S %Z')}")
    report_sections_generale_sintetico.append(f"Cartella Schede Analizzata: {FOLDER_PATH_DEFAULT if FOLDER_PATH_DEFAULT else 'N/D'} (File Candidati: {len(candidate_files)}, Schede Validate: {validated_file_count})")
    total_cert_usages = len(all_extracted_data); congru_u = sum(1 for i in all_extracted_data if i.get('is_congruent') is True)
    incongru_total_u = sum(1 for i in all_extracted_data if i.get('is_congruent') is False); not_verif_u = sum(1 for i in all_extracted_data if i.get('is_congruent') is None)
    used_before_em_u = sum(1 for i in all_extracted_data if i.get('used_before_emission') is True)
    expired_pure_u = sum(1 for i in all_extracted_data if i.get('is_expired_at_use') and not i.get('used_before_emission')) 
    incongru_rules_u = incongru_total_u - used_before_em_u 
    cert_stats_summary = (
        f"Utilizzi Certificati Tot: {total_cert_usages} "
        f"(Congrui: {congru_u}, Non Congrui Tot: {incongru_total_u} "
        f"[Prima Emiss: {used_before_em_u}, Solo Regole: {incongru_rules_u}], "
        f"Scaduti (non per emiss.): {expired_pure_u}, Non Verif: {not_verif_u})"
    )
    report_sections_generale_sintetico.append(cert_stats_summary)
    total_compilation_errors = len(human_errors_details_list) 
    report_sections_generale_sintetico.append(f"Errori Compilazione Scheda (strutturali e anagrafici): {total_compilation_errors} (Dettagliati nei tab e nel report Word)")
    if LOG_FILEPATH : report_sections_generale_sintetico.append(f"Log Completo Salvato in: {LOG_FILEPATH}")
    
    temporal_alerts_entries = [] 
    for item in all_extracted_data:
        item_copy = item.copy() 
        item_copy['card_date_str'] = item_copy.get('card_date', datetime.min.replace(tzinfo=timezone.utc)).strftime('%d/%m/%Y')
        item_copy['range'] = item_copy.get('instrument_range_on_card', 'N/D') 
        
        if item_copy.get('used_before_emission'):
            found_c = next((s for s in strumenti_campione_list if s['id_certificato'] == item_copy['certificate_id']), None) if strumenti_campione_list else None
            item_copy['data_emissione_presunta'] = found_c['data_emissione'].strftime('%d/%m/%Y') if found_c and found_c.get('data_emissione') else "N/A"
            item_copy['alert_type'] = 'premature_emission'; temporal_alerts_entries.append(item_copy)
        elif item_copy.get('is_expired_at_use'): 
            item_copy['expiry_date_str'] = item_copy['certificate_expiry'].strftime('%d/%m/%Y') if item_copy.get('certificate_expiry') else "N/P"
            item_copy['alert_type'] = 'expired_at_use'; temporal_alerts_entries.append(item_copy)
    temporal_alerts_entries.sort(key=lambda x: (x.get('card_date', datetime.min.replace(tzinfo=timezone.utc)), x.get('file_name', '')))
    
    incongruent_rules_list = [{**i, 'range': i.get('instrument_range_on_card', 'N/D'), 
                              'card_date_str': i.get('card_date', datetime.min.replace(tzinfo=timezone.utc)).strftime('%d/%m/%Y'),
                              'certificate_expiry_str': i['certificate_expiry'].strftime('%d/%m/%Y') if i.get('certificate_expiry') else "N/P"}
                             for i in all_extracted_data if i.get('is_congruent') is False and not i.get('used_before_emission')]
    incongruent_rules_list.sort(key=lambda x: (x.get('card_date', datetime.min.replace(tzinfo=timezone.utc)), x.get('file_name', '')))
    
    cert_details_map = defaultdict(lambda: {
        'id': "", 'utilizzi': 0, 'date_scadenza_raw_set': set(), 'date_utilizzo_obj_set': set(),
        'modelli_su_scheda_counter': Counter(), 'range_su_scheda_counter': Counter(),
        'tipologie_scheda_associate_counter': Counter(),
        'usi_congrui': 0, 'usi_total_incongrui': 0, 'usi_congruita_non_verificata': 0,
        'usi_prima_emissione': 0, 'usi_scaduti_puri': 0, 'max_giorni_scaduto_uso': 0,
        'dettaglio_usi_list': []
    })
    for item in all_extracted_data:
        details = cert_details_map[item['certificate_id']]; details['id'] = item['certificate_id']; details['utilizzi'] += 1
        if item.get('certificate_expiry_raw') and str(item['certificate_expiry_raw']).lower() != 'nan': details['date_scadenza_raw_set'].add(item['certificate_expiry_raw'])
        if item.get('card_date'): details['date_utilizzo_obj_set'].add(item['card_date'])
        model_card = item.get('instrument_model_on_card', "N/D"); range_card = item.get('instrument_range_on_card', "N/D"); tip_sch = item.get('tipologia_strumento_scheda', "N/D")
        if model_card != "N/D": details['modelli_su_scheda_counter'][model_card] += 1
        if range_card != "N/D": details['range_su_scheda_counter'][range_card] += 1
        if tip_sch != "N/D": details['tipologie_scheda_associate_counter'][tip_sch] += 1
        
        if item.get('is_congruent') is True: details['usi_congrui'] += 1
        elif item.get('is_congruent') is False: details['usi_total_incongrui'] += 1
        else: details['usi_congruita_non_verificata'] += 1
        
        if item.get('used_before_emission'): details['usi_prima_emissione'] += 1
        elif item.get('is_expired_at_use') and item.get('certificate_expiry') and item.get('card_date'): 
            details['usi_scaduti_puri'] += 1
            card_date_comp = item['card_date'].replace(tzinfo=None) if item['card_date'].tzinfo else item['card_date']
            cert_exp_comp = item['certificate_expiry'].replace(tzinfo=None) if item['certificate_expiry'].tzinfo else item['certificate_expiry']
            giorni_scad = (card_date_comp - cert_exp_comp).days
            details['max_giorni_scaduto_uso'] = max(details['max_giorni_scaduto_uso'], giorni_scad)
            
        uso_det = {**item, 'range_su_scheda': range_card, 'modello_su_scheda': model_card}
        if item.get('is_expired_at_use') and item.get('certificate_expiry') and item.get('card_date'):
            card_date_comp = item['card_date'].replace(tzinfo=None) if item['card_date'].tzinfo else item['card_date']
            cert_exp_comp = item['certificate_expiry'].replace(tzinfo=None) if item['certificate_expiry'].tzinfo else item['certificate_expiry']
            uso_det['giorni_scaduto_all_uso'] = (card_date_comp - cert_exp_comp).days
        else:
            uso_det['giorni_scaduto_all_uso'] = 0
        details['dettaglio_usi_list'].append(uso_det)
        
    detailed_certs_table_data = [] 
    for cert_id, details in cert_details_map.items():
        expiries = {d.get('certificate_expiry') for d in details['dettaglio_usi_list'] if d.get('certificate_expiry')}
        scad_rec = "N/D"
        if expiries:
            valid_expiries = [dt for dt in expiries if isinstance(dt, datetime)]
            if valid_expiries:
                scad_rec = max(valid_expiries).strftime('%d/%m/%Y')

        range_p = details['range_su_scheda_counter'].most_common(1)[0][0] if details['range_su_scheda_counter'] else "N/D"
        tip_p = details['tipologie_scheda_associate_counter'].most_common(1)[0][0] if details['tipologie_scheda_associate_counter'] else "N/D"
        incongr_r_calc = details['usi_total_incongrui'] - details['usi_prima_emissione'] 
        detailed_certs_table_data.append({
            "ID Certificato": cert_id, "Utilizzi": details['utilizzi'], 
            "Tipologia Scheda Principale": tip_p,
            "Usi Congrui": details['usi_congrui'], 
            "Usi NON Congrui (Solo Regole)": incongr_r_calc, 
            "Usi Prima Emissione": details['usi_prima_emissione'],
            "Usi da Scaduto (non per emissione)": details['usi_scaduti_puri'], 
            "Scadenza (Più Recente Vista Parsata)": scad_rec,
            "Range Principale Usato (su Scheda)": range_p
        })
    detailed_certs_table_data.sort(key=lambda x: (-x["Usi Prima Emissione"], -x["Usi NON Congrui (Solo Regole)"], -x["Usi da Scaduto (non per emissione)"], -x["Utilizzi"])) 
    
    log_to_all(log_text_widget, "Aggregazione dati completata.", "DEBUG")
    log_to_all(log_text_widget, "Creazione finestra risultati Tkinter...", "DEBUG")
    
    result_window = tk.Toplevel(root); result_window_ref_store['ref'] = result_window 
    result_window.title(f"Analisi Schede Taratura - Report del {ANALYSIS_DATETIME.astimezone().strftime('%d/%m/%Y')}")
    result_window.geometry("1750x980") 
    style = ttk.Style(result_window)
    try: 
        sel_theme = 'vista' if 'vista' in style.theme_names() else 'clam' if 'clam' in style.theme_names() else 'default'
        style.theme_use(sel_theme); logger.info(f"Tema ttk: {sel_theme}")
    except tk.TclError: logger.warning(f"Tema ttk '{sel_theme}' non trovato.")
    style.configure("Treeview.Heading", font=('Segoe UI', 10, 'bold'), relief="groove")
    style.configure("Treeview", rowheight=28, font=('Segoe UI', 9)) 
    style.configure("TNotebook.Tab", font=('Segoe UI', 10, 'bold'), padding=[12, 6])
    style.configure("TLabelframe.Label", font=('Segoe UI', 11, 'bold'), padding=(0,0,0,5)) 
    style.configure("Accent.TButton", font=('Segoe UI', 10, 'bold'), padding=8) 
    style.configure("Hyperlink.TLabel", foreground="blue", font=('Segoe UI', 9, 'underline')) 
    
    notebook = ttk.Notebook(result_window, style="TNotebook")
    notebook.pack(expand=True, fill='both', padx=10, pady=(10,0)) 
    sugg_tab_interface_details['notebook_widget'] = notebook 
    
    tab_cruscotto = ttk.Frame(notebook, padding=10); notebook.add(tab_cruscotto, text=' Cruscotto Riepilogativo ')
    cruscotto_main_frame = ttk.Frame(tab_cruscotto) 
    cruscotto_main_frame.pack(expand=True, fill='both', padx=5, pady=5)
    
    cruscotto_main_frame.rowconfigure(0, weight=0) 
    cruscotto_main_frame.rowconfigure(1, weight=0) 
    cruscotto_main_frame.rowconfigure(2, weight=1) 
    cruscotto_main_frame.rowconfigure(3, weight=3) 
    cruscotto_main_frame.columnconfigure(0, weight=1) 

    stats_frame = ttk.LabelFrame(cruscotto_main_frame, text="Statistiche Generali", padding=(10,5))
    stats_frame.grid(row=0, column=0, sticky="ew", padx=5, pady=(0,5)) 
    stats_frame.columnconfigure(0, weight=1) 
    for item_txt in report_sections_generale_sintetico:
        def update_stats_wraplength(widget_to_update, parent_frame_for_width):
            parent_width = parent_frame_for_width.winfo_width()
            if parent_width > 1: 
                widget_to_update.config(wraplength=parent_width - 20) 
            else: 
                widget_to_update.after(100, lambda w=widget_to_update, p=parent_frame_for_width: update_stats_wraplength(w,p))

        lbl = ttk.Label(stats_frame, text=item_txt, anchor=tk.W, font=('Segoe UI',9))
        lbl.pack(fill=tk.X, padx=5, pady=1, anchor=tk.W) 
        lbl.after(50, lambda w=lbl, p=stats_frame: update_stats_wraplength(w,p)) 


    action_buttons_frame = ttk.Frame(cruscotto_main_frame)
    action_buttons_frame.grid(row=1, column=0, sticky="ew", padx=5, pady=(5,5))

    btn_stampa_anomalie = ttk.Button(action_buttons_frame, text="Stampa Report Anomalie (Word)",
                                     command=lambda: crea_e_apri_report_anomalie_word(
                                         log_text_widget, 
                                         result_window_ref_store.get('ref'), 
                                         human_errors_details_list, 
                                         temporal_alerts_entries,
                                         incongruent_rules_list,
                                         len(candidate_files), 
                                         validated_file_count 
                                     ), style="Accent.TButton")
    btn_stampa_anomalie.pack(side=tk.LEFT, padx=5, pady=5)

    def create_scrollable_section_frame(parent_widget, section_title, title_font_color="black"): 
        section_lf = ttk.LabelFrame(parent_widget, text=section_title, padding=(10,5))
        canvas = tk.Canvas(section_lf, borderwidth=0, highlightthickness=0) 
        vsb = ttk.Scrollbar(section_lf, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        
        vsb.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)
        
        inner_frame = ttk.Frame(canvas, padding=(5,5,15,5)) 
        inner_frame_window_id = canvas.create_window((0,0), window=inner_frame, anchor="nw", tags=f"inner_frame_tag_{section_title.replace(' ','_')}")

        def _on_inner_frame_configure(event_ignored): 
            canvas.configure(scrollregion=canvas.bbox("all"))
        
        def _on_canvas_configure(event_canvas_local): 
            canvas_width = event_canvas_local.width
            if canvas.find_withtag(inner_frame_window_id): 
                canvas.itemconfig(inner_frame_window_id, width=canvas_width) 
            canvas.configure(scrollregion=canvas.bbox("all"))

        inner_frame.bind("<Configure>", _on_inner_frame_configure)
        canvas.bind("<Configure>", _on_canvas_configure)
        inner_frame.columnconfigure(0, weight=1) 
        return section_lf, inner_frame

    alerts_cert_container_frame = ttk.Frame(cruscotto_main_frame)
    alerts_cert_container_frame.grid(row=2, column=0, sticky="nsew", padx=5, pady=(5,5)) 
    alerts_cert_container_frame.columnconfigure(0, weight=1); alerts_cert_container_frame.columnconfigure(1, weight=1) 
    alerts_cert_container_frame.rowconfigure(0, weight=1) 

    temporal_issues_lf, temp_inner = create_scrollable_section_frame(alerts_cert_container_frame, "Avvisi: Uso Certificati Temporalmente Non Validi")
    temporal_issues_lf.grid(row=0, column=0, sticky="nsew", padx=(0,5), pady=5) 
    if temporal_alerts_entries:
        for idx, entry in enumerate(temporal_alerts_entries):
            ef = ttk.Frame(temp_inner); ef.pack(fill=tk.X, pady=(5,0)); tf = ttk.Frame(ef); tf.pack(fill=tk.X) 
            fl = ttk.Label(tf, text=f"File: {entry['file_name']}", style="Hyperlink.TLabel", cursor="hand2")
            fl.pack(side=tk.LEFT, padx=(0,10))
            _fp = entry.get('file_path'); _fn = entry.get('file_name',"file") 
            fl.bind("<Button-1>", lambda event, p=_fp, fn=_fn: on_file_click(p, fn, open_file_direct=True, event=event))
            
            sb = ttk.Button(tf, text="Alternative", width=12, style="Accent.TButton", command=partial(launch_suggestion_tab_and_search, entry))
            sb.pack(side=tk.RIGHT, padx=5)
            
            pd_text = f"USATO PRIMA DELL'EMISSIONE (Emiss.Cert: {entry.get('data_emissione_presunta','N/A')})" if entry['alert_type'] == 'premature_emission' else f"USATO DA SCADUTO (Scad.Cert: {entry.get('expiry_date_str','N/P')})"
            color_text = "red" if entry['alert_type'] == 'premature_emission' else "darkorange"
            tk.Label(ef, text=f"PROBLEMA: {pd_text}", fg=color_text, font=('Segoe UI', 9, 'bold'), anchor=tk.W, background=style.lookup('TFrame','background')).pack(fill=tk.X, anchor=tk.W, padx=10, pady=(5,0))
            
            congr_notes_display_temp = entry.get('congruency_notes','')
            if entry['alert_type'] == 'premature_emission' and congr_notes_display_temp.startswith("Dettaglio:"):
                 original_congr_note = congr_notes_display_temp 
                 congr_notes_display_temp = original_congr_note.split("ma il certificato è stato emesso")[0].replace(f"Dettaglio: errato per EMISSIONE. Cert.'{entry['certificate_id']}'({entry['modello_strumento_campione_usato']}) usato il {entry['card_date']:%d/%m/%Y}", "").strip()
                 # CORREZIONE SINTASSI:
                 if congr_notes_display_temp.endswith(".") or not congr_notes_display_temp:
                     congr_notes_display_temp = "Verificato per regole."
                 else:
                     congr_notes_display_temp = f"Regole: {congr_notes_display_temp}"
            else: 
                congr_notes_display_temp = f"Congr.Regole: {'OK' if entry.get('is_congruent') else 'NON OK' if entry.get('is_congruent') is False else 'N/V'} ({entry.get('congruency_notes','')})"

            wrap_len = temp_inner.winfo_width() - 50 if temp_inner.winfo_width() > 50 else 500 
            tk.Label(ef, text=f"  Scheda: {entry.get('card_date_str','N/D')} - Cert: {entry.get('cert_id','N/D')} (Mod.Camp: {entry.get('modello_strumento_campione_usato','N/D')}, Range: {entry.get('range','N/D')})\n  {congr_notes_display_temp}", font=('Consolas',9), anchor=tk.W, justify=tk.LEFT, background=style.lookup('TFrame','background'), wraplength=wrap_len).pack(fill=tk.X, anchor=tk.W, padx=20, pady=(0,5))
            if idx < len(temporal_alerts_entries) - 1: ttk.Separator(ef, orient='horizontal').pack(fill='x', pady=10, padx=5)
    else:
        tk.Label(temp_inner, text="Nessun uso prima emissione o da scaduti rilevato.", fg="darkgreen", font=('Segoe UI',10,'bold'), anchor=tk.W, background=style.lookup('TFrame','background')).pack(anchor=tk.W)
    
    incongr_lf, incongr_inner = create_scrollable_section_frame(alerts_cert_container_frame, "Avvisi: Certificati Non Congrui (per Regole)")
    incongr_lf.grid(row=0, column=1, sticky="nsew", padx=(5,0), pady=5) 
    if incongruent_rules_list:
        for idx, item_inc in enumerate(incongruent_rules_list):
            ef_inc = ttk.Frame(incongr_inner); ef_inc.pack(fill=tk.X, pady=(5,0)); tf_inc = ttk.Frame(ef_inc); tf_inc.pack(fill=tk.X)
            fl_inc = ttk.Label(tf_inc, text=f"File: {item_inc['file_name']}", style="Hyperlink.TLabel", cursor="hand2")
            fl_inc.pack(side=tk.LEFT, padx=(0,10))
            _fp_item_inc = item_inc.get('file_path'); _fn_item_inc = item_inc.get('file_name',"file")
            fl_inc.bind("<Button-1>", lambda event, p=_fp_item_inc, fn=_fn_item_inc: on_file_click(p, fn, open_file_direct=True, event=event))
            
            sb_inc = ttk.Button(tf_inc, text="Alternative", width=12, style="Accent.TButton", command=partial(launch_suggestion_tab_and_search, item_inc))
            sb_inc.pack(side=tk.RIGHT, padx=5)
            
            scad_info_text_inc = " (USATO SCADUTO!)" if item_inc.get('is_expired_at_use') else "" 
            wrap_len_inc = incongr_inner.winfo_width() - 50 if incongr_inner.winfo_width() > 50 else 500
            
            congr_notes_display_inc = item_inc.get('congruency_notes','')
            reason_text_inc = f"  MOTIVO NON CONGRUO: {congr_notes_display_inc}" if not congr_notes_display_inc.startswith("Dettaglio:") else f"  {congr_notes_display_inc}"

            tk.Label(ef_inc, text=f"  Scheda: {item_inc.get('card_date_str','N/D')} (Tip: {item_inc.get('tipologia_strumento_scheda','N/D')}, L9: {item_inc.get('modello_L9_scheda','N/A')})\n  Cert: {item_inc.get('cert_id','N/D')} (Mod.Camp: {item_inc.get('modello_strumento_campione_usato','N/D')}, Scad: {item_inc.get('certificate_expiry_str','N/P')}{scad_info_text_inc})\n{reason_text_inc}", font=('Consolas',9), anchor=tk.W, justify=tk.LEFT, background=style.lookup('TFrame','background'), wraplength=wrap_len_inc).pack(fill=tk.X, anchor=tk.W, padx=10, pady=(5,5))
            if idx < len(incongruent_rules_list) - 1: ttk.Separator(ef_inc, orient='horizontal').pack(fill='x', pady=10, padx=5)
    else:
        tk.Label(incongr_inner, text="Nessun uso non congruo per regole rilevato.", fg="darkgreen", font=('Segoe UI',10,'bold'), anchor=tk.W, background=style.lookup('TFrame','background')).pack(anchor=tk.W)
    
    comp_errors_lf, comp_errors_inner_frame = create_scrollable_section_frame(cruscotto_main_frame, "Errori di Compilazione Scheda (Range/UM, SP, L9)")
    comp_errors_lf.grid(row=3, column=0, sticky="nsew", padx=5, pady=(10,5)) 
    
    comp_errors_classici_list = [err for err in human_errors_details_list if not err['key'].startswith("COMP_")]

    if comp_errors_classici_list:
        comp_errors_classici_list.sort(key=lambda x: (x['file'], x['key'])) 
        for idx, error_detail in enumerate(comp_errors_classici_list):
            err_entry_frame = ttk.Frame(comp_errors_inner_frame)
            err_entry_frame.pack(fill=tk.X, pady=(2,0), padx=5)
            err_entry_frame.columnconfigure(1, weight=1) 

            file_label_err = ttk.Label(err_entry_frame, text=f"{error_detail['file']}:", style="Hyperlink.TLabel", cursor="hand2", font=('Segoe UI', 9, 'bold'))
            file_label_err.grid(row=0, column=0, sticky=tk.NW, padx=(0,5))
            _fp_err = error_detail.get('path'); _fn_err = error_detail.get('file',"file")
            file_label_err.bind("<Button-1>", lambda event, p=_fp_err, fn=_fn_err: on_file_click(p, fn, open_file_direct=True, event=event))

            error_message_desc = human_error_messages_map_descriptive.get(error_detail['key'], f"Errore sconosciuto: {error_detail['key']}")
            
            msg_label = tk.Label(err_entry_frame, text=error_message_desc, fg="firebrick", font=('Segoe UI', 9),
                                 anchor=tk.W, justify=tk.LEFT,
                                 background=style.lookup('TFrame','background')) 
            msg_label.grid(row=0, column=1, sticky=tk.EW, padx=(5,0))
            msg_label.after(50, lambda w=msg_label, p=err_entry_frame, fl=file_label_err: update_label_wraplength(w, p, fl))
            
            if idx < len(comp_errors_classici_list) - 1:
                ttk.Separator(comp_errors_inner_frame, orient='horizontal').pack(fill='x', pady=(5,3), padx=10)
    else:
        tk.Label(comp_errors_inner_frame, text="Nessun errore di compilazione (Range/UM, SP, L9) rilevato.", fg="darkgreen", font=('Segoe UI', 10, 'bold'), anchor=tk.W, background=style.lookup('TFrame','background')).pack(anchor=tk.W, padx=5, pady=5)
    log_to_all(log_text_widget, "Tab Cruscotto Riepilogativo creata/aggiornata.", "DEBUG")

# --- FINE PARTE 4 ---
# --- PARTE 5 (Continuazione GUI e main) ---
    # Funzione helper per aggiornare il wraplength delle label, definita una volta qui.
    # Usata nel tab Cruscotto e potenzialmente in altri.
    def update_label_wraplength(widget, parent_for_width, fixed_width_widget=None):
        # Assicura che i widget esistano prima di chiamare winfo_width
        if not widget.winfo_exists() or not parent_for_width.winfo_exists():
            return
        if fixed_width_widget and not fixed_width_widget.winfo_exists():
            fixed_width_widget = None # Tratta come se non ci fosse se non esiste

        parent_width = parent_for_width.winfo_width()
        offset_width = fixed_width_widget.winfo_width() if fixed_width_widget else 0
        
        # Calcola il wraplength solo se la larghezza del parent è significativa
        if parent_width > 20 : # Evita calcoli con larghezze iniziali nulle o troppo piccole
            new_wraplength = parent_width - offset_width - 30 # Sottrai larghezza widget fisso e un po' di padding
            if new_wraplength > 50 : # Imposta solo se il risultato è ragionevole
                widget.config(wraplength=new_wraplength)
        else: # Se la larghezza non è ancora pronta, riprova dopo un breve ritardo
            widget.after(100, lambda w=widget, p=parent_for_width, fw=fixed_width_widget: update_label_wraplength(w, p, fw))


    # --- Tab Dettaglio Utilizzo Certificati ---
    tab_cert_details = ttk.Frame(notebook, padding=10)
    notebook.add(tab_cert_details, text=' Dettaglio Utilizzo Certificati ')
    cols_tree_new = ["ID Certificato", "Utilizzi", "Tipologia Scheda Principale",
                     "Usi Congrui", "Usi NON Congrui (Solo Regole)", "Usi Prima Emissione",
                     "Usi da Scaduto (non per emissione)", "Scadenza (Più Recente Vista Parsata)",
                     "Range Principale Usato (su Scheda)"]
    tree_cert = ttk.Treeview(tab_cert_details, columns=cols_tree_new, show='headings', selectmode="browse")

    last_clicked_item_id_for_toggle = [None] # Usa una lista per permettere la modifica nella closure

    def on_tree_item_interaction(event_arg, is_double_click):
        nonlocal last_clicked_item_id_for_toggle # Per gestire l'apri/chiudi con singolo click
        current_result_window = result_window_ref_store.get('ref') # Per parent messagebox
        item_id = tree_cert.identify_row(event_arg.y) # Identifica l'item cliccato
        if not item_id: # Click su area vuota
            logger.debug("on_tree_item_interaction: nessun item identificato a y=" + str(event_arg.y))
            return

        logger.debug(f"on_tree_item_interaction: item_id='{item_id}', is_double_click={is_double_click}")
        current_item_id = item_id
        is_child = bool(tree_cert.parent(current_item_id)) # Verifica se è un item figlio (dettaglio uso)
        logger.debug(f"current_item_id='{current_item_id}', is_child={is_child}")

        def _handle_child_click(item_id_for_child): # Azione per click su item figlio
            tags = tree_cert.item(item_id_for_child, 'tags')
            logger.debug(f"Child item '{item_id_for_child}' tags: {tags}")
            if len(tags) > 0 : # L'ultimo tag dovrebbe essere il file_path
                file_path_from_tag = tags[-1] # Assume che il file_path sia l'ultimo tag

                # Validazione del percorso estratto dal tag
                if not file_path_from_tag or not isinstance(file_path_from_tag, str) or not os.path.exists(file_path_from_tag) :
                    logger.warning(f"Percorso file non valido o non esistente associato all'item '{item_id_for_child}' (tag: '{file_path_from_tag}').")
                    if current_result_window and current_result_window.winfo_exists():
                            messagebox.showwarning("Percorso Non Valido", f"Il percorso associato a questo item non è valido o il file non esiste:\n{file_path_from_tag}", parent=current_result_window)
                    return

                # Estrai il nome del file dal testo visualizzato per un messaggio più user-friendly
                display_text_child_raw = tree_cert.item(item_id_for_child, 'values')[0] if tree_cert.item(item_id_for_child, 'values') else ""
                match_child = re.search(r"File:\s*([^(\s]+)", display_text_child_raw) # Cerca "File: nomefile"
                filename_for_display = match_child.group(1) if match_child else "File Sconosciuto"

                logger.info(f"Azione click su child: File='{filename_for_display}', Path='{file_path_from_tag}'. Chiamata a on_file_click.")
                on_file_click(file_path_from_tag, filename_for_display, open_file_direct=True, event=event_arg if is_child else None)
            else:
                logger.warning(f"Nessun tag trovato per l'item child '{item_id_for_child}'. Impossibile estrarre il percorso.")


        if is_double_click:
            if not is_child: # Doppio click su un item PARENT (ID Certificato)
                values = tree_cert.item(current_item_id, 'values')
                if values and sugg_tab_interface_details.get('cert_id_entry_widget'): # Assicura che il tab suggerimenti sia pronto
                    cert_id_to_check = values[cols_tree_new.index("ID Certificato")]
                    utilizzi_count = int(values[cols_tree_new.index("Utilizzi")])
                    range_to_suggest = values[cols_tree_new.index("Range Principale Usato (su Scheda)")]
                    
                    # Determina la data di riferimento più recente per questo certificato
                    data_riferimento_per_suggerimento_dt = ANALYSIS_DATETIME.astimezone() # Default alla data di analisi
                    if cert_id_to_check in cert_details_map:
                        utilizzi_date_cert_obj = sorted(list(cert_details_map[cert_id_to_check]['date_utilizzo_obj_set']), reverse=True)
                        if utilizzi_date_cert_obj: # Se ci sono date di utilizzo registrate
                            data_riferimento_per_suggerimento_dt = utilizzi_date_cert_obj[0] # Prendi la più recente
                    data_riferimento_str = data_riferimento_per_suggerimento_dt.strftime('%d/%m/%Y')

                    # Popola i campi nel tab suggerimenti
                    sugg_tab_interface_details['cert_id_entry_widget'].delete(0, tk.END)
                    sugg_tab_interface_details['cert_id_entry_widget'].insert(0, cert_id_to_check)
                    sugg_tab_interface_details['range_entry_widget'].delete(0, tk.END)
                    sugg_tab_interface_details['range_entry_widget'].insert(0, range_to_suggest if range_to_suggest != "N/D" else "")
                    sugg_tab_interface_details['date_entry_widget'].delete(0, tk.END)
                    sugg_tab_interface_details['date_entry_widget'].insert(0, data_riferimento_str)

                    # Attiva il tab suggerimenti e, se necessario, esegui la ricerca
                    sugg_tab_interface_details['notebook_widget'].select(sugg_tab_interface_details['suggerimenti_tab_widget'])
                    # Condizioni per ricerca automatica suggerimenti
                    usi_prima_emissione_count = int(values[cols_tree_new.index("Usi Prima Emissione")])
                    usi_da_scaduto_count = int(values[cols_tree_new.index("Usi da Scaduto (non per emissione)")])
                    usi_non_congrui_regole_count = int(values[cols_tree_new.index("Usi NON Congrui (Solo Regole)")])

                    if utilizzi_count > SOGLIA_PER_SUGGERIMENTO_ALTERNATIVO or usi_prima_emissione_count > 0 or usi_da_scaduto_count > 0 or usi_non_congrui_regole_count > 0 :
                        sugg_tab_interface_details['search_function_widget']() # Esegui la ricerca
                    else: # Informa l'utente che la ricerca non è automatica sotto soglia
                        sugg_text_widget_ref = sugg_tab_interface_details.get('results_text_widget')
                        if sugg_text_widget_ref:
                            sugg_text_widget_ref.config(state=tk.NORMAL); sugg_text_widget_ref.delete("1.0", tk.END)
                            sugg_text_widget_ref.insert(tk.END, f"Certificato '{cert_id_to_check}' (Utilizzi: {utilizzi_count}, Emiss.: {usi_prima_emissione_count}, Scad.: {usi_da_scaduto_count}, Incongr.: {usi_non_congrui_regole_count}).\nNon supera soglia ({SOGLIA_PER_SUGGERIMENTO_ALTERNATIVO} utilizzi) o problemi per suggerimento automatico.\nClicca 'Cerca Alternative' per forzare.", "sugg_info")
                            sugg_text_widget_ref.config(state=tk.DISABLED)
            elif is_child: # Doppio click su un item FIGLIO
                logger.debug(f"Doppio click su child item: '{current_item_id}'. Chiamata a _handle_child_click.")
                _handle_child_click(current_item_id)

        else: # Singolo click
            if not is_child: # Singolo click su un item PARENT: apri/chiudi figli
                if current_item_id == last_clicked_item_id_for_toggle[0]: # Click sullo stesso item, inverti lo stato
                    tree_cert.item(current_item_id, open=not tree_cert.item(current_item_id, 'open'))
                    last_clicked_item_id_for_toggle[0] = None # Resetta per il prossimo click
                else: # Click su un nuovo item parent
                    if last_clicked_item_id_for_toggle[0] and tree_cert.exists(last_clicked_item_id_for_toggle[0]): # Chiudi il precedente se aperto
                        tree_cert.item(last_clicked_item_id_for_toggle[0], open=False)
                    tree_cert.item(current_item_id, open=True) # Apri il nuovo
                    last_clicked_item_id_for_toggle[0] = current_item_id
            else: # Singolo click su un item FIGLIO: selezionalo
                logger.debug(f"Singolo click su child item: '{current_item_id}'.")
                tree_cert.selection_set(current_item_id) # Evidenzia l'item figlio


    tree_cert.bind("<Double-1>", partial(on_tree_item_interaction, is_double_click=True))
    tree_cert.bind("<Button-1>", partial(on_tree_item_interaction, is_double_click=False)) # Gestisce anche apertura/chiusura parent

    # Menu contestuale per item figli (click destro)
    child_item_menu = tk.Menu(result_window_ref_store.get('ref'), tearoff=0)
    def show_child_item_menu(event_arg): 
        current_result_window = result_window_ref_store.get('ref')
        item_id = tree_cert.identify_row(event_arg.y) 
        if item_id and tree_cert.parent(item_id): # Solo per item figli
            tree_cert.selection_set(item_id) # Seleziona l'item cliccato
            tags = tree_cert.item(item_id, 'tags')
            if len(tags) > 0:
                file_path_for_menu = tags[-1] 
                if not file_path_for_menu or not isinstance(file_path_for_menu, str) or not os.path.exists(file_path_for_menu):
                    logger.warning(f"Menu contestuale: percorso non valido o file non esistente per item '{item_id}' (tag: '{file_path_for_menu}')")
                    return

                display_text_raw = tree_cert.item(item_id, 'values')[0] if tree_cert.item(item_id, 'values') else ""
                match = re.search(r"File:\s*([^(\s]+)", display_text_raw)
                filename_for_menu = match.group(1) if match else "File Sconosciuto"

                child_item_menu.delete(0, tk.END) # Pulisci menu precedente
                child_item_menu.add_command(label=f"Copia percorso: {filename_for_menu}", command=lambda p=file_path_for_menu: pyperclip.copy(p) or (messagebox.showinfo("Info", f"Percorso copiato:\n{p}", parent=current_result_window) if current_result_window and current_result_window.winfo_exists() else None) )
                child_item_menu.add_command(label=f"Apri file: {filename_for_menu}", command=lambda p=file_path_for_menu, f_name=filename_for_menu: on_file_click(p, f_name, open_file_direct=True, event=event_arg)) 
                child_item_menu.add_command(label="Apri cartella del file", command=lambda p=file_path_for_menu, f_name=filename_for_menu: on_file_click(p, f_name, open_file_direct=False, event=event_arg)) 
                try: child_item_menu.tk_popup(event_arg.x_root, event_arg.y_root)
                finally: child_item_menu.grab_release()
    tree_cert.bind("<Button-3>", show_child_item_menu) # <Button-3> è il click destro

    # Configurazione colonne e popolamento Treeview
    col_widths = {"ID Certificato":180, "Utilizzi":60, "Tipologia Scheda Principale":170,
                  "Usi Congrui":70, "Usi NON Congrui (Solo Regole)":140, "Usi Prima Emissione":110,
                  "Usi da Scaduto (non per emissione)":150, "Scadenza (Più Recente Vista Parsata)":170,
                  "Range Principale Usato (su Scheda)":200}
    for col_name in cols_tree_new:
        width = col_widths.get(col_name, 120); min_w = max(70, width - 50)
        stretch_col = col_name in ["ID Certificato", "Range Principale Usato (su Scheda)", "Tipologia Scheda Principale"]
        tree_cert.heading(col_name, text=col_name, anchor=tk.W)
        tree_cert.column(col_name, width=width, minwidth=min_w, stretch=tk.YES if stretch_col else tk.NO, anchor=tk.W)
    
    tree_cert.tag_configure('oddrow', background='#E8E8E8' if style.theme_use() != 'vista' else '#F3F3F3')
    tree_cert.tag_configure('evenrow', background='white')
    tree_cert.tag_configure('parent_has_premature_uses', foreground='red', font=tkFont.Font(weight='bold'))
    tree_cert.tag_configure('parent_has_incongruent_rules_uses', foreground='purple', font=tkFont.Font(slant='italic'))
    tree_cert.tag_configure('parent_has_expired_uses', foreground='darkorange')
    child_font = tkFont.Font(family='Consolas', size=8); tree_cert.tag_configure('child_base', font=child_font, background='#FAFAFA') 
    tree_cert.tag_configure('child_ok', foreground='darkgreen')
    tree_cert.tag_configure('child_premature', foreground='red', font=tkFont.Font(family='Consolas', size=8, weight='bold'))
    tree_cert.tag_configure('child_incongruent_rule', foreground='purple', font=tkFont.Font(family='Consolas', size=8, slant='italic'))
    tree_cert.tag_configure('child_expired_pure', foreground='chocolate')
    tree_cert.tag_configure('child_congruency_nv', foreground='gray50')

    for i, row_data in enumerate(detailed_certs_table_data):
        display_vals = [row_data.get(col, "N/A") for col in cols_tree_new]
        parent_base_tag = 'evenrow' if i % 2 == 0 else 'oddrow'; parent_prob_tags = []
        if row_data.get("Usi Prima Emissione", 0) > 0: parent_prob_tags.append('parent_has_premature_uses')
        if row_data.get("Usi NON Congrui (Solo Regole)", 0) > 0: parent_prob_tags.append('parent_has_incongruent_rules_uses')
        if row_data.get("Usi da Scaduto (non per emissione)", 0) > 0: parent_prob_tags.append('parent_has_expired_uses')
        final_parent_tags = tuple([parent_base_tag] + parent_prob_tags)
        parent_item_id = tree_cert.insert("", "end", values=display_vals, tags=final_parent_tags, open=False) 
        
        cert_id_curr = row_data["ID Certificato"]; usi_dett = cert_details_map.get(cert_id_curr, {}).get('dettaglio_usi_list', [])
        if usi_dett:
            usi_dett.sort(key=lambda x: (x.get('card_date', datetime.min.replace(tzinfo=timezone.utc)), x.get('file_name', '')), reverse=True) 
            for uso_info in usi_dett:
                child_style_tags = [] 
                if uso_info.get('used_before_emission'): child_style_tags.append('child_premature')
                elif uso_info.get('is_congruent') is False: child_style_tags.append('child_incongruent_rule')
                elif uso_info.get('is_expired_at_use', False): child_style_tags.append('child_expired_pure')
                elif uso_info.get('is_congruent') is None: child_style_tags.append('child_congruency_nv')
                else: child_style_tags.append('child_ok')

                file_path_val = uso_info.get('file_path', "") 
                final_child_tags = ('child_base',) + tuple(child_style_tags) + (file_path_val,) 

                scad_str = f"SCADUTO di {uso_info['giorni_scaduto_all_uso']}gg" if uso_info.get('giorni_scaduto_all_uso',0) > 0 else 'OK'
                congr_str = "N/D"; congr_notes_short = uso_info.get('congruency_notes','')[:50]+'...' if len(uso_info.get('congruency_notes','')) > 50 else uso_info.get('congruency_notes','')
                if uso_info.get('is_congruent') is True: congr_str = "Congruo"
                elif uso_info.get('is_congruent') is False: congr_str = f"NON Congruo ({congr_notes_short})"
                else: congr_str = f"Congr. N/V ({congr_notes_short})"
                card_dt_str = uso_info['card_date'].strftime('%d/%m/%Y') if uso_info.get('card_date') else "N/D"
                
                child_vals = [""] * len(cols_tree_new) 
                child_vals[0] = f"  └─File: {uso_info['file_name']} (Scheda: {card_dt_str})" 
                child_vals[1] = f"Range Scheda: {uso_info.get('range_su_scheda','N/A')}" 
                child_vals[2] = f"Tip.Strum: {uso_info.get('tipologia_strumento_scheda','N/A')} (L9: {uso_info.get('modello_L9_scheda','N/A')})"
                child_vals[3] = f"Mod.Campione: {uso_info.get('modello_strumento_campione_usato','N/A')}"
                child_vals[4] = f"Stato Scadenza: {scad_str}"
                child_vals[5] = f"Stato Congruità: {congr_str}"
                tree_cert.insert(parent_item_id, "end", values=child_vals, tags=final_child_tags)

    tree_vsb = ttk.Scrollbar(tab_cert_details, orient="vertical", command=tree_cert.yview)
    tree_hsb = ttk.Scrollbar(tab_cert_details, orient="horizontal", command=tree_cert.xview)
    tree_cert.configure(yscrollcommand=tree_vsb.set, xscrollcommand=tree_hsb.set)
    tree_vsb.pack(side=tk.RIGHT, fill=tk.Y); tree_hsb.pack(side=tk.BOTTOM, fill=tk.X)
    tree_cert.pack(expand=True, fill='both')
    log_to_all(log_text_widget, "Tab Dettaglio Certificati creata.", "DEBUG")

    # --- Tab Compilatore ---
    tab_compilatore = ttk.Frame(notebook, padding=10)
    notebook.add(tab_compilatore, text=' Compilatore Automatico ')
    compilatore_main_frame = ttk.Frame(tab_compilatore)
    compilatore_main_frame.pack(expand=True, fill='both', padx=5, pady=5)
    compilatore_main_frame.rowconfigure(0, weight=0) 
    compilatore_main_frame.rowconfigure(1, weight=1) 
    compilatore_main_frame.columnconfigure(0, weight=1)

    compilatore_action_frame = ttk.Frame(compilatore_main_frame)
    compilatore_action_frame.grid(row=0, column=0, sticky="ew", pady=(0,10))
    
    def _get_cell_coordinate_for_compilation(campo_logico, tipo_file_scheda_locale):
        if tipo_file_scheda_locale == 'analogico':
            return {
                'DATA_COMPILAZIONE': SCHEDA_ANA_CELL_DATA_COMPILAZIONE, 'ESECUTORE': SCHEDA_ANA_CELL_ESECUTORE,
                'SUPERVISORE_ISAB': SCHEDA_ANA_CELL_SUPERVISORE_ISAB, 'ODC': SCHEDA_ANA_CELL_ODC,
                'PDL': SCHEDA_ANA_CELL_PDL, 'CONTRATTO_COEMI': SCHEDA_ANA_CELL_CONTRATTO_COEMI
            }.get(campo_logico)
        elif tipo_file_scheda_locale == 'digitale':
            return {
                'DATA_COMPILAZIONE': SCHEDA_DIG_CELL_DATA_COMPILAZIONE, 'ESECUTORE': SCHEDA_DIG_CELL_ESECUTORE,
                'SUPERVISORE_ISAB': SCHEDA_DIG_CELL_SUPERVISORE_ISAB, 'ODC': SCHEDA_DIG_CELL_ODC,
                'PDL': SCHEDA_DIG_CELL_PDL, 'CONTRATTO_COEMI': SCHEDA_DIG_CELL_CONTRATTO_COEMI
            }.get(campo_logico)
        return None

    def _estrai_odc_normalizzato_da_sorgente(valore_sorgente_odc):
        if pd.isna(valore_sorgente_odc):
            return None
        s = str(valore_sorgente_odc).strip()
        
        match_5400_ext = re.search(r"(5400\d{6}(?:[-\s]?\d{1,2})?)", s)
        if match_5400_ext:
            extracted = match_5400_ext.group(1)
            extracted = re.sub(r"(\d)\s+(\d{1,2})$", r"\1-\2", extracted)
            return extracted.replace("/", "-") 
            
        match_slash = re.search(r"(\d{3,})[-\s/]*(\d{1,2})", s) 
        if match_slash:
            if not match_slash.group(1).startswith("5400") or len(match_slash.group(1)) < 10 :
                 return f"{match_slash.group(1)}-{match_slash.group(2)}"
        
        first_line = s.split('\n')[0].strip() 
        if re.match(r"^[0-9]+(?:-[0-9]+)*$", first_line) and not re.match(r"^\d{3}-\d{2}$", first_line) :
            if first_line.startswith("5400") and len(re.sub(r"[^0-9]","",first_line)) >= 10 :
                 return first_line.replace("/", "-") 
            elif not first_line.startswith("5400"): 
                 return first_line.replace("/", "-")
        return None 


    def esegui_compilazione_schede():
        log_to_all(log_text_widget, "Avvio compilazione automatica schede...", "INFO")
        
        if not FILE_DATI_COMPILAZIONE_SCHEDE or not os.path.exists(FILE_DATI_COMPILAZIONE_SCHEDE):
            msg = f"File dati compilazione ({FILE_DATI_COMPILAZIONE_SCHEDE}) non trovato o non specificato."
            log_to_all(log_text_widget, msg, "ERROR")
            messagebox.showerror("Errore File Compilazione", f"{msg}\nControllare parametri.xlsm (B4).", parent=result_window_ref_store.get('ref'))
            return

        log_to_all(log_text_widget, f"Lettura dati da: {FILE_DATI_COMPILAZIONE_SCHEDE}, Foglio: '{NOME_FOGLIO_DATI_COMPILAZIONE}'", "INFO")

        try:
            df_sorgente_dati = pd.read_excel(FILE_DATI_COMPILAZIONE_SCHEDE,
                                             sheet_name=NOME_FOGLIO_DATI_COMPILAZIONE,
                                             engine='openpyxl',
                                             header=0) 
            
            log_to_all(log_text_widget, f"Lette {len(df_sorgente_dati)} righe (dati, escluso header) dal foglio '{NOME_FOGLIO_DATI_COMPILAZIONE}'.", "INFO")
            
            col_indici_necessari_max = max(COL_IDX_COMP_DATA, COL_IDX_COMP_ESECUTORE, 
                                           COL_IDX_COMP_SUPERVISORE, COL_IDX_COMP_ODC, COL_IDX_COMP_PDL)
            if df_sorgente_dati.shape[1] <= col_indici_necessari_max:
                msg = f"Il file sorgente dati '{NOME_FOGLIO_DATI_COMPILAZIONE}' non ha abbastanza colonne. Trovate {df_sorgente_dati.shape[1]}, necessarie almeno {col_indici_necessari_max + 1} per accedere a tutti i campi specificati (A,B,D,E,F)."
                log_to_all(log_text_widget, msg, "ERROR")
                messagebox.showerror("Errore Struttura File Sorgente", msg, parent=result_window_ref_store.get('ref'))
                return

            modifiche_effettuate_conteggio = 0
            global schede_analizzate_info_list 

            chiavi_rilevanti_per_ricerca = { # Definito qui come richiesto
                KEY_COMP_ANA_DATA_COMP_MANCANTE, KEY_COMP_DIG_DATA_COMP_MANCANTE,
                KEY_COMP_ANA_ODC_MANCANTE, KEY_COMP_DIG_ODC_MANCANTE,
                KEY_COMP_ANA_PDL_MANCANTE, KEY_COMP_DIG_PDL_MANCANTE,
                KEY_COMP_ANA_ESECUTORE_MANCANTE, KEY_COMP_DIG_ESECUTORE_MANCANTE,
                KEY_COMP_ANA_SUPERVISORE_MANCANTE, KEY_COMP_DIG_SUPERVISORE_MANCANTE
            }

            for scheda_info_originale in schede_analizzate_info_list: 
                scheda_info = scheda_info_originale.copy() 
                
                file_da_modificare_originale_xls_o_xlsx = scheda_info["file_path"]
                nome_file_scheda_log_corrente = scheda_info["base_filename"] # Nome file per i log
                tipo_file_scheda = scheda_info["file_type"]
                campi_da_compilare_su_scheda = scheda_info["campi_mancanti"].copy() 
                
                if not campi_da_compilare_su_scheda:
                    log_to_all(log_text_widget, f"Nessun campo rilevato per la compilazione in '{nome_file_scheda_log_corrente}'. Salto.", "DEBUG")
                    continue

                log_to_all(log_text_widget, f"Processo scheda: {nome_file_scheda_log_corrente}. Campi rilevati per compilazione: {campi_da_compilare_su_scheda}", "DEBUG")
                
                dati_effettivi_da_scrivere = {} 
                scheda_effettivamente_modificata_in_questo_ciclo = False

                key_contr_manc_tipo = KEY_COMP_ANA_CONTRATTO_MANCANTE if tipo_file_scheda == 'analogico' else KEY_COMP_DIG_CONTRATTO_MANCANTE
                key_contr_div_tipo = KEY_COMP_ANA_CONTRATTO_DIVERSO if tipo_file_scheda == 'analogico' else KEY_COMP_DIG_CONTRATTO_DIVERSO
                
                coord_contratto = _get_cell_coordinate_for_compilation('CONTRATTO_COEMI', tipo_file_scheda)
                if coord_contratto and (key_contr_manc_tipo in campi_da_compilare_su_scheda or key_contr_div_tipo in campi_da_compilare_su_scheda):
                    dati_effettivi_da_scrivere[coord_contratto] = VALORE_ATTESO_CONTRATTO_COEMI
                    log_to_all(log_text_widget, f"[{nome_file_scheda_log_corrente}] Contratto Coemi sarà impostato a '{VALORE_ATTESO_CONTRATTO_COEMI}' in {coord_contratto}.", "INFO")
                    campi_da_compilare_su_scheda.discard(key_contr_manc_tipo)
                    campi_da_compilare_su_scheda.discard(key_contr_div_tipo)

                riga_sorgente_trovata_series = None
                campi_che_richiedono_ricerca_sorgente = any(k in campi_da_compilare_su_scheda for k in chiavi_rilevanti_per_ricerca)

                if campi_che_richiedono_ricerca_sorgente:
                    pdl_scheda = scheda_info.get("pdl_val")
                    odc_scheda = scheda_info.get("odc_val_scheda")

                    if pdl_scheda and not pd.isna(pdl_scheda) and str(pdl_scheda).strip():
                        try:
                            pdl_da_cercare = str(pdl_scheda).strip()
                            colonna_pdl_sorgente_str = df_sorgente_dati.iloc[:, COL_IDX_COMP_PDL].astype(str).str.strip()
                            corrispondenze_pdl_df = df_sorgente_dati[colonna_pdl_sorgente_str.str.contains(re.escape(pdl_da_cercare), case=False, na=False, regex=True)]
                            if not corrispondenze_pdl_df.empty:
                                if len(corrispondenze_pdl_df) > 1 and COL_IDX_COMP_DATA < df_sorgente_dati.shape[1] :
                                    corrispondenze_pdl_df_copy = corrispondenze_pdl_df.copy()
                                    corrispondenze_pdl_df_copy.loc[:, 'parsed_date_sorgente'] = pd.to_datetime(corrispondenze_pdl_df_copy.iloc[:, COL_IDX_COMP_DATA], errors='coerce')
                                    corrispondenze_pdl_df_copy.dropna(subset=['parsed_date_sorgente'], inplace=True)
                                    if not corrispondenze_pdl_df_copy.empty:
                                        riga_sorgente_trovata_series = corrispondenze_pdl_df_copy.loc[corrispondenze_pdl_df_copy['parsed_date_sorgente'].idxmax()]
                                    else: riga_sorgente_trovata_series = corrispondenze_pdl_df.iloc[0]
                                else: riga_sorgente_trovata_series = corrispondenze_pdl_df.iloc[0]
                                log_to_all(log_text_widget, f"[{nome_file_scheda_log_corrente}] Trovata riga sorgente per PDL '{pdl_da_cercare}'.", "DEBUG")
                        except Exception as e_pdl_s: log_to_all(log_text_widget, f"[{nome_file_scheda_log_corrente}] Errore ricerca PDL: {e_pdl_s}", "WARNING")

                    if riga_sorgente_trovata_series is None and odc_scheda and not pd.isna(odc_scheda) and str(odc_scheda).strip():
                        try:
                            odc_scheda_norm_ricerca = _estrai_odc_normalizzato_da_sorgente(odc_scheda) 
                            if odc_scheda_norm_ricerca:
                                for _, row_s_odc in df_sorgente_dati.iterrows():
                                    odc_sorgente_val = row_s_odc.iloc[COL_IDX_COMP_ODC]
                                    extracted_odc_sorgente_norm = _estrai_odc_normalizzato_da_sorgente(odc_sorgente_val)
                                    if extracted_odc_sorgente_norm and extracted_odc_sorgente_norm == odc_scheda_norm_ricerca:
                                        riga_sorgente_trovata_series = row_s_odc
                                        log_to_all(log_text_widget, f"[{nome_file_scheda_log_corrente}] Trovata riga sorgente per ODC esatto '{odc_scheda_norm_ricerca}'.", "DEBUG")
                                        break 
                                if riga_sorgente_trovata_series is None: 
                                    odc_scheda_str_ricerca = str(odc_scheda).replace("/","-").strip() # Usa ODC originale (normalizzato per /) per 'contains'
                                    col_odc_sorg_str = df_sorgente_dati.iloc[:, COL_IDX_COMP_ODC].astype(str).str.strip()
                                    corr_odc_df = df_sorgente_dati[col_odc_sorg_str.str.contains(re.escape(odc_scheda_str_ricerca), case=False, na=False, regex=True)]
                                    if not corr_odc_df.empty:
                                         if len(corr_odc_df) > 1 and COL_IDX_COMP_DATA < df_sorgente_dati.shape[1] :
                                            corr_odc_df_copy = corr_odc_df.copy(); corr_odc_df_copy.loc[:, 'parsed_date_sorgente'] = pd.to_datetime(corr_odc_df_copy.iloc[:, COL_IDX_COMP_DATA], errors='coerce')
                                            corr_odc_df_copy.dropna(subset=['parsed_date_sorgente'], inplace=True)
                                            if not corr_odc_df_copy.empty: riga_sorgente_trovata_series = corr_odc_df_copy.loc[corr_odc_df_copy['parsed_date_sorgente'].idxmax()]
                                            else: riga_sorgente_trovata_series = corr_odc_df.iloc[0]
                                         else: riga_sorgente_trovata_series = corr_odc_df.iloc[0]
                                         log_to_all(log_text_widget, f"[{nome_file_scheda_log_corrente}] Trovata riga sorgente per ODC (contains) '{odc_scheda_str_ricerca}'.", "DEBUG")
                        except Exception as e_odc_s: log_to_all(log_text_widget, f"[{nome_file_scheda_log_corrente}] Errore ricerca ODC: {e_odc_s}", "WARNING")


                    if riga_sorgente_trovata_series is not None:
                        try:
                            map_key_to_field_col = {
                                (KEY_COMP_ANA_DATA_COMP_MANCANTE if tipo_file_scheda == 'analogico' else KEY_COMP_DIG_DATA_COMP_MANCANTE): ('DATA_COMPILAZIONE', COL_IDX_COMP_DATA),
                                (KEY_COMP_ANA_ESECUTORE_MANCANTE if tipo_file_scheda == 'analogico' else KEY_COMP_DIG_ESECUTORE_MANCANTE): ('ESECUTORE', COL_IDX_COMP_ESECUTORE),
                                (KEY_COMP_ANA_SUPERVISORE_MANCANTE if tipo_file_scheda == 'analogico' else KEY_COMP_DIG_SUPERVISORE_MANCANTE): ('SUPERVISORE_ISAB', COL_IDX_COMP_SUPERVISORE),
                                (KEY_COMP_ANA_ODC_MANCANTE if tipo_file_scheda == 'analogico' else KEY_COMP_DIG_ODC_MANCANTE): ('ODC', COL_IDX_COMP_ODC),
                                (KEY_COMP_ANA_PDL_MANCANTE if tipo_file_scheda == 'analogico' else KEY_COMP_DIG_PDL_MANCANTE): ('PDL', COL_IDX_COMP_PDL),
                            }
                            for key_errore_campo, (nome_campo_logico, idx_col_sorgente) in map_key_to_field_col.items():
                                if key_errore_campo in campi_da_compilare_su_scheda: 
                                    valore_sorgente = riga_sorgente_trovata_series.iloc[idx_col_sorgente]
                                    coord_target = _get_cell_coordinate_for_compilation(nome_campo_logico, tipo_file_scheda)
                                    if coord_target:
                                        if nome_campo_logico == 'ODC': 
                                            valore_sorgente_normalizzato = _estrai_odc_normalizzato_da_sorgente(valore_sorgente)
                                            if valore_sorgente_normalizzato:
                                                dati_effettivi_da_scrivere[coord_target] = valore_sorgente_normalizzato
                                            else:
                                                log_to_all(log_text_widget, f"[{nome_file_scheda_log_corrente}] ODC sorgente '{valore_sorgente}' per {coord_target} non normalizzabile, non scritto.", "WARNING")
                                        else:
                                            dati_effettivi_da_scrivere[coord_target] = valore_sorgente
                        except IndexError: log_to_all(log_text_widget, f"[{nome_file_scheda_log_corrente}] Errore indice estrazione dati sorgente.", "ERROR")
                        except Exception as e_extr_src: log_to_all(log_text_widget, f"[{nome_file_scheda_log_corrente}] Errore estrazione dati specifici: {e_extr_src}", "WARNING")
                
                file_obiettivo_per_scrittura = file_da_modificare_originale_xls_o_xlsx # Inizializza con il percorso originale
                nome_file_log_finale = nome_file_scheda_log_corrente


                if dati_effettivi_da_scrivere:
                    file_scheda_originale_ext = os.path.splitext(file_da_modificare_originale_xls_o_xlsx)[1].lower()
                    
                    if file_scheda_originale_ext == '.xls':
                        path_master_da_usare = None
                        if tipo_file_scheda == 'digitale' and FILE_MASTER_DIGITALE_XLSX:
                            path_master_da_usare = FILE_MASTER_DIGITALE_XLSX
                        elif tipo_file_scheda == 'analogico' and FILE_MASTER_ANALOGICO_XLSX:
                            path_master_da_usare = FILE_MASTER_ANALOGICO_XLSX
                        
                        if path_master_da_usare:
                            log_to_all(log_text_widget, f"[{nome_file_log_finale}] File .xls ('{file_da_modificare_originale_xls_o_xlsx}') necessita scrittura. Uso master: '{path_master_da_usare}'.", "INFO")
                            base_xls_orig, _ = os.path.splitext(file_da_modificare_originale_xls_o_xlsx)
                            path_xlsx_target_da_master = base_xls_orig + ".xlsx"

                            try:
                                # 1. Copia il file master nella destinazione
                                shutil.copy2(path_master_da_usare, path_xlsx_target_da_master)
                                log_to_all(log_text_widget, f"[{nome_file_log_finale}] Creato file '{path_xlsx_target_da_master}' da master.", "INFO")

                                # 2. Leggi i dati dal file .xls originale (solo il primo foglio)
                                df_xls_content = pd.read_excel(file_da_modificare_originale_xls_o_xlsx, sheet_name=0, header=None, engine='xlrd')

                                # 3. Apri il nuovo file .xlsx (copia del master) e scrivi i dati
                                wb_target_xlsx = load_workbook(path_xlsx_target_da_master)
                                ws_target_xlsx = wb_target_xlsx.active # Assumiamo di scrivere nel primo/foglio attivo del master

                                for r_idx, xls_row_data in df_xls_content.iterrows():
                                    for c_idx, xls_cell_value in enumerate(xls_row_data):
                                        if not pd.isna(xls_cell_value): # Non scrivere NaN, preserva il contenuto del master se .xls è vuoto lì
                                            ws_target_xlsx.cell(row=r_idx + 1, column=c_idx + 1, value=xls_cell_value) # openpyxl è 1-based

                                wb_target_xlsx.save(path_xlsx_target_da_master)
                                log_to_all(log_text_widget, f"[{nome_file_log_finale}] Dati da '{file_da_modificare_originale_xls_o_xlsx}' copiati in '{path_xlsx_target_da_master}'.", "SUCCESS")
                                
                                # 4. Cancella il file .xls originale
                                try:
                                    os.remove(file_da_modificare_originale_xls_o_xlsx)
                                    log_to_all(log_text_widget, f"[{nome_file_log_finale}] File .xls originale '{file_da_modificare_originale_xls_o_xlsx}' cancellato.", "INFO")
                                except Exception as e_remove_xls_master:
                                    log_to_all(log_text_widget, f"[{nome_file_log_finale}] ERRORE cancellazione file .xls originale '{file_da_modificare_originale_xls_o_xlsx}': {e_remove_xls_master}", "WARNING")
                                
                                file_obiettivo_per_scrittura = path_xlsx_target_da_master
                                nome_file_log_finale = os.path.basename(path_xlsx_target_da_master)
                                
                                # Aggiorna lista globale
                                for idx_lista, item_lista in enumerate(schede_analizzate_info_list):
                                    if item_lista["file_path"] == file_da_modificare_originale_xls_o_xlsx:
                                        schede_analizzate_info_list[idx_lista]["file_path"] = path_xlsx_target_da_master
                                        schede_analizzate_info_list[idx_lista]["base_filename"] = nome_file_log_finale
                                        break
                                log_to_all(log_text_widget, f"[{nome_file_scheda_log_corrente} -> {nome_file_log_finale}] Info file aggiornate in memoria post-conversione via master.", "DEBUG")

                            except FileNotFoundError as e_fnf_master:
                                log_to_all(log_text_widget, f"[{nome_file_log_finale}] ERRORE: File master '{path_master_da_usare}' o file sorgente .xls '{file_da_modificare_originale_xls_o_xlsx}' non trovato durante conversione: {e_fnf_master}. Scrittura saltata.", "ERROR")
                                continue
                            except Exception as e_convert_master:
                                log_to_all(log_text_widget, f"[{nome_file_log_finale}] ERRORE durante conversione via master per '{file_da_modificare_originale_xls_o_xlsx}': {e_convert_master}. Scrittura saltata.", "ERROR", exc_info=True)
                                continue
                        else:
                            log_to_all(log_text_widget, f"[{nome_file_log_finale}] File .xls ('{file_da_modificare_originale_xls_o_xlsx}') necessita scrittura ma il file master .xlsx appropriato (B5/B6 in parametri) non è configurato o non valido. Scrittura saltata.", "ERROR")
                            continue # Salta scrittura per questo file

                    # Procedi con la scrittura dei campi anagrafici sul file .xlsx (originale o convertito via master)
                    try:
                        wb_scheda_finale = load_workbook(file_obiettivo_per_scrittura)
                        ws_scheda_finale = wb_scheda_finale.active 

                        for cella_coord_str_final, valore_da_scrivere_final_val in dati_effettivi_da_scrivere.items():
                            if cella_coord_str_final and valore_da_scrivere_final_val is not None and not (isinstance(valore_da_scrivere_final_val, float) and pd.isna(valore_da_scrivere_final_val)):
                                valore_pulito_final_val = str(valore_da_scrivere_final_val).strip()
                                
                                is_contratto_cell_final = (cella_coord_str_final == _get_cell_coordinate_for_compilation('CONTRATTO_COEMI', tipo_file_scheda))
                                if not valore_pulito_final_val and not is_contratto_cell_final : 
                                    log_to_all(log_text_widget, f"[{nome_file_log_finale}] Valore vuoto per {cella_coord_str_final}, non scritto.", "DEBUG")
                                    continue
                                elif not valore_pulito_final_val and is_contratto_cell_final and valore_da_scrivere_final_val != VALORE_ATTESO_CONTRATTO_COEMI:
                                    log_to_all(log_text_widget, f"[{nome_file_log_finale}] Valore vuoto per Contratto {cella_coord_str_final} e non è valore atteso. Non scritto.", "DEBUG")
                                    continue

                                if cella_coord_str_final == _get_cell_coordinate_for_compilation('DATA_COMPILAZIONE', tipo_file_scheda):
                                    data_dt_val_final = parse_date_robust(valore_pulito_final_val, context_filename=os.path.basename(FILE_DATI_COMPILAZIONE_SCHEDE if FILE_DATI_COMPILAZIONE_SCHEDE else "FileSorgenteN/D"))
                                    if data_dt_val_final:
                                        ws_scheda_finale[cella_coord_str_final] = data_dt_val_final 
                                        log_to_all(log_text_widget, f"[{nome_file_log_finale}] Scritto DATA '{data_dt_val_final.strftime('%d/%m/%Y')}' in {cella_coord_str_final}.", "INFO")
                                        scheda_effettivamente_modificata_in_questo_ciclo = True
                                    else:
                                        log_to_all(log_text_widget, f"[{nome_file_log_finale}] Data '{valore_pulito_final_val}' non valida per {cella_coord_str_final}. Non scritta.", "WARNING")
                                else: 
                                    ws_scheda_finale[cella_coord_str_final] = valore_pulito_final_val
                                    log_to_all(log_text_widget, f"[{nome_file_log_finale}] Scritto '{valore_pulito_final_val}' in {cella_coord_str_final}.", "INFO")
                                    scheda_effettivamente_modificata_in_questo_ciclo = True
                        
                        if scheda_effettivamente_modificata_in_questo_ciclo:
                            wb_scheda_finale.save(file_obiettivo_per_scrittura)
                            modifiche_effettuate_conteggio +=1
                            log_to_all(log_text_widget, f"Scheda '{nome_file_log_finale}' salvata con aggiornamenti anagrafici.", "SUCCESS")

                    except FileNotFoundError:
                        log_to_all(log_text_widget, f"File scheda '{file_obiettivo_per_scrittura}' non trovato durante tentativo di scrittura campi anagrafici.", "ERROR")
                    except Exception as e_write_anagrafica:
                        log_to_all(log_text_widget, f"Errore scrittura campi anagrafici su scheda '{nome_file_log_finale}': {e_write_anagrafica}", "ERROR", exc_info=True)
                elif campi_che_richiedono_ricerca_sorgente and riga_sorgente_trovata_series is None and not dati_effettivi_da_scrivere:
                    log_to_all(log_text_widget, f"Nessun dato di compilazione trovato per '{nome_file_log_finale}' e nessun altro campo da scrivere.", "WARNING")
            
            if modifiche_effettuate_conteggio > 0:
                messagebox.showinfo("Compilazione Completata", f"{modifiche_effettuate_conteggio} schede sono state potenzialmente aggiornate.\nControllare il log per i dettagli.", parent=result_window_ref_store.get('ref'))
            else:
                messagebox.showinfo("Compilazione Completata", "Nessuna scheda è stata modificata. Controllare il log per i dettagli (es. nessun campo mancante idoneo, problemi con file master, o dati sorgente non trovati).", parent=result_window_ref_store.get('ref'))
            log_to_all(log_text_widget, "Processo di compilazione automatica terminato.", "INFO")

        except FileNotFoundError: 
            msg = f"File sorgente dati compilazione '{FILE_DATI_COMPILAZIONE_SCHEDE}' non trovato."
            log_to_all(log_text_widget, msg, "ERROR")
            messagebox.showerror("Errore File Sorgente", msg, parent=result_window_ref_store.get('ref'))
        except ValueError as ve: 
            msg = f"Errore durante la lettura del file sorgente dati '{FILE_DATI_COMPILAZIONE_SCHEDE}': {ve}. Assicurarsi che il foglio '{NOME_FOGLIO_DATI_COMPILAZIONE}' esista e che gli header siano nella prima riga."
            log_to_all(log_text_widget, msg, "ERROR")
            messagebox.showerror("Errore Lettura File Sorgente", msg, parent=result_window_ref_store.get('ref'))
        except Exception as e_comp_glob_final:
            log_to_all(log_text_widget, f"Errore imprevisto durante la compilazione automatica: {e_comp_glob_final}", "ERROR", exc_info=True)
            messagebox.showerror("Errore Imprevisto Compilazione", f"Si è verificato un errore imprevisto.\n{e_comp_glob_final}", parent=result_window_ref_store.get('ref'))


    btn_compila_schede = ttk.Button(compilatore_action_frame, text="Compila Schede Automaticamente",
                                     command=esegui_compilazione_schede,
                                     style="Accent.TButton")
    btn_compila_schede.pack(side=tk.LEFT, padx=5)
    if not FILE_DATI_COMPILAZIONE_SCHEDE: 
        btn_compila_schede.config(state=tk.DISABLED)
        info_no_comp_file = ttk.Label(compilatore_action_frame, text="File dati compilazione (B4) non specificato in parametri.xlsm.", foreground="orange")
        info_no_comp_file.pack(side=tk.LEFT, padx=10)


    comp_addizionali_lf, comp_addizionali_inner = create_scrollable_section_frame(compilatore_main_frame, "Anomalie Compilazione Scheda (Campi ODC, Data, PDL, etc.)")
    comp_addizionali_lf.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)

    errori_compilazione_anagrafici = [err for err in human_errors_details_list if err['key'].startswith("COMP_")]
    if errori_compilazione_anagrafici:
        errori_compilazione_anagrafici.sort(key=lambda x: (x['file'], x['key']))
        for idx, error_detail_comp in enumerate(errori_compilazione_anagrafici):
            err_entry_frame_comp = ttk.Frame(comp_addizionali_inner)
            err_entry_frame_comp.pack(fill=tk.X, pady=(2,0), padx=5)
            err_entry_frame_comp.columnconfigure(1, weight=1)

            file_label_comp_err = ttk.Label(err_entry_frame_comp, text=f"{error_detail_comp['file']}:", style="Hyperlink.TLabel", cursor="hand2", font=('Segoe UI', 9, 'bold'))
            file_label_comp_err.grid(row=0, column=0, sticky=tk.NW, padx=(0,5))
            _fp_comp_err_tab = error_detail_comp.get('path'); _fn_comp_err_tab = error_detail_comp.get('file',"file")
            file_label_comp_err.bind("<Button-1>", lambda event, p=_fp_comp_err_tab, fn=_fn_comp_err_tab: on_file_click(p, fn, open_file_direct=True, event=event))

            error_message_desc_comp_tab = human_error_messages_map_descriptive.get(error_detail_comp['key'], f"Errore sconosciuto: {error_detail_comp['key']}")
            
            msg_label_comp_tab = tk.Label(err_entry_frame_comp, text=error_message_desc_comp_tab, fg="maroon", font=('Segoe UI', 9),
                                     anchor=tk.W, justify=tk.LEFT,
                                     background=style.lookup('TFrame','background'))
            msg_label_comp_tab.grid(row=0, column=1, sticky=tk.EW, padx=(5,0))
            msg_label_comp_tab.after(50, lambda w=msg_label_comp_tab, p=err_entry_frame_comp, fl=file_label_comp_err: update_label_wraplength(w, p, fl)) 

            if idx < len(errori_compilazione_anagrafici) - 1:
                ttk.Separator(comp_addizionali_inner, orient='horizontal').pack(fill='x',pady=(5,3),padx=10)
    else:
        tk.Label(comp_addizionali_inner, text="Nessuna anomalia di compilazione (campi ODC, Data, PDL, etc.) rilevata.", fg="darkgreen", font=('Segoe UI', 10, 'bold'), anchor=tk.W, background=style.lookup('TFrame','background')).pack(anchor=tk.W, padx=5, pady=5)
    log_to_all(log_text_widget,"Tab Compilatore creata.","DEBUG")
    # --- Fine Tab Compilatore ---


    # --- Tab Suggerimenti Strumenti Alternativi ---
    tab_suggerimenti_generale = ttk.Frame(notebook, padding=10)
    notebook.add(tab_suggerimenti_generale, text=' Suggerimenti Strumenti Alternativi ')
    sugg_main_frame = ttk.Frame(tab_suggerimenti_generale); sugg_main_frame.pack(expand=True, fill='both')
    sugg_input_frame = ttk.LabelFrame(sugg_main_frame, text="Parametri Ricerca", padding=10); sugg_input_frame.pack(fill=tk.X, pady=5)
    cert_id_sugg_entry_local = ttk.Entry(sugg_input_frame, width=25)
    range_sugg_entry_local = ttk.Entry(sugg_input_frame, width=30)
    date_sugg_entry_local = ttk.Entry(sugg_input_frame, width=15); date_sugg_entry_local.insert(0, ANALYSIS_DATETIME.astimezone().strftime('%d/%m/%Y'))
    ttk.Label(sugg_input_frame, text="ID Cert. (opz.):", font=('Segoe UI',10)).grid(row=0, column=0, padx=(0,5), pady=5, sticky=tk.W)
    cert_id_sugg_entry_local.grid(row=0, column=1, padx=5, pady=5, sticky=tk.EW)
    ttk.Label(sugg_input_frame, text="Range Richiesto:", font=('Segoe UI',10)).grid(row=0, column=2, padx=(10,5), pady=5, sticky=tk.W)
    range_sugg_entry_local.grid(row=0, column=3, padx=5, pady=5, sticky=tk.EW)
    ttk.Label(sugg_input_frame, text="Data Rif. (gg/mm/aaaa):", font=('Segoe UI',10)).grid(row=1, column=0, padx=(0,5), pady=5, sticky=tk.W)
    date_sugg_entry_local.grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)
    sugg_input_frame.columnconfigure(1, weight=1); sugg_input_frame.columnconfigure(3, weight=1) 
    
    sugg_results_text_frame = ttk.Frame(sugg_main_frame); sugg_results_text_frame.pack(expand=True, fill='both', pady=(5,0))
    sugg_results_text_local = tk.Text(sugg_results_text_frame, wrap=tk.WORD, height=15, font=("Consolas",10), relief=tk.SUNKEN, borderwidth=1)
    sugg_vsb_res = ttk.Scrollbar(sugg_results_text_frame, orient="vertical", command=sugg_results_text_local.yview)
    sugg_results_text_local.configure(yscrollcommand=sugg_vsb_res.set); sugg_vsb_res.pack(side=tk.RIGHT, fill=tk.Y); sugg_results_text_local.pack(side=tk.LEFT, expand=True, fill='both')
    sugg_results_text_local.config(state=tk.DISABLED) 
    sugg_results_text_local.tag_configure("sugg_header", font=('Segoe UI',11,'bold'), spacing1=5, spacing3=10, foreground="navy")
    sugg_results_text_local.tag_configure("sugg_item_main", font=('Consolas',10,'bold'), lmargin1=15, lmargin2=15, foreground="darkgreen")
    sugg_results_text_local.tag_configure("sugg_item_detail", font=('Consolas',10), lmargin1=30, lmargin2=30, spacing1=3)
    sugg_results_text_local.tag_configure("sugg_error", foreground="red", font=('Segoe UI',10,'bold'))
    sugg_results_text_local.tag_configure("sugg_warning", foreground="darkorange", font=('Segoe UI',10))
    sugg_results_text_local.tag_configure("sugg_info", foreground="blue", font=('Segoe UI',10))

    def cerca_e_mostra_suggerimenti_generale_impl_local():
        cert_id_target = cert_id_sugg_entry_local.get().strip()
        range_target_raw = range_sugg_entry_local.get().strip()
        date_ref_str = date_sugg_entry_local.get().strip()
        sugg_text = sugg_results_text_local; sugg_text.config(state=tk.NORMAL); sugg_text.delete("1.0", tk.END)

        if not cert_id_target and not range_target_raw:
            sugg_text.insert(tk.END, "Inserire ID Certificato (per derivare il range) o un Range specifico.\n", "sugg_error")
            sugg_text.config(state=tk.DISABLED); return
        if not strumenti_campione_list: 
            sugg_text.insert(tk.END, "Registro strumenti campione non caricato o vuoto. Impossibile cercare alternative.\n", "sugg_warning")
            sugg_text.config(state=tk.DISABLED); return
        
        date_riferimento_dt = parse_date_robust(date_ref_str, "UI Suggerimenti")
        if not date_riferimento_dt:
            sugg_text.insert(tk.END, f"Data di riferimento '{date_ref_str}' non valida. Usare formato gg/mm/aaaa.\n", "sugg_error")
            sugg_text.config(state=tk.DISABLED); return

        final_range_search = range_target_raw
        if not final_range_search and cert_id_target: 
            if cert_id_target in cert_details_map: 
                range_info_counter = cert_details_map[cert_id_target]['range_su_scheda_counter']
                if range_info_counter:
                    final_range_search = range_info_counter.most_common(1)[0][0]
                    if final_range_search != "N/D":
                        sugg_text.insert(tk.END, f"Range più comune per ID '{cert_id_target}': '{final_range_search}'. Uso questo per la ricerca.\n\n", "sugg_info")
                    else: 
                        sugg_text.insert(tk.END, f"Nessun range valido ('N/D' escluso) trovato per ID Cert. '{cert_id_target}'. Specificare un Range manualmente.\n", "sugg_warning")
                        final_range_search = "" 
                else:
                    sugg_text.insert(tk.END, f"Nessun range comune trovato per ID Cert. '{cert_id_target}'. Specificare un Range manualmente.\n", "sugg_warning")
            else:
                sugg_text.insert(tk.END, f"ID Cert. '{cert_id_target}' non trovato nei dati analizzati per derivare il range. Specificare Range.\n", "sugg_warning")
        
        if not final_range_search: 
            sugg_text.insert(tk.END, "Range non specificato e non derivabile. Inserire un Range per la ricerca.\n", "sugg_error")
            sugg_text.config(state=tk.DISABLED); return
        
        alternative = trova_strumenti_alternativi(final_range_search, date_riferimento_dt, strumenti_campione_list)
        sugg_text.insert(tk.END, f"Suggerimenti per Range '{final_range_search}' validi al {date_riferimento_dt:%d/%m/%Y}:\n(Esclude strumento con ID '{cert_id_target}' se fornito e trovato tra le alternative)\n", "sugg_header")
        if alternative:
            count_s = 0
            for strum in alternative:
                if strum.get('id_certificato') == cert_id_target and cert_id_target: continue 
                count_s += 1
                scad_str_alt = strum['scadenza'].strftime('%d/%m/%Y') if strum.get('scadenza') else "N/D"
                emiss_str_alt = strum.get('data_emissione').strftime('%d/%m/%Y') if strum.get('data_emissione') else "N/D"
                sugg_text.insert(tk.END, f"\n{count_s}. Modello: ", "sugg_item_detail"); sugg_text.insert(tk.END, f"{strum.get('modello_strumento','N/D')}", "sugg_item_main")
                sugg_text.insert(tk.END, f" (Cert: ", "sugg_item_detail"); sugg_text.insert(tk.END, f"{strum.get('id_certificato','N/D')}", "sugg_item_main"); sugg_text.insert(tk.END, ")\n", "sugg_item_detail")
                sugg_text.insert(tk.END, f"\tRange Campione: {strum.get('range','N/D')}\n", "sugg_item_detail")
                sugg_text.insert(tk.END, f"\tEmissione: {emiss_str_alt}\n", "sugg_item_detail"); sugg_text.insert(tk.END, f"\tScadenza: {scad_str_alt}\n", "sugg_item_detail")
            if count_s == 0: sugg_text.insert(tk.END, f"\nNessun alternativo (diverso da '{cert_id_target if cert_id_target else 'N/A'}') trovato per i criteri specificati.\n", "sugg_warning")
        else:
            sugg_text.insert(tk.END, f"\nNessuno strumento alternativo trovato per i criteri specificati.\n", "sugg_warning")
        sugg_text.config(state=tk.DISABLED)

    btn_cerca_sugg = ttk.Button(sugg_input_frame, text="Cerca Alternative", command=cerca_e_mostra_suggerimenti_generale_impl_local, style="Accent.TButton")
    btn_cerca_sugg.grid(row=1, column=2, columnspan=2, padx=10, pady=5, sticky=tk.E) 
    log_to_all(log_text_widget, "Tab Suggerimenti creata.", "DEBUG")
    sugg_tab_interface_details['suggerimenti_tab_widget'] = tab_suggerimenti_generale
    sugg_tab_interface_details['cert_id_entry_widget'] = cert_id_sugg_entry_local
    sugg_tab_interface_details['range_entry_widget'] = range_sugg_entry_local
    sugg_tab_interface_details['date_entry_widget'] = date_sugg_entry_local
    sugg_tab_interface_details['search_function_widget'] = cerca_e_mostra_suggerimenti_generale_impl_local
    sugg_tab_interface_details['results_text_widget'] = sugg_results_text_local
    
    log_to_all(log_text_widget, "Tab Grafici non implementata in questa versione.", "INFO") 
    
    def close_app_confirmed():
        nonlocal root, progress_window 
        current_result_window = result_window_ref_store.get('ref')
        confirm_close = True
        parent_for_confirm = current_result_window if current_result_window and current_result_window.winfo_exists() else root
        
        if progress_window.winfo_exists():
            if not messagebox.askokcancel("Chiudi Tutto", "Finestra Log/Progresso ancora aperta.\nChiudere l'applicazione?", parent=parent_for_confirm, icon='question'):
                confirm_close = False
        
        if confirm_close:
            log_to_all(log_text_widget, "--- APPLICAZIONE IN CHIUSURA ---", "INFO")
            if progress_window.winfo_exists(): progress_window.destroy()
            if current_result_window and current_result_window.winfo_exists(): current_result_window.destroy()
            if root.winfo_exists(): root.quit(); root.destroy() 

    current_result_window_for_btn = result_window_ref_store.get('ref')
    close_button_parent_frame = current_result_window_for_btn if current_result_window_for_btn and current_result_window_for_btn.winfo_exists() else root
    
    main_app_bottom_frame = ttk.Frame(close_button_parent_frame) 
    main_app_bottom_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(0,5)) 
    close_button_main = ttk.Button(main_app_bottom_frame, text="Chiudi Applicazione", command=close_app_confirmed, style="Accent.TButton")
    close_button_main.pack(pady=(5,5)) 

    if current_result_window_for_btn and current_result_window_for_btn.winfo_exists():
        current_result_window_for_btn.protocol("WM_DELETE_WINDOW", close_app_confirmed)

    log_to_all(log_text_widget, "Finestra risultati pronta. Analisi completa.", "SUCCESS")
    
    final_res_win = result_window_ref_store.get('ref')
    if progress_window.winfo_exists() and final_res_win and final_res_win.winfo_exists():
        if progress_window.state() == 'iconic': progress_window.deiconify() 
        final_res_win.lift() 
        progress_window.lift() 
        
    root.mainloop() 


if __name__ == "__main__":
    main_logger = logging.getLogger('SchedeAnalyzer') 
    
    try:
        main_logger.info(f"App Analisi Schede avviata. Python: {sys.version.split()[0]}, OS: {sys.platform}, Log: {LOG_FILEPATH if LOG_FILEPATH else 'NON DISP.'}, Data: {datetime.now():%d/%m/%Y %H:%M:%S}")
        run_analysis_and_show_results()

    except (FileNotFoundError, ValueError, RuntimeError, NotADirectoryError) as e_fatal_config: # Aggiunto NotADirectoryError
        if main_logger.hasHandlers() and any(isinstance(h, logging.FileHandler) for h in main_logger.handlers):
            main_logger.critical(f"Errore FATALE durante l'inizializzazione o configurazione: {e_fatal_config}", exc_info=True)
        else: 
            print(f"ERRORE FATALE (log su file non pronto): {e_fatal_config}", file=sys.stderr)
            import traceback
            traceback.print_exc(file=sys.stderr)

        try:
            err_root = tk.Tk(); err_root.withdraw() 
            messagebox.showerror("Errore Fatale Applicazione",
                                 f"Si è verificato un errore critico durante l'inizializzazione:\n{e_fatal_config}\n\n"
                                 f"L'applicazione sarà chiusa.\nControllare la console e il log (se disponibile):\n{LOG_FILEPATH if LOG_FILEPATH else 'Output console'}",
                                 parent=None) 
            if err_root.winfo_exists(): err_root.destroy()
        except Exception as e_msg_box_config:
            if main_logger.hasHandlers(): 
                main_logger.error(f"Impossibile mostrare messagebox per errore fatale di configurazione: {e_msg_box_config}")
            print(f"ERRORE FATALE AGGIUNTIVO (config): Impossibile mostrare messagebox: {e_msg_box_config}", file=sys.stderr)
            
    except Exception as e_glob: 
        if main_logger.hasHandlers():
            main_logger.critical(f"Errore FATALE non gestito nell'applicazione: {e_glob}", exc_info=True)
        else:
            print(f"ERRORE FATALE (logger non pronto): {e_glob}", file=sys.stderr)
            import traceback
            traceback.print_exc(file=sys.stderr)
        try:
            err_root_glob = tk.Tk(); err_root_glob.withdraw()
            messagebox.showerror("Errore Fatale Applicazione",
                                 f"Si è verificato un errore critico non gestito:\n{e_glob}\n\n"
                                 f"L'applicazione sarà chiusa.\nControllare log:\n{LOG_FILEPATH if LOG_FILEPATH else 'Output console'}",
                                 parent=None)
            if err_root_glob.winfo_exists(): err_root_glob.destroy()
        except Exception as e_msg_glob:
            if main_logger.hasHandlers():
                main_logger.error(f"Impossibile mostrare messagebox per errore fatale globale: {e_msg_glob}")
            print(f"ERRORE FATALE AGGIUNTIVO (globale): Impossibile mostrare messagebox: {e_msg_glob}", file=sys.stderr)
    finally:
        if main_logger.hasHandlers():
            main_logger.info(f"Applicazione Analisi Schede terminata. Data: {datetime.now():%d/%m/%Y %H:%M:%S}")
        else: 
            print(f"Applicazione Analisi Schede terminata. Data: {datetime.now():%d/%m/%Y %H:%M:%S}")

        if logging.getLogger().hasHandlers(): 
            logging.shutdown()
# --- FINE PARTE 5 E FINE SCRIPT ---