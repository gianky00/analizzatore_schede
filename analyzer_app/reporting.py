import os
import sys
import subprocess
import tempfile
import logging
from typing import List, Dict
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import config

logger = logging.getLogger(__name__)

def crea_e_apri_report_anomalie_word(
    errors_list: List[Dict],
    temporal_list: List[Dict],
    incongruent_list: List[Dict],
    candidate_files_count: int,
    validated_file_count: int
) -> str:
    """
    Crea un report Word con le anomalie, lo salva in una cartella temporanea,
    tenta di aprirlo e restituisce il percorso del file.
    Restituisce None se non ci sono anomalie o in caso di errore.
    """
    logger.info("Inizio creazione report anomalie Word.")

    anomalie_da_riportare = bool(errors_list or any(item['alert_type'] == 'premature_emission' for item in temporal_list) or any(item['alert_type'] == 'expired_at_use' for item in temporal_list) or incongruent_list)

    if not anomalie_da_riportare:
        logger.info("Nessuna anomalia significativa da riportare nel file Word.")
        return None

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    title = doc.add_heading("Report Anomalie Analisi Schede Taratura", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().add_run(f"Data Analisi: {config.ANALYSIS_DATETIME.astimezone().strftime('%d/%m/%Y %H:%M:%S %Z')}").italic = True
    p_folder = doc.add_paragraph()
    p_folder.add_run("Cartella Analizzata: ").italic = True
    p_folder.add_run(f"{config.FOLDER_PATH_DEFAULT if config.FOLDER_PATH_DEFAULT else 'N/D'}").bold = True
    p_folder.add_run(f" (File Candidati: {candidate_files_count}, Schede Validate: {validated_file_count})").italic = True
    doc.add_paragraph()

    # Sezione 1: Errori di Compilazione Scheda (Strutturali)
    errori_comp_scheda_strutturali = [err for err in errors_list if not err['key'].startswith("COMP_")]
    if errori_comp_scheda_strutturali:
        doc.add_heading("1. Errori di Compilazione Scheda (Strutturali: Range/UM, SP, L9)", level=1)
        errori_comp_scheda_strutturali.sort(key=lambda x: (x['file'], x['key']))
        current_file_comp_err = None
        for error in errori_comp_scheda_strutturali:
            if error['file'] != current_file_comp_err:
                if current_file_comp_err is not None: doc.add_paragraph()
                p_file = doc.add_paragraph()
                p_file.add_run("File: ").bold = True
                p_file.add_run(f"{error['file']} (Percorso: {error['path']})")
                current_file_comp_err = error['file']

            error_desc = config.human_error_messages_map_descriptive.get(error['key'], f"Codice Errore Sconosciuto: {error['key']}")
            doc.add_paragraph(f"  • {error_desc}", style='ListBullet')
        doc.add_paragraph()

    # Sezione 2: Certificati Utilizzati Prima dell'Emissione
    premature_uses = [item for item in temporal_list if item['alert_type'] == 'premature_emission']
    if premature_uses:
        doc.add_heading("2. Certificati Campione Utilizzati Prima della Loro Data di Emissione", level=1)
        premature_uses.sort(key=lambda x: (x.get('file_name', ''), x.get('card_date', datetime.min)))
        for item in premature_uses:
            p = doc.add_paragraph()
            p.add_run("File Scheda: ").bold = True
            p.add_run(f"{item['file_name']} ({item['file_path']})")
            doc.add_paragraph(f"  • Data Scheda: {item['card_date_str']}", style='ListBullet')
            doc.add_paragraph(f"  • Certificato ID: {item['certificate_id']}", style='ListBullet')
            p_cert_detail = doc.add_paragraph(f"  • CERTIFICATO CAMPIONE (da Registro): {item['modello_strumento_campione_usato']}", style='ListBullet')
            p_cert_detail_em = doc.add_paragraph(f"  • Data Emissione Certificato Campione: {item.get('data_emissione_presunta','N/A')}", style='ListBullet')
            run_em_note = p_cert_detail_em.add_run(" - USATO PRIMA DELL'EMISSIONE!")
            run_em_note.bold = True; run_em_note.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            doc.add_paragraph()
        doc.add_paragraph()

    # Sezione 3: Certificati Scaduti Utilizzati
    expired_at_use_pure = [item for item in temporal_list if item['alert_type'] == 'expired_at_use']
    if expired_at_use_pure:
        doc.add_heading("3. Certificati Campione Scaduti al Momento dell'Uso", level=1)
        expired_at_use_pure.sort(key=lambda x: (x.get('file_name', ''), x.get('card_date', datetime.min)))
        for item in expired_at_use_pure:
            p = doc.add_paragraph()
            p.add_run("File Scheda: ").bold = True
            p.add_run(f"{item['file_name']} ({item['file_path']})")
            doc.add_paragraph(f"  • Data Scheda: {item['card_date_str']}", style='ListBullet')
            doc.add_paragraph(f"  • Certificato ID: {item['certificate_id']}", style='ListBullet')
            p_cert_detail_exp = doc.add_paragraph(f"  • CERTIFICATO CAMPIONE (da Registro): {item['modello_strumento_campione_usato']}", style='ListBullet')
            p_cert_detail_scad = doc.add_paragraph(f"  • Data Scadenza Certificato Campione: {item.get('expiry_date_str','N/P')}", style='ListBullet')
            run_scad_note = p_cert_detail_scad.add_run(" - SCADUTO ALL'USO!")
            run_scad_note.bold = True; run_scad_note.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
            doc.add_paragraph(f"  • Note Congruità (se presenti): {item['congruency_notes']}", style='ListBullet')
            doc.add_paragraph()
        doc.add_paragraph()

    # Sezione 4: Certificati Non Congrui per Regole
    if incongruent_list:
        doc.add_heading("4. Certificati Campione Non Congrui (per Regole Tipologia/Modello)", level=1)
        incongruent_list.sort(key=lambda x: (x.get('file_name', ''), x.get('card_date', datetime.min)))
        for item in incongruent_list:
            p = doc.add_paragraph()
            p.add_run("File Scheda: ").bold = True
            p.add_run(f"{item['file_name']} ({item['file_path']})")
            p_reason = doc.add_paragraph(f"  • Motivo Non Congruità: ", style='ListBullet')
            run_reason = p_reason.add_run(item.get('congruency_notes','Non specificato'))
            run_reason.bold = True; run_reason.font.color.rgb = RGBColor(0x80, 0x00, 0x80)
            doc.add_paragraph()
        doc.add_paragraph()

    # Sezione 5: Errori di Compilazione Campi Anagrafici
    errori_comp_campi_scheda = [err for err in errors_list if err['key'].startswith("COMP_")]
    if errori_comp_campi_scheda:
        doc.add_heading("5. Errori/Mancanze nei Campi Anagrafici Scheda (ODC, Data, PDL, etc.)", level=1)
        errori_comp_campi_scheda.sort(key=lambda x: (x['file'], x['key']))
        current_file_comp_campi_err = None
        for error in errori_comp_campi_scheda:
            if error['file'] != current_file_comp_campi_err:
                if current_file_comp_campi_err is not None: doc.add_paragraph()
                p_file_campi = doc.add_paragraph()
                p_file_campi.add_run("File: ").bold = True
                p_file_campi.add_run(f"{error['file']} (Percorso: {error['path']})")
                current_file_comp_campi_err = error['file']

            error_desc_campi = config.human_error_messages_map_descriptive.get(error['key'], f"Codice Errore Sconosciuto: {error['key']}")
            doc.add_paragraph(f"  • {error_desc_campi}", style='ListBullet')
        doc.add_paragraph()

    # Salvataggio e apertura file
    try:
        temp_dir = tempfile.gettempdir()
        timestamp_rep = config.ANALYSIS_DATETIME.astimezone().strftime("%Y%m%d_%H%M%S")
        word_filename = f"Report_Anomalie_Schede_{timestamp_rep}.docx"
        word_file_path = os.path.join(temp_dir, word_filename)
        doc.save(word_file_path)
        logger.info(f"Report Word salvato in: {word_file_path}")

        if sys.platform == "win32":
            os.startfile(word_file_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", word_file_path])
        else:
            subprocess.Popen(["xdg-open", word_file_path])

        return word_file_path

    except Exception as e_word:
        logger.error(f"Errore durante la creazione o apertura del report Word: {e_word}", exc_info=True)
        return None
